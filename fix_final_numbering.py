#!/usr/bin/env python3
"""Rebuild FINAL doc body with correct list/section numbering and left alignment."""
from __future__ import annotations

import os
import re
import shutil
import sys

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph

sys.path.insert(0, os.path.dirname(__file__))
from build_marco_v2 import (  # noqa: E402
    ANNEX_MARKER,
    BODY_START_MARKER,
    SRC,
    build_v2_spec,
    delete_paragraph,
    extract_articles,
    find_body_bounds,
)

FINAL_PATH = (
    "/Users/claudiolarrea/Library/CloudStorage/OneDrive-Personal/"
    "16 Secretaría de Investigación/60 Observatorio de Inteligencia Artificial/"
    "MARCO DE GOBERNANZA DE LA IA/Documento FINAL/"
    "Marco de Gobernanza IA - UCCuyo FINAL.docx"
)

LIST_SPLIT = re.compile(
    r"\n(?=[a-z]\)\s|\d+\.\d+\.\s|\([ivxlc]+\)\s|—\s|-\s)",
    re.IGNORECASE,
)
SUBSECTION = re.compile(r"^(\d+\.\d+)\.\s*(.*)$", re.S)
LETTER_ITEM = re.compile(r"^([a-z])\)\s*(.*)$", re.S | re.IGNORECASE)


def split_body_into_paragraphs(body: str) -> list[str]:
    paragraphs: list[str] = []
    for block in re.split(r"\n\s*\n", body.strip()):
        block = block.strip()
        if not block:
            continue
        for part in LIST_SPLIT.split(block):
            part = part.strip()
            if part:
                paragraphs.append(part)
    return explode_merged_paragraphs(paragraphs)


def explode_merged_paragraphs(paragraphs: list[str]) -> list[str]:
    out: list[str] = []
    for p in paragraphs:
        if re.search(r"\n[a-z]\)\s", p, re.I):
            for part in re.split(r"\n(?=[a-z]\)\s)", p):
                part = part.strip()
                if part:
                    out.extend(explode_merged_paragraphs([part]))
            continue

        m = re.match(r"^(\d+\.\d+\.\s+[^\n]+)\n([\s\S]+)$", p)
        if m:
            out.append(m.group(1).strip())
            rest = m.group(2).strip()
            if rest:
                out.extend(explode_merged_paragraphs([rest]))
            continue

        if "\n" in p and not SUBSECTION.match(p):
            lines = [ln.strip() for ln in p.split("\n") if ln.strip()]
            if len(lines) > 1 and not any(SUBSECTION.match(ln) for ln in lines[1:]):
                out.append(lines[0])
                out.extend(explode_merged_paragraphs(["\n".join(lines[1:])]))
                continue

        out.append(p)
    return out


def renumber_subsections(paragraphs: list[str], art_num: int) -> tuple[list[str], dict[str, str]]:
    out: list[str] = []
    mapping: dict[str, str] = {}
    sub = 0
    for p in paragraphs:
        m = SUBSECTION.match(p)
        if m:
            old = m.group(1)
            sub += 1
            new = f"{art_num}.{sub}"
            mapping[old] = new
            out.append(f"{new}. {m.group(2).strip()}")
        else:
            out.append(p)
    return out, mapping


def renumber_letter_lists(paragraphs: list[str]) -> list[str]:
    """Ensure a), b), c)... sequences restart after subsections and prose blocks."""
    out: list[str] = []
    letter_idx = 0

    def reset():
        nonlocal letter_idx
        letter_idx = 0

    for p in paragraphs:
        if SUBSECTION.match(p):
            reset()
            out.append(p)
            continue

        m = LETTER_ITEM.match(p)
        if m:
            expected = chr(ord("a") + letter_idx)
            letter_idx += 1
            out.append(f"{expected}) {m.group(2).strip()}")
            continue

        if p.strip():
            reset()
        out.append(p)
    return out


def apply_subsection_reference_map(text: str, mapping: dict[str, str]) -> str:
    if not mapping:
        return text
    for old, new in sorted(mapping.items(), key=lambda x: len(x[0]), reverse=True):
        text = re.sub(rf"\bArtículo\s+{re.escape(old)}\b", f"numeral {new}", text, flags=re.I)
        text = re.sub(rf"\bel\s+{re.escape(old)}\b", f"el numeral {new}", text, flags=re.I)
        text = re.sub(rf"\bal\s+{re.escape(old)}\b", f"al numeral {new}", text, flags=re.I)
        text = re.sub(rf"\bdel\s+{re.escape(old)}\b", f"del numeral {new}", text, flags=re.I)
    return text


def add_para(doc, anchor, text, *, bold=False, heading=None):
    new_el = OxmlElement("w:p")
    anchor._p.addnext(new_el)
    p = Paragraph(new_el, anchor._parent)
    if heading == 1:
        p.style = "Heading 1"
    elif heading == 2:
        p.style = "Heading 2"
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.line_spacing = 1.5
    if text:
        run = p.add_run(text)
        run.bold = bold
        run.font.name = "Times New Roman"
    return p


def set_left_align(doc, start_marker=BODY_START_MARKER):
    in_scope = False
    for p in doc.paragraphs:
        t = p.text.strip()
        if t == start_marker and "\t" not in t:
            in_scope = True
            continue
        if not in_scope:
            continue
        if p.style.name.startswith("toc"):
            continue
        if t in ("ÍNDICE RESUMIDO", "ÍNDICE GENERAL"):
            continue
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT


def rebuild_body(doc, spec):
    start, end = find_body_bounds(doc)
    for idx in range(end - 1, start - 1, -1):
        delete_paragraph(doc.paragraphs[idx])

    anchor = doc.paragraphs[start - 1]
    note = (
        "VERSIÓN 2.0 DEPURADA — Este texto reorganiza el Marco anterior (v1) conforme a criterios de "
        "técnica normativa, reducción de redundancias y mayor operabilidad, preservando principios, "
        "garantías y anexos operativos. Las referencias internas corresponden a la numeración de esta versión."
    )
    current = add_para(doc, anchor, note, bold=True)
    current = add_para(doc, current, "")

    global_ref_map: dict[str, str] = {}

    for item in spec:
        if item["type"] == "chapter":
            current = add_para(doc, current, item["title"], bold=True, heading=1)
            current = add_para(doc, current, "")
            continue

        art_num = item["n"]
        heading = f"Artículo {art_num}°. {item['title']}"
        current = add_para(doc, current, heading, bold=True, heading=2)

        paragraphs = split_body_into_paragraphs(item["body"])
        paragraphs, ref_map = renumber_subsections(paragraphs, art_num)
        global_ref_map.update(ref_map)
        paragraphs = renumber_letter_lists(paragraphs)

        for para in paragraphs:
            para = apply_subsection_reference_map(para, ref_map)
            current = add_para(doc, current, para)
        current = add_para(doc, current, "")


def main():
    if not os.path.exists(FINAL_PATH):
        raise SystemExit(f"No existe: {FINAL_PATH}")

    backup = FINAL_PATH.replace(".docx", " - backup pre-numeracion.docx")
    shutil.copy2(FINAL_PATH, backup)

    arts = extract_articles(Document(SRC))
    spec = build_v2_spec(arts)

    doc = Document(FINAL_PATH)
    rebuild_body(doc, spec)
    set_left_align(doc)
    doc.save(FINAL_PATH)

    print(f"Backup: {backup}")
    print(f"Rebuilt body in: {FINAL_PATH}")

    # Rebuild clickable index on the fixed file
    import fix_toc_word_display as toc_fix

    toc_fix.PATHS = [FINAL_PATH]
    toc_fix.main()
    print("Index rebuilt.")


if __name__ == "__main__":
    main()
