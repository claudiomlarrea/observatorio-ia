#!/usr/bin/env python3
"""Re-expand Capítulo III (gobernanza) in Marco 2026."""
from __future__ import annotations

import re
import shutil
from datetime import date

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph

TARGET = (
    "/Users/claudiolarrea/Library/CloudStorage/OneDrive-Personal/"
    "16 Secretaría de Investigación/60 Observatorio de Inteligencia Artificial/"
    "MARCO DE GOBERNANZA DE LA IA/Documento FINAL/"
    "Marco de Gobernanza IA - UCCuyo 2026.docx"
)

INSERTIONS = [
    (
        "Artículo 12°. Órganos competentes, roles y responsabilidades",
        "La gobernanza de la IA no recae en un único órgano: se articula como un flujo de "
        "autoridad, orientación y ejecución. El Consejo Superior fija la política institucional; "
        "el Rectorado impulsa e implementa; el Observatorio orienta, evalúa y propone con "
        "independencia técnica y ética; las unidades académicas y de gestión aplican los "
        "criterios en sus ámbitos; las áreas de TIC aseguran condiciones técnicas de seguridad "
        "y trazabilidad. Ninguna de estas instancias absorbe las responsabilidades de las demás.",
    ),
    (
        "12.1. Consejo Superior",
        "El Consejo Superior ejerce la máxima dirección política del Marco. Sus decisiones "
        "tienen alcance de Universidad y no pueden ser sustituidas por lineamientos técnicos "
        "ni por prácticas locales de sede o facultad.",
    ),
    (
        "12.2. Rectorado",
        "El Rectorado es el órgano de implementación y coordinación ordinaria. Asegura que "
        "el Marco se aplique de modo transversal en todas las sedes y que las autorizaciones "
        "de sistemas institucionales de alto impacto no queden fragmentadas entre unidades.",
    ),
    (
        "12.3. Observatorio de Inteligencia Artificial de la Universidad",
        "El Observatorio no gobierna por sí mismo ni decide en lugar de las autoridades. Su "
        "función es articular, asesorar, producir evidencia y proponer criterios, de modo que "
        "las decisiones institucionales sobre IA sean informadas, trazables y coherentes con "
        "este Marco.",
    ),
    (
        "Artículo 13°. Responsabilidad humana explícita",
        "Ninguna recomendación, puntuación, alerta o contenido generado por IA desplaza la "
        "imputación de la decisión. Este artículo exige que toda intervención algorítmica de "
        "impacto relevante tenga un responsable humano identificable, con competencia efectiva "
        "para validar, explicar y responder por sus efectos.",
    ),
    (
        "Artículo 14°. Grado de delegación y matriz institucional",
        "No todo uso de IA implica el mismo grado de transferencia de una tarea o decisión. "
        "Este artículo distingue lo no delegable, lo condicionado y lo delegable, y ordena "
        "esa clasificación en la Matriz institucional (Anexo B), a fin de evitar tanto la "
        "prohibición absoluta como la delegación implícita e incontrolada.",
    ),
    (
        "Artículo 15°. Mapa, registro y trazabilidad de sistemas",
        "La Matriz, el Mapa y el Registro forman un sistema único de trazabilidad: los datos "
        "comunes se consignan una sola vez y alimentan los tres instrumentos sin cargas "
        "duplicadas. El Mapa visibiliza decisiones críticas y zonas grises; el Registro "
        "documenta los sistemas institucionales autorizados a lo largo de su ciclo de vida.",
    ),
    (
        "Artículo 16°. Evaluación de riesgo e impacto",
        "Antes de adoptar, contratar, desarrollar o ampliar un sistema de IA institucional "
        "con potencial impacto sobre personas, evaluación, investigación o gestión, la unidad "
        "responsable deberá evaluar riesgo e impacto, con acompañamiento del Observatorio y "
        "de las áreas técnicas competentes. La evaluación es condición de autorización, no "
        "un trámite posterior al despliegue.",
    ),
    (
        "Artículo 17°. Autorización y ciclo de vida de sistemas institucionales",
        "El ciclo de identificación, evaluación, autorización, registro, monitoreo y baja "
        "evita que un sistema institucional se incorpore de hecho, sin responsable, sin "
        "evaluación ni posibilidad de retiro. Los tramos operativos de registro y seguimiento "
        "se integran con el Artículo 15° y se desarrollan en el Anexo D.",
    ),
    (
        "Artículo 18°. Usos individuales y zonas grises",
        "No todo uso de IA es un sistema institucional. Los usos individuales de docentes, "
        "estudiantes o personal se rigen por los Capítulos IV y siguientes. Este artículo "
        "precisa que las prácticas extendidas “de hecho” deben visibilizarse, y que una zona "
        "gris no equivale a autorización implícita.",
    ),
    (
        "Artículo 19°. Coordinación entre sedes",
        "Dada la estructura multi-sede de la Universidad, la gobernanza de la IA se ejercerá "
        "con criterios comunes de Universidad. Podrán admitirse adecuaciones locales solo "
        "cuando no contradigan este Marco ni debiliten la protección de las personas, la "
        "integridad académica o la seguridad de la información.",
    ),
    (
        "Artículo 20°. Aprendizaje institucional",
        "Gobernar la IA no se agota en autorizar o prohibir: exige registrar lo aprendido, "
        "corregir prácticas y actualizar instrumentos. Este artículo incorpora esa función "
        "de mejora continua como parte del modelo de gobernanza, y no como un capítulo "
        "aislado de seguimiento.",
    ),
]

ART19_EXTRA = [
    "El Observatorio promoverá la coordinación entre sedes, facultades y niveles educativos "
    "dependientes, evitando regulaciones contradictorias o vacíos de responsabilidad.",
    "A esos efectos, se observará el siguiente procedimiento mínimo: (i) las unidades y sedes "
    "informarán al Observatorio los usos institucionales relevantes, incidentes y criterios "
    "locales que pretendan adoptar; (ii) el Observatorio identificará divergencias, redundancias "
    "o vacíos y propondrá un criterio común o una adecuación fundada; (iii) cuando la "
    "divergencia afecte derechos de personas, integridad académica o seguridad de la "
    "información, el Rectorado resolverá con alcance de Universidad; (iv) las adecuaciones "
    "locales se documentarán y no podrán ser más permisivas que este Marco.",
    "Las facultades, institutos y colegios dependientes podrán dictar orientaciones propias "
    "más restrictivas para sus ámbitos, siempre que sean coherentes con este Marco, se "
    "comuniquen a la comunidad destinataria y se pongan en conocimiento del Observatorio.",
]

LIST_FIXES = {
    "Finalidad explícita;": "a) Finalidad explícita;",
    "Supervisión humana efectiva;": "b) Supervisión humana efectiva;",
    "Declaración o registro cuando corresponda;": "c) Declaración o registro cuando corresponda;",
    "Posibilidad de revisión o impugnación en decisiones con efecto sobre personas.": (
        "d) Posibilidad de revisión o impugnación en decisiones con efecto sobre personas."
    ),
}

TEXT_REPLACES = [
    (
        "Las competencias del Consejo Superior y del Rectorado en materia de IA son las establecidas en el Artículo 14°. Las responsabilidades institucionales se ejercerán con carácter de rendición de cuentas, coordinación interáreas y respuesta oportuna ante riesgos relevantes, sin perjuicio de las atribuciones estatutarias de cada órgano.",
        "Las competencias del Consejo Superior y del Rectorado en materia de IA son las establecidas en este artículo. Las responsabilidades institucionales se ejercerán con carácter de rendición de cuentas, coordinación interáreas y respuesta oportuna ante riesgos relevantes, sin perjuicio de las atribuciones estatutarias de cada órgano.",
    ),
]


def delete_paragraph(p: Paragraph):
    p._element.getparent().remove(p._element)


def insert_after(anchor: Paragraph, text: str) -> Paragraph:
    new_el = OxmlElement("w:p")
    anchor._p.addnext(new_el)
    p = Paragraph(new_el, anchor._parent)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.line_spacing = 1.5
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    return p


def body_start_index(doc: Document) -> int:
    for i, p in enumerate(doc.paragraphs):
        if "VERSIÓN 2026" in p.text or "VERSIÓN 2.0 DEPURADA" in p.text:
            return i
    return 0


def already_has_next(doc: Document, idx: int, snippet: str) -> bool:
    for j in range(idx + 1, min(idx + 4, len(doc.paragraphs))):
        if snippet[:40] in doc.paragraphs[j].text:
            return True
    return False


def find_heading(doc: Document, prefix: str, start: int = 0):
    for i, p in enumerate(doc.paragraphs):
        if i < start or p.style.name.startswith("toc"):
            continue
        t = p.text.strip().split("\t")[0]
        if t == prefix or (t.startswith(prefix) and len(t) < 160):
            if (
                p.style.name.startswith("Heading")
                or t.startswith("Artículo")
                or re.match(r"^\d+\.\d+\.", t)
            ):
                return i, p
    return None, None


def apply_insertions(doc: Document) -> int:
    n = 0
    start = body_start_index(doc)
    for anchor_text, prose in INSERTIONS:
        idx, p = find_heading(doc, anchor_text, start)
        if p is None:
            print("MISSING", anchor_text)
            continue
        if already_has_next(doc, idx, prose):
            start = idx + 1
            continue
        insert_after(p, prose)
        n += 1
        start = idx + 1
    return n


def expand_art19(doc: Document) -> int:
    start = body_start_index(doc)
    idx, p = find_heading(doc, "Artículo 19°. Coordinación entre sedes", start)
    if p is None:
        return 0
    # skip existing two short paragraphs if they match v1; replace/add procedure
    # Find the paragraph that starts with "El Observatorio promoverá"
    obs = None
    obs_idx = None
    for j in range(idx + 1, idx + 8):
        t = doc.paragraphs[j].text.strip()
        if t.startswith("El Observatorio promoverá la coordinación"):
            obs = doc.paragraphs[j]
            obs_idx = j
            break
        if t.startswith("Artículo 20°"):
            break
    n = 0
    if obs is None:
        # insert after heading intro
        cur = p
        for para in ART19_EXTRA:
            cur = insert_after(cur, para)
            n += 1
        return n
    if already_has_next(doc, obs_idx, "A esos efectos, se observará"):
        return 0
    insert_after(obs, ART19_EXTRA[1])
    # find the new last extra and add third
    for j, q in enumerate(doc.paragraphs):
        if "A esos efectos, se observará" in q.text:
            insert_after(q, ART19_EXTRA[2])
            n += 2
            break
    return n or 1


def fix_lists_and_refs(doc: Document) -> int:
    n = 0
    start = body_start_index(doc)
    in_iii = False
    for i, p in enumerate(doc.paragraphs):
        if i < start or p.style.name.startswith("toc"):
            continue
        t = p.text.strip()
        if t.startswith("CAPÍTULO III."):
            in_iii = True
        if t.startswith("CAPÍTULO IV."):
            break
        if not in_iii:
            continue
        if t in LIST_FIXES:
            new = LIST_FIXES[t]
            if p.runs:
                p.runs[0].text = new
                for r in p.runs[1:]:
                    r.text = ""
            else:
                p.text = new
            n += 1
        for old, new in TEXT_REPLACES:
            if old in p.text:
                nt = p.text.replace(old, new)
                if p.runs:
                    p.runs[0].text = nt
                    for r in p.runs[1:]:
                        r.text = ""
                else:
                    p.text = nt
                n += 1
    return n


def main():
    backup = TARGET.replace(".docx", f" - backup pre-capIII {date.today().isoformat()}.docx")
    shutil.copy2(TARGET, backup)
    doc = Document(TARGET)
    stats = {
        "insertions": apply_insertions(doc),
        "art19": expand_art19(doc),
        "lists_refs": fix_lists_and_refs(doc),
    }
    doc.save(TARGET)
    print(f"Backup: {backup}")
    print(f"Saved: {TARGET}")
    for k, v in stats.items():
        print(f"  {k}: {v}")


if __name__ == "__main__":
    main()
