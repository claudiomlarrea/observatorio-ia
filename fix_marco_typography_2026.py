#!/usr/bin/env python3
"""Typography-only pass for Marco de Gobernanza IA - UCCuyo 2026.docx.

Unifies fonts, sizes, heading colors and body rhythm. Does not change text.
"""
from __future__ import annotations

import shutil
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

TARGET = Path(
    "/Users/claudiolarrea/Library/CloudStorage/OneDrive-Personal/"
    "16 Secretaría de Investigación/60 Observatorio de Inteligencia Artificial/"
    "MARCO DE GOBERNANZA DE LA IA/Documento FINAL/"
    "Marco de Gobernanza IA - UCCuyo 2026.docx"
)
BACKUP_DIR = TARGET.parent / "Versiones previas"
WORKSPACE_COPY = Path(
    "/Users/claudiolarrea/Documents/Observatorio/"
    "Marco de Gobernanza IA - UCCuyo FINAL.docx"
)

FONT = "Times New Roman"
BORDO = RGBColor(0x7A, 0x15, 0x32)  # institucional UCCuyo
INK = RGBColor(0x1A, 0x1A, 0x1A)
H3_INK = RGBColor(0x2E, 0x2E, 0x2E)


def set_style_font(style, *, name, size_pt=None, bold=None, italic=None, color=None):
    style.font.name = name
    # Ensure ASCII/hAnsi both point to the same face (Word quirk)
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    rfonts.set(qn("w:ascii"), name)
    rfonts.set(qn("w:hAnsi"), name)
    rfonts.set(qn("w:cs"), name)
    if size_pt is not None:
        style.font.size = Pt(size_pt)
    if bold is not None:
        style.font.bold = bold
    if italic is not None:
        style.font.italic = italic
    if color is not None:
        style.font.color.rgb = color


def clear_run_typeface(run, *, keep_bold_italic=True):
    """Remove direct typeface/size/color so the paragraph style wins."""
    run.font.name = None
    run.font.size = None
    try:
        run.font.color.rgb = None
    except Exception:
        pass
    # Clear theme color if present
    rpr = run._element.get_or_add_rPr()
    solid = rpr.find(qn("w:color"))
    if solid is not None:
        rpr.remove(solid)
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is not None:
        rpr.remove(rfonts)
    sz = rpr.find(qn("w:sz"))
    if sz is not None:
        rpr.remove(sz)
    szCs = rpr.find(qn("w:szCs"))
    if szCs is not None:
        rpr.remove(szCs)
    if not keep_bold_italic:
        run.bold = None
        run.italic = None


def configure_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    set_style_font(normal, name=FONT, size_pt=12, color=INK)
    pf = normal.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    pf.space_after = Pt(6)
    pf.space_before = Pt(0)
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    h1 = doc.styles["Heading 1"]
    set_style_font(h1, name=FONT, size_pt=14, bold=True, color=BORDO)
    h1.paragraph_format.space_before = Pt(22)
    h1.paragraph_format.space_after = Pt(10)
    h1.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    h1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    h1.paragraph_format.keep_with_next = True

    h2 = doc.styles["Heading 2"]
    set_style_font(h2, name=FONT, size_pt=12, bold=True, color=BORDO)
    h2.paragraph_format.space_before = Pt(14)
    h2.paragraph_format.space_after = Pt(6)
    h2.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    h2.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    h2.paragraph_format.keep_with_next = True

    h3 = doc.styles["Heading 3"]
    set_style_font(h3, name=FONT, size_pt=12, bold=True, color=H3_INK)
    h3.paragraph_format.space_before = Pt(10)
    h3.paragraph_format.space_after = Pt(4)
    h3.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    h3.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    h3.paragraph_format.keep_with_next = True

    for toc_name, size in (("toc 1", 11), ("toc 2", 11)):
        try:
            toc = doc.styles[toc_name]
        except KeyError:
            continue
        set_style_font(toc, name=FONT, size_pt=size, color=INK)
        toc.paragraph_format.space_after = Pt(3)
        toc.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE


def normalize_paragraphs(doc: Document) -> dict:
    stats = {"headings_cleaned": 0, "normal_justified": 0, "typeface_cleared": 0}

    for p in doc.paragraphs:
        st = p.style.name if p.style else ""

        if st.startswith("Heading"):
            for r in p.runs:
                # Drop Arial / 11pt / theme-blue overrides; keep emphasis
                clear_run_typeface(r, keep_bold_italic=True)
                r.bold = True
            stats["headings_cleaned"] += 1
            continue

        if st.startswith("toc"):
            for r in p.runs:
                # TOC entries: clear rogue faces; keep as style
                if r.font.name and r.font.name != FONT:
                    clear_run_typeface(r, keep_bold_italic=True)
                    stats["typeface_cleared"] += 1
            continue

        if st == "Normal" and p.text.strip():
            # Justify body (skip very short labels already bold index titles — still ok)
            if p.alignment != WD_ALIGN_PARAGRAPH.JUSTIFY:
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                stats["normal_justified"] += 1
            for r in p.runs:
                # Only strip alien typefaces (Arial) or odd explicit sizes on body
                if r.font.name and r.font.name not in (FONT, None):
                    r.font.name = FONT
                    stats["typeface_cleared"] += 1
                # Remove explicit 11pt on long body that should follow Normal 12
                if r.font.size and abs(r.font.size.pt - 11.0) < 0.1 and len(p.text) > 80:
                    r.font.size = None
                    stats["typeface_cleared"] += 1

    return stats


def normalize_body_section(doc: Document) -> None:
    """Keep cover (section 0) as-is; normalize body page margins to 2.5 cm."""
    if len(doc.sections) < 2:
        return
    body = doc.sections[1]
    body.left_margin = Cm(2.5)
    body.right_margin = Cm(2.5)
    body.top_margin = Cm(2.5)
    body.bottom_margin = Cm(2.5)


def main() -> None:
    BACKUP_DIR.mkdir(parents=True, exist_ok=True)
    backup = (
        BACKUP_DIR
        / f"Marco de Gobernanza IA - UCCuyo 2026.pre-typo-{date.today().isoformat()}.docx"
    )
    shutil.copy2(TARGET, backup)
    print("Backup:", backup)

    doc = Document(str(TARGET))
    configure_styles(doc)
    stats = normalize_paragraphs(doc)
    normalize_body_section(doc)
    print("Paragraph stats:", stats)

    doc.save(str(TARGET))
    shutil.copy2(TARGET, WORKSPACE_COPY)
    print("Saved:", TARGET)

    # Verify
    doc2 = Document(str(TARGET))
    for name in ("Normal", "Heading 1", "Heading 2", "Heading 3"):
        st = doc2.styles[name]
        print(
            f"{name}: {st.font.name} {st.font.size.pt if st.font.size else '-'}pt "
            f"bold={st.font.bold} color={st.font.color.rgb}"
        )
    aliens = 0
    for p in doc2.paragraphs:
        for r in p.runs:
            if r.font.name and r.font.name not in (FONT,):
                aliens += 1
    print("Alien run fonts left:", aliens)


if __name__ == "__main__":
    main()
