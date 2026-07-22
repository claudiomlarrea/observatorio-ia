#!/usr/bin/env python3
"""Plantilla Word (.docx) para resúmenes de las Jornadas IA 2026 (banner estilo PPT)."""
from __future__ import annotations

import shutil
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
OUT_REPO = ROOT / "docs/plantillas/plantilla-resumen-jornadas-ia-2026.docx"
OUT_ONEDRIVE = Path(
    "/Users/claudiolarrea/Library/CloudStorage/OneDrive-Personal/16 Secretaría de Investigación/"
    "60 Observatorio de Inteligencia Artificial/Jornadas de IA 2026/"
    "plantilla-resumen-jornadas-ia-2026.docx"
)
LOGO = ROOT / "assets/logo-observatorio-ia.png"
LOGO_CIRCLE = ROOT / "assets/logo-observatorio-ia-circle.png"

GREEN = RGBColor(0x04, 0x2F, 0x23)
GOLD = RGBColor(0xC9, 0xA2, 0x27)
MUTED = RGBColor(0x5C, 0x4F, 0x54)


def _ensure_circle_logo() -> Path:
    if LOGO_CIRCLE.is_file():
        return LOGO_CIRCLE
    from PIL import Image, ImageDraw

    im = Image.open(LOGO).convert("RGBA")
    size = min(im.size)
    im = im.resize((size, size), Image.Resampling.LANCZOS)
    mask = Image.new("L", (size, size), 0)
    ImageDraw.Draw(mask).ellipse((0, 0, size - 1, size - 1), fill=255)
    out = Image.new("RGBA", (size, size), (0, 0, 0, 0))
    out.paste(im, (0, 0))
    out.putalpha(mask)
    LOGO_CIRCLE.parent.mkdir(parents=True, exist_ok=True)
    out.save(LOGO_CIRCLE, "PNG")
    return LOGO_CIRCLE


def _set_run_font(run, *, size=10, bold=False, italic=False, color=None, name="Arial"):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color is not None:
        run.font.color.rgb = color


def _shade_cell(cell, hex_color: str) -> None:
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    shd.set(qn("w:val"), "clear")
    tcPr.append(shd)


def _set_cell_margins(cell, top=40, bottom=40, left=60, right=60) -> None:
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement("w:tcMar")
    for m, v in (("top", top), ("left", left), ("bottom", bottom), ("right", right)):
        node = OxmlElement(f"w:{m}")
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")
        tcMar.append(node)
    tcPr.append(tcMar)


def _add_banner(doc: Document) -> None:
    logo = _ensure_circle_logo()
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    # Anchos aprox. página A4 con márgenes
    cell_logo = table.cell(0, 0)
    cell_text = table.cell(0, 1)
    cell_logo.width = Cm(2.6)
    cell_text.width = Cm(14.5)

    for cell in (cell_logo, cell_text):
        _shade_cell(cell, "042F23")
        _set_cell_margins(cell, top=80, bottom=80, left=80, right=80)

    # Logo
    p_logo = cell_logo.paragraphs[0]
    p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_logo.add_run()
    run.add_picture(str(logo), width=Cm(1.9))

    # Textos banner (como en PPT)
    lines = [
        ("UNIVERSIDAD CATÓLICA DE CUYO · Observatorio de IA", 9, False),
        ("1° Jornadas internas de Inteligencia Artificial 2026", 13, True),
        ("Plantilla de resumen · Encuentro virtual · 6 de octubre de 2026 · 15:00 h", 9, False),
    ]
    first = True
    for text, size, bold in lines:
        p = cell_text.paragraphs[0] if first else cell_text.add_paragraph()
        first = False
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(text)
        _set_run_font(r, size=size, bold=bold, color=RGBColor(0xFF, 0xFF, 0xFF))

    # Franja dorada
    gold = doc.add_table(rows=1, cols=1)
    gold.alignment = WD_TABLE_ALIGNMENT.CENTER
    gcell = gold.cell(0, 0)
    gcell.width = Cm(17.1)
    _shade_cell(gcell, "C9A227")
    _set_cell_margins(gcell, top=0, bottom=0, left=0, right=0)
    gp = gcell.paragraphs[0]
    gp.paragraph_format.space_before = Pt(0)
    gp.paragraph_format.space_after = Pt(0)
    gr = gp.add_run(" ")
    _set_run_font(gr, size=4, color=GOLD)


def _p(
    doc: Document,
    text: str,
    *,
    size=10,
    bold=False,
    italic=False,
    center=False,
    color=None,
    space_after=6,
    space_before=0,
):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    r = p.add_run(text)
    _set_run_font(r, size=size, bold=bold, italic=italic, color=color)
    return p


def build() -> Path:
    doc = Document()

    for section in doc.sections:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.8)
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(2.0)
        # Pie
        footer = section.footer
        footer.is_linked_to_previous = False
        fp = footer.paragraphs[0]
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        fr = fp.add_run(
            "Observatorio de IA · UCCuyo  ·  "
            "https://claudiomlarrea.github.io/observatorio-ia/#jornadas-ia  ·  "
            "observatorioia@uccuyo.edu.ar"
        )
        _set_run_font(fr, size=8, color=MUTED)

    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(10)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")

    _add_banner(doc)

    _p(
        doc,
        "INSTRUCCIONES (borrar este párrafo al enviar): Arial 10 pt, interlineado sencillo, "
        "máximo 250 palabras en el cuerpo del resumen. Sin cuadros, gráficos ni bibliografía. "
        "Subrayá al autor que presentará. Nombre del archivo: Area_Universidad_Apellido_Titulo.docx",
        size=9,
        italic=True,
        color=MUTED,
        space_before=10,
        space_after=12,
    )

    _p(
        doc,
        "Título de la ponencia (mayúsculas/minúsculas tipo oración)",
        size=10,
        bold=True,
        center=True,
        space_after=12,
    )

    _p(
        doc,
        "N. Apellido¹, N. Apellido²  (subrayar al expositor; sin título de grado)",
        size=10,
        bold=True,
        center=True,
        space_after=6,
    )

    _p(
        doc,
        "Unidad académica / Facultad o Instituto",
        size=10,
        italic=True,
        center=True,
        space_after=2,
    )
    _p(
        doc,
        "correo@uccuyo.edu.ar",
        size=10,
        italic=True,
        center=True,
        space_after=14,
    )

    _p(doc, "Texto del resumen", size=10, bold=True, center=False, space_after=8)

    for label, hint in [
        ("Introducción.", "Breve contexto y objetivo del trabajo."),
        ("Metodología.", "Enfoque, datos, participantes o procedimiento."),
        ("Resultados.", "Hallazgos principales, de forma concisa."),
        ("Conclusiones.", "Mensaje clave e implicancias."),
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.space_after = Pt(8)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        r1 = p.add_run(label + " ")
        _set_run_font(r1, size=10, bold=True)
        r2 = p.add_run(hint + " [reemplazar este texto]")
        _set_run_font(r2, size=10, italic=True, color=MUTED)

    _p(
        doc,
        "Cierre de carga: 10 de septiembre de 2026 · Observatorio de IA — "
        "https://claudiomlarrea.github.io/observatorio-ia/#jornadas-ia",
        size=9,
        italic=True,
        color=MUTED,
        space_before=16,
        center=True,
    )

    OUT_REPO.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUT_REPO))
    return OUT_REPO


def main() -> None:
    path = build()
    print(f"Wrote {path} ({path.stat().st_size} bytes)")
    if OUT_ONEDRIVE.parent.is_dir():
        shutil.copy2(path, OUT_ONEDRIVE)
        print(f"Copied to {OUT_ONEDRIVE}")
    else:
        print(f"Skip OneDrive: {OUT_ONEDRIVE.parent}")


if __name__ == "__main__":
    main()
