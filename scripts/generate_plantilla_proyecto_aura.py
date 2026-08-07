#!/usr/bin/env python3
"""Documento único: Convocatoria AURA 2026 + Anexo I (plantilla Secretaría).

Parte A — Bases de la convocatoria (investigación y extensión).
Parte B — Formulario de presentación (Anexo I de la Secretaría de Investigación
          y Vinculación Tecnológica, adaptado al Plan AURA).
"""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
OUT_COMBINED = (
    ROOT
    / "instituto-del-agua/documentos"
    / "Convocatoria_y_Plantilla_Proyectos_AURA_2026.docx"
)
OUT_PLANTILLA = (
    ROOT
    / "instituto-del-agua/documentos/plantillas"
    / "plantilla-proyecto-investigacion-extension-AURA-2026.docx"
)
LOGO_GOTA = ROOT / "instituto-del-agua/assets/logo-aura-gota.png"
LOGO_UCCUYO = ROOT / "instituto-del-agua/assets/logo-uccuyo.png"

GREEN = RGBColor(0x04, 0x2F, 0x23)
WATER = RGBColor(0x0E, 0x4F, 0x72)
INST = RGBColor(0x7A, 0x15, 0x32)
MUTED = RGBColor(0x5C, 0x4F, 0x54)
BLACK = RGBColor(0x1F, 0x14, 0x18)


def _set_run_font(run, *, size=10, bold=False, italic=False, color=None, name="Arial"):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color is not None:
        run.font.color.rgb = color


def _shade_cell(cell, hex_color: str) -> None:
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    shd.set(qn("w:val"), "clear")
    tcPr.append(shd)


def _set_cell_margins(cell, top=40, bottom=40, left=60, right=60) -> None:
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for child in list(tcPr):
        if child.tag == qn("w:tcMar"):
            tcPr.remove(child)
    tcMar = OxmlElement("w:tcMar")
    for m, v in (("top", top), ("left", left), ("bottom", bottom), ("right", right)):
        node = OxmlElement(f"w:{m}")
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")
        tcMar.append(node)
    tcPr.append(tcMar)


def _p(
    doc: Document,
    text: str,
    *,
    size=10,
    bold=False,
    italic=False,
    center=False,
    color=None,
    space_after=6,
    space_before=0,
    justify=True,
):
    p = doc.add_paragraph()
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif justify:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    else:
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    r = p.add_run(text)
    _set_run_font(r, size=size, bold=bold, italic=italic, color=color)
    return p


def _heading(doc: Document, text: str, *, color=WATER) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    _set_run_font(r, size=11, bold=True, color=color)


def _bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text)
    _set_run_font(r, size=10, color=BLACK)


def _field(doc: Document, label: str, hint: str = "") -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.space_before = Pt(2)
    r1 = p.add_run(label)
    _set_run_font(r1, size=10, bold=True, color=BLACK)
    if hint:
        r2 = p.add_run("  " + hint)
        _set_run_font(r2, size=9, italic=True, color=MUTED)


def _blank_lines(doc: Document, n: int = 2) -> None:
    for _ in range(n):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run("___________________________________________________________________________")
        _set_run_font(r, size=9, color=RGBColor(0xC8, 0xC0, 0xC4))


def _hint_box(doc: Document, lines: list[str]) -> None:
    _p(
        doc,
        "Recomendaciones / completar:",
        size=9,
        bold=True,
        color=GREEN,
        space_before=4,
        space_after=2,
        justify=False,
    )
    for line in lines:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(1)
        r = p.add_run(line)
        _set_run_font(r, size=9, italic=True, color=MUTED)


def _page_break(doc: Document) -> None:
    p = doc.add_paragraph()
    run = p.add_run()
    run.add_break(WD_BREAK.PAGE)


def _add_banner(doc: Document, subtitle: str) -> None:
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    cell_logo = table.cell(0, 0)
    cell_text = table.cell(0, 1)
    cell_logo.width = Cm(2.6)
    cell_text.width = Cm(14.5)

    for cell in (cell_logo, cell_text):
        _shade_cell(cell, "0E4F72")
        _set_cell_margins(cell, top=80, bottom=80, left=80, right=80)

    p_logo = cell_logo.paragraphs[0]
    p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    logo = LOGO_GOTA if LOGO_GOTA.is_file() else LOGO_UCCUYO
    if logo.is_file():
        run = p_logo.add_run()
        run.add_picture(str(logo), width=Cm(1.85))

    lines = [
        ("UNIVERSIDAD CATÓLICA DE CUYO · Secretaría de Investigación y Extensión", 8, False),
        ("Plan Integral AURA · Convocatoria 2026", 13, True),
        (subtitle, 9, False),
    ]
    first = True
    for text, size, bold in lines:
        p = cell_text.paragraphs[0] if first else cell_text.add_paragraph()
        first = False
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(text)
        _set_run_font(r, size=size, bold=bold, color=RGBColor(0xFF, 0xFF, 0xFF))

    bar = doc.add_table(rows=1, cols=1)
    bar.alignment = WD_TABLE_ALIGNMENT.CENTER
    gcell = bar.cell(0, 0)
    gcell.width = Cm(17.1)
    _shade_cell(gcell, "1A7FB5")
    _set_cell_margins(gcell, top=0, bottom=0, left=0, right=0)
    gp = gcell.paragraphs[0]
    gp.paragraph_format.space_before = Pt(0)
    gp.paragraph_format.space_after = Pt(0)
    gr = gp.add_run(" ")
    _set_run_font(gr, size=4, color=WATER)


def _add_cronograma_table(doc: Document) -> None:
    rows = [
        ("Etapa / Actividad", "Fecha"),
        (
            "Presentación de la propuesta ante Consejos y Consejo Superior",
            "Agosto de 2026",
        ),
        ("Apertura y período de presentación de proyectos", "15 sep. – 2 oct. 2026"),
        ("Cierre de la recepción de proyectos", "2 de octubre de 2026"),
        ("Evaluación de los proyectos presentados", "5 – 16 de octubre de 2026"),
        ("Presentación de aprobados ante Consejo Superior", "Viernes 30 de octubre de 2026"),
        ("Oficialización del listado de proyectos aprobados", "6 de noviembre de 2026"),
        ("Inicio de ejecución de los proyectos aprobados", "Una vez aprobados por CS"),
        ("Informe de avance (a los 6 meses)", "Mayo de 2027"),
        ("Finalización de la ejecución / informe final", "Hasta octubre de 2027"),
    ]
    table = doc.add_table(rows=len(rows), cols=2)
    table.style = "Table Grid"
    for i, (a, b) in enumerate(rows):
        c0, c1 = table.rows[i].cells
        if i == 0:
            _shade_cell(c0, "0E4F72")
            _shade_cell(c1, "0E4F72")
            color = RGBColor(0xFF, 0xFF, 0xFF)
            bold = True
        else:
            color = BLACK
            bold = False
        r0 = c0.paragraphs[0].add_run(a)
        r1 = c1.paragraphs[0].add_run(b)
        _set_run_font(r0, size=9, bold=bold, color=color)
        _set_run_font(r1, size=9, bold=bold, color=color)


def _add_convocatoria(doc: Document) -> None:
    _add_banner(
        doc,
        "Documento único: bases de la convocatoria + Anexo I de presentación de proyectos",
    )

    _p(
        doc,
        "CONVOCATORIA A PROYECTOS DE INVESTIGACIÓN Y EXTENSIÓN",
        size=13,
        bold=True,
        center=True,
        color=INST,
        space_before=14,
        space_after=4,
    )
    _p(
        doc,
        "Vinculados al Ahorro y Uso Responsable del Agua · Plan Integral AURA",
        size=11,
        bold=True,
        center=True,
        color=WATER,
        space_after=10,
    )
    _p(
        doc,
        "Secretaría de Investigación – Secretaría de Extensión · Universidad Católica de Cuyo",
        size=10,
        center=True,
        color=MUTED,
        space_after=12,
    )

    _heading(doc, "1. Fundamentación", color=INST)
    for para in [
        "La Universidad Católica de Cuyo, a través de la Resolución N.º 418-CS-2024, "
        "aprobó el “Plan Integral de Ahorro y Uso Responsable del Agua” (Plan AURA), "
        "en consonancia con el Plan Estratégico Institucional 2023-2027, reconociendo "
        "el contexto crítico de escasez hídrica que atraviesa la provincia de San Juan "
        "y asumiendo el compromiso institucional de promover una gestión eficiente, "
        "responsable y sostenible del recurso hídrico en las tres funciones sustantivas "
        "de la universidad: docencia, investigación y extensión.",
        "Posteriormente, mediante la Resolución N.º 767-CS-2025, el Consejo Superior "
        "aprobó el “Plan de Integración Curricular del Proyecto AURA”, estableciendo "
        "entre sus ejes de desarrollo la integración curricular del eje “Ahorro y Uso "
        "Responsable del Agua” en programas académicos, talleres y proyectos de "
        "investigación y/o extensión, con el propósito de consolidar la sostenibilidad "
        "hídrica como un eje transversal del proyecto educativo institucional y "
        "posicionar a la UCCuyo como referente regional en la construcción de una "
        "cultura hídrica responsable.",
        "A su vez, la Resolución N.º 849-CS-2026 aprobó el “Plan de Integración "
        "Curricular AURA – Planificación 2026”, documento que sistematiza las "
        "propuestas formuladas por las distintas Unidades Académicas, niveles "
        "educativos y áreas de la Sede San Juan. Dicho Plan incorpora, en el apartado "
        "correspondiente a Secretarías/Rectorado, la acción institucional “Convocatoria "
        "para Proyectos de Investigación y Extensión”, bajo la responsabilidad "
        "conjunta de la Secretaría de Investigación y la Secretaría de Extensión, "
        "con un plazo estimado de ejecución de 12 meses y sujeta a la aprobación del "
        "Consejo Superior.",
        "En cumplimiento de dicha planificación y en el marco normativo descripto, "
        "se emite la presente convocatoria, orientada a promover la presentación de "
        "proyectos de investigación y extensión que aporten a la generación de "
        "conocimiento, la innovación aplicada y la transferencia a la comunidad en "
        "relación con el uso responsable y sostenible del recurso hídrico.",
    ]:
        _p(doc, para, space_after=8)

    _heading(doc, "2. Objetivo general", color=INST)
    _p(
        doc,
        "Promover la presentación de proyectos de investigación y extensión vinculados "
        "al ahorro y uso responsable del agua, en el marco del Plan Integral AURA, que "
        "contribuyan a la generación de conocimiento aplicado, la sensibilización "
        "institucional y comunitaria, y la construcción de una cultura hídrica "
        "responsable en la Universidad Católica de Cuyo y su área de influencia.",
        space_after=8,
    )

    _heading(doc, "3. Objetivos específicos", color=INST)
    for item in [
        "Fomentar la producción de conocimiento científico y aplicado sobre la "
        "problemática hídrica en contextos áridos, con especial referencia a la "
        "provincia de San Juan.",
        "Impulsar acciones de extensión y vinculación que transfieran a la comunidad "
        "los conocimientos generados en el ámbito académico.",
        "Fortalecer la interacción entre la investigación y la extensión, de modo que "
        "ambas funciones trabajen de manera articulada.",
        "Promover la conformación de equipos interdisciplinarios, con participación "
        "de docentes, investigadores y estudiantes.",
        "Generar productos de investigación y extensión (informes, publicaciones, "
        "materiales de transferencia, propuestas de intervención, entre otros).",
        "Contribuir al cumplimiento de los indicadores institucionales previstos en "
        "el Plan de Integración Curricular AURA – Planificación 2026.",
    ]:
        _bullet(doc, item)

    _heading(doc, "4. Destinatarios", color=INST)
    _p(
        doc,
        "Podrán presentar proyectos los docentes, investigadores, equipos de cátedra, "
        "estudiantes avanzados y equipos interdisciplinarios pertenecientes a las "
        "Unidades Académicas, Institutos y Colegios en todos sus niveles educativos "
        "de la Universidad Católica de Cuyo, Sede San Juan.",
        space_after=8,
    )

    _heading(doc, "5. Ejes temáticos", color=INST)
    _p(
        doc,
        "Los proyectos deberán enmarcarse en uno o más de los siguientes ejes "
        "(indicar la letra correspondiente en el Anexo I):",
        space_after=4,
    )
    ejes = [
        "a) Estudios sobre consumo y comportamiento hídrico (hábitos, percepciones y prácticas).",
        "b) Innovación tecnológica aplicada a la medición, ahorro o reutilización del agua.",
        "c) Impacto económico, social y productivo del recurso hídrico en San Juan.",
        "d) Investigación aplicada en salud vinculada al uso racional del agua.",
        "e) Investigación ambiental y territorial sobre la situación hídrica regional.",
        "f) Programas de formación, sensibilización y transferencia comunitaria.",
        "g) Prácticas socio-comunitarias e intervenciones territoriales.",
        "h) Educación y uso responsable del recurso hídrico (todos los niveles).",
        "i) Legislación, marco normativo y políticas públicas del recurso hídrico.",
        "j) Otras temáticas transversales pertinentes al Plan AURA, a criterio del Comité Evaluador.",
    ]
    for e in ejes:
        _bullet(doc, e)

    _heading(doc, "6. Cronograma de la convocatoria", color=INST)
    _add_cronograma_table(doc)
    _p(
        doc,
        "Los proyectos aprobados tendrán una duración de doce (12) meses, con inicio "
        "una vez oficializado el listado por el Consejo Superior y finalización en "
        "octubre de 2027. Deberán presentar informe de avance a los seis (6) meses "
        "e informe final al concluir.",
        size=9,
        italic=True,
        color=MUTED,
        space_before=8,
        space_after=8,
    )

    _heading(doc, "7. Requisitos de presentación", color=INST)
    _p(
        doc,
        "Los proyectos se presentan ante la Secretaría de Investigación y la Secretaría "
        "de Extensión mediante el formulario institucional (Anexo I de este documento) "
        "e incluyen como mínimo:",
        space_after=4,
    )
    for item in [
        "Título del proyecto y eje temático (letra a–j).",
        "Fundamentación y objetivos (general y específicos).",
        "Equipo responsable y unidad académica de pertenencia.",
        "Metodología y plan de trabajo, con cronograma de actividades.",
        "Resultados esperados e indicadores de seguimiento.",
        "Recursos necesarios para su ejecución.",
        "Vinculación explícita con el Plan Integral AURA y con la Planificación 2026.",
    ]:
        _bullet(doc, item)

    _heading(doc, "8. Requisitos de directores y equipos", color=INST)
    for item in [
        "La dirección deberá estar a cargo de un/a docente de la UCCuyo.",
        "El equipo deberá integrar docentes y, como mínimo, dos (2) estudiantes por proyecto.",
    ]:
        _bullet(doc, item)

    _heading(doc, "9. Comité Evaluador", color=INST)
    for item in [
        "3 especialistas en temática hídrica (propuesta del equipo coordinador AURA).",
        "2 miembros del Consejo de Investigación.",
        "2 miembros del Consejo de Extensión.",
    ]:
        _bullet(doc, item)

    _heading(doc, "10. Circuito de presentación y aprobación", color=INST)
    for item in [
        "Presentación del proyecto en la Unidad Académica de pertenencia.",
        "Aprobación por el Consejo Directivo de la Unidad Académica.",
        "Carga del proyecto en el Drive institucional constituido a tal fin.",
        "Evaluación por el Comité Evaluador, con emisión de dictamen.",
        "Elevación a los Consejos de Extensión y de Investigación.",
        "Elevación al Consejo Superior para aprobación definitiva.",
        "Oficialización del listado, inicio de ejecución, informes de avance y final, "
        "y rendición de cuentas.",
    ]:
        _bullet(doc, item)

    _heading(doc, "11. Criterios de evaluación", color=INST)
    for item in [
        "Pertinencia y coherencia con los ejes temáticos de la convocatoria.",
        "Rigurosidad metodológica y viabilidad en el plazo de 12 meses.",
        "Idoneidad del/la director/a y conformación del equipo (mín. 2 estudiantes).",
        "Potencial de impacto institucional, académico y/o comunitario.",
        "Grado de integración entre docencia, investigación y extensión.",
        "Sostenibilidad y proyección de continuidad de la propuesta.",
    ]:
        _bullet(doc, item)

    _heading(doc, "12. Presupuesto y financiamiento", color=INST)
    _p(
        doc,
        "Presupuesto total de la convocatoria: veinte millones de pesos ($20.000.000). "
        "Tope máximo estimado por proyecto: un millón de pesos ($1.000.000). El monto "
        "efectivamente asignado dependerá del presupuesto presentado y de la evaluación "
        "del Comité Evaluador, pudiendo ser inferior al tope.",
        space_after=6,
    )
    for item in [
        "La transferencia de fondos queda condicionada a la aprobación del Comité "
        "Evaluador y del Consejo Superior, y a la presentación de informes.",
        "Toda modificación presupuestaria debe aprobarse antes de ejecutar el gasto.",
        "Los gastos se canalizan por la Unidad Académica e imputan al centro de costo "
        "que se genere al aprobar el proyecto.",
    ]:
        _bullet(doc, item)

    _heading(doc, "13. Marco normativo", color=INST)
    for item in [
        "Resolución N.º 418-CS-2024 – Aprobación del Plan Integral AURA.",
        "Resolución N.º 767-CS-2025 – Plan de Integración Curricular del Proyecto AURA.",
        "Resolución N.º 849-CS-2026 – Plan de Integración Curricular AURA – Planificación 2026.",
    ]:
        _bullet(doc, item)

    _p(
        doc,
        "Consultas: luisjimenez@uccuyo.edu.ar (Instituto del Agua) · "
        "investigacion@uccuyo.edu.ar (Secretaría de Investigación).",
        size=9,
        italic=True,
        color=MUTED,
        space_before=12,
        center=True,
    )


def _add_anexo_plantilla(doc: Document) -> None:
    _add_banner(
        doc,
        "Parte B · Anexo I — Presentación de proyecto de investigación / extensión "
        "(plantilla Secretaría de Investigación, adaptada a AURA)",
    )

    _p(
        doc,
        "ANEXO I · PRESENTACIÓN DE PROYECTO",
        size=13,
        bold=True,
        center=True,
        color=INST,
        space_before=14,
        space_after=4,
    )
    _p(
        doc,
        "Formulario institucional de la Secretaría de Investigación y Vinculación "
        "Tecnológica · Convocatoria Plan AURA 2026",
        size=10,
        center=True,
        color=MUTED,
        space_after=8,
    )
    _p(
        doc,
        "INSTRUCCIONES (borrar este párrafo al enviar): Completar todos los campos. "
        "Nombre sugerido del archivo: AURA_UA_ApellidoDirector_TituloCorto.docx. "
        "Presentación: 15 de septiembre al 2 de octubre de 2026.",
        size=9,
        italic=True,
        color=MUTED,
        space_after=10,
    )

    for label, hint in [
        ("Título del Proyecto:", "[sin abreviaturas]"),
        ("Director/a:", "[docente UCCuyo]"),
        ("Co-director/a:", "[si corresponde]"),
        ("Unidad Académica:", "[Facultad / Instituto / Colegio]"),
        ("Convocatoria / Año:", "Plan AURA · Investigación y Extensión · 2026"),
        ("Eje temático (letra a–j):", "[según punto 5 de las bases]"),
        ("Tipo de proyecto:", "Investigación / Extensión / Investigación y extensión"),
        ("Fecha de presentación:", "[dd/mm/aaaa]"),
    ]:
        _field(doc, label, hint)
        _blank_lines(doc, 1)

    _heading(doc, "1. Identificación del proyecto")
    for label, hint in [
        ("Denominación del Proyecto:", ""),
        ("Denominación abreviada (opcional):", ""),
        ("Tipo de Proyecto:", "Básico / Aplicado / Innovación / Extensión / Mixto"),
        ("Duración estimada (meses):", "12 meses (según convocatoria AURA)"),
        ("Unidad Académica / Instituto / Centro:", ""),
        ("Área disciplinar:", ""),
        ("Eje temático AURA (a–j):", "Indicar letra y título del eje"),
    ]:
        _field(doc, label, hint)
        _blank_lines(doc, 1)
    _hint_box(
        doc,
        [
            "Título claro y sin abreviaturas.",
            "Indicar de forma explícita el eje temático de la convocatoria AURA.",
            "Área disciplinar según clasificaciones institucionales.",
        ],
    )

    _heading(doc, "2. Datos del Director/a y Co-director/a")
    for label, hint in [
        ("Director/a:", ""),
        ("Cargo académico:", ""),
        ("Formación académica:", ""),
        ("Categoría de investigador/a:", ""),
        ("Correo institucional:", "@uccuyo.edu.ar"),
        ("Co-director/a (si corresponde):", ""),
        ("Cargo académico:", ""),
        ("Formación académica:", ""),
        ("Categoría de investigador/a:", ""),
        ("Correo institucional:", ""),
    ]:
        _field(doc, label, hint)
        _blank_lines(doc, 1)
    _hint_box(
        doc,
        [
            "La dirección debe estar a cargo de un/a docente de la UCCuyo.",
            "Correo institucional obligatorio.",
            "Cumplir requisitos para dirigir según Ordenanza vigente.",
        ],
    )

    _heading(doc, "3. Equipo de investigación / extensión")
    _p(
        doc,
        "Integrantes, roles, filiación institucional y funciones específicas. "
        "La convocatoria AURA exige como mínimo dos (2) estudiantes por proyecto.",
        size=10,
        italic=True,
        color=MUTED,
        space_after=6,
    )
    table = doc.add_table(rows=5, cols=4)
    table.style = "Table Grid"
    headers = ["Nombre completo", "Filiación institucional", "Rol", "Función en el proyecto"]
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        _shade_cell(cell, "0E4F72")
        r = cell.paragraphs[0].add_run(h)
        _set_run_font(r, size=9, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
    for row in table.rows[1:]:
        for cell in row.cells:
            cell.paragraphs[0].add_run(" ")
    _hint_box(
        doc,
        [
            "Incluir docentes y al menos dos estudiantes.",
            "Funciones alineadas con la metodología y el plan de trabajo.",
        ],
    )

    _heading(doc, "4. Resumen del proyecto (≈ 300 palabras)")
    _p(
        doc,
        "Síntesis del proyecto (problema, objetivo general, enfoque metodológico, "
        "población/unidad de análisis y resultados esperados):",
        size=10,
        space_after=4,
    )
    _blank_lines(doc, 5)

    _heading(doc, "5. Fundamentación (≈ 400–600 palabras)")
    _p(
        doc,
        "Justificación teórica, empírica y contextual del proyecto:",
        size=10,
        space_after=4,
    )
    _blank_lines(doc, 5)
    _hint_box(
        doc,
        [
            "Antecedentes empíricos relevantes y brechas de conocimiento.",
            "Relevancia institucional y marco conceptual.",
            "Justificación del valor agregado en contextos áridos / San Juan.",
        ],
    )

    _heading(doc, "6. Pertinencia, relevancia y alineación institucional (≈ 200–400 palabras)")
    _p(
        doc,
        "Justificar la importancia científica, social e institucional del proyecto.",
        size=10,
        space_after=4,
    )
    for label in [
        "Correspondencia con el PEI (citar ejes y líneas):",
        "Correspondencia con el Plan Estratégico de la Función I+D:",
        "Aporte institucional (carreras, estudiantes, comunidad):",
        "Contribución a líneas prioritarias de la unidad académica:",
        "Vinculación explícita con el Plan Integral AURA (Res. 418-CS-2024):",
        "Vinculación con el Plan de Integración Curricular AURA – Planificación 2026 "
        "(Res. 849-CS-2026):",
    ]:
        _field(doc, label)
        _blank_lines(doc, 2)

    _heading(doc, "7. Planteo del problema y objetivos (≈ 200–400 palabras)")
    for label in [
        "Planteo del problema de investigación / intervención:",
        "Objetivo general:",
        "Objetivos específicos:",
    ]:
        _field(doc, label)
        _blank_lines(doc, 3)
    _hint_box(
        doc,
        [
            "Regla de coherencia: Problema → Objetivo general → Objetivos específicos → Metodología.",
            "Identificar variables o dimensiones cuando corresponda.",
        ],
    )

    _heading(doc, "8. Originalidad y aporte al conocimiento (≈ 100–200 palabras)")
    _blank_lines(doc, 4)

    _heading(doc, "9. Marco teórico y estado del arte (≈ 600–800 palabras)")
    _blank_lines(doc, 6)
    _hint_box(
        doc,
        [
            "Conceptos, enfoques y teorías fundamentales.",
            "Estudios recientes (últimos 5 años); 30–40% de actualización recomendada.",
            "Estilo de citación único (APA 7ª, Vancouver, ISO 690, IEEE o Chicago).",
        ],
    )

    _heading(doc, "10. Metodología (≈ 800–900 palabras)")
    for label in [
        "Diseño del estudio / de la intervención:",
        "Población y muestra / destinatarios:",
        "Instrumentos y técnicas de recolección de datos:",
        "Procedimientos de análisis:",
        "Validación y fiabilidad:",
        "Consideraciones éticas:",
    ]:
        _field(doc, label)
        _blank_lines(doc, 2)

    _heading(doc, "11. Factibilidad y cronograma de actividades (≈ 200–300 palabras)")
    _p(
        doc,
        "Recursos disponibles, viabilidad temporal y plan de trabajo (12 meses):",
        size=10,
        space_after=6,
    )
    crono = doc.add_table(rows=7, cols=3)
    crono.style = "Table Grid"
    for i, h in enumerate(["Etapa / Actividad", "Período estimado", "Responsables"]):
        cell = crono.rows[0].cells[i]
        _shade_cell(cell, "064A38")
        r = cell.paragraphs[0].add_run(h)
        _set_run_font(r, size=9, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
    for row in crono.rows[1:]:
        for cell in row.cells:
            cell.paragraphs[0].add_run(" ")
    _p(
        doc,
        "Incluir informe de avance a los 6 meses e informe final (oct. 2027).",
        size=9,
        italic=True,
        color=MUTED,
        space_before=6,
    )

    _heading(doc, "12. Impacto esperado y plan de difusión / transferencia (≈ 200–300 palabras)")
    for label in [
        "Impacto académico o científico:",
        "Impacto social o institucional:",
        "Indicadores de seguimiento:",
        "Plan de difusión (publicaciones, congresos, jornadas, etc.):",
        "Plan de transferencia (productos, herramientas, informes, etc.):",
    ]:
        _field(doc, label)
        _blank_lines(doc, 2)

    _heading(doc, "13. Presupuesto, sostenibilidad y alineación institucional (≈ 200–300 palabras)")
    _p(
        doc,
        "Tope estimado de la convocatoria AURA: hasta $1.000.000 por proyecto "
        "(presupuesto total de la convocatoria: $20.000.000).",
        size=9,
        italic=True,
        color=WATER,
        space_after=6,
    )
    for label in [
        "Presupuesto total estimado ($):",
        "Detalle por rubro (personal, insumos, viajes, difusión, otros):",
        "Fuentes de financiamiento (internas / externas / convocatoria AURA):",
        "Plan de sostenibilidad del proyecto:",
        "Alineación con PEI / líneas prioritarias / Ordenanza:",
    ]:
        _field(doc, label)
        _blank_lines(doc, 2)

    _heading(doc, "14. Bibliografía")
    _p(
        doc,
        "Citar las fuentes principales en un único estilo (APA 7ª edición u otro admitido).",
        size=10,
        italic=True,
        color=MUTED,
        space_after=4,
    )
    _blank_lines(doc, 6)

    _heading(doc, "15. Firmas")
    for label in [
        "Firma y aclaración del Director/a:",
        "Unidad Académica:",
        "Fecha:",
        "Conformidad del Consejo Directivo de la Unidad Académica (fecha / resolución):",
    ]:
        _field(doc, label)
        _blank_lines(doc, 2)

    _p(
        doc,
        "Documento combinado · Convocatoria AURA 2026 + Anexo I Secretaría de Investigación · UCCuyo",
        size=9,
        italic=True,
        color=MUTED,
        center=True,
        space_before=16,
    )


def _new_doc() -> Document:
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.8)
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(2.0)
        footer = section.footer
        footer.is_linked_to_previous = False
        fp = footer.paragraphs[0]
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        fr = fp.add_run(
            "Plan AURA · Instituto del Agua / Secretarías de Investigación y Extensión · UCCuyo  ·  "
            "https://claudiomlarrea.github.io/observatorio-ia/instituto-del-agua/#convocatoria"
        )
        _set_run_font(fr, size=8, color=MUTED)

    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(10)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    return doc


def build() -> tuple[Path, Path]:
    doc = _new_doc()
    _add_convocatoria(doc)
    _page_break(doc)
    _add_anexo_plantilla(doc)

    OUT_COMBINED.parent.mkdir(parents=True, exist_ok=True)
    OUT_PLANTILLA.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUT_COMBINED))

    # Copia del Anexo I solo (Parte B) para quien ya tenga las bases
    solo = _new_doc()
    _add_anexo_plantilla(solo)
    solo.save(str(OUT_PLANTILLA))

    return OUT_COMBINED, OUT_PLANTILLA


def main() -> None:
    combined, plantilla = build()
    print(f"Wrote {combined} ({combined.stat().st_size} bytes)")
    print(f"Wrote {plantilla} ({plantilla.stat().st_size} bytes)")


if __name__ == "__main__":
    main()
