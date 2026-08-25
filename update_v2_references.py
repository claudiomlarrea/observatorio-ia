#!/usr/bin/env python3
"""Update v2.0 Marco: annex cross-refs, subsection renumbering, Anexo K equivalencias."""
from docx import Document
from docx.text.paragraph import Paragraph
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re

PATHS = [
    "/Users/claudiolarrea/Library/CloudStorage/OneDrive-Personal/16 Secretaría de Investigación/60 Observatorio de Inteligencia Artificial/MARCO DE GOBERNANZA DE LA IA/Marco de Gobernanza, Ética y Uso Responsable de la IA - UCCuyo - v2.0 depurada.docx",
    "/Users/claudiolarrea/Documents/Observatorio/Marco de Gobernanza, Ética y Uso Responsable de la IA - v2.0 depurada.docx",
]

# v1 artículo -> v2 artículo
V1_TO_V2 = {
    1: 1, 2: 2, 3: 3, 4: 3, 5: 4, 6: 5, 7: 6, 8: 7, 9: 8, 10: 9, 11: 4, 12: 10,
    13: 11, 14: 12, 15: 13, 16: 14, 17: 14, 18: 15, 19: 16, 20: 17, 21: 18, 22: 19, 23: 20,
    24: 21, 25: 22, 26: 22, 27: 23, 28: 24, 29: 25, 30: 25, 31: 25, 32: 26, 33: 27,
    34: 28, 35: 28, 36: 29, 37: 29, 38: 30, 39: 30, 40: 30, 41: 31, 42: 31, 43: 32, 44: 32,
    45: 33, 46: 33, 47: 34, 48: 34, 49: 35, 50: 35, 51: 36, 52: 42, 53: 37, 54: 37,
    55: 38, 56: 38, 57: 38, 58: 39, 59: 39, 60: 40, 61: 40, 62: 41, 63: 41, 64: 42,
    65: 43, 66: 43, 67: 44, 68: 44, 69: 45, 70: 45, 71: 45, 72: 46, 73: 46, 74: 12, 75: 12,
    76: 47, 77: 47, 78: 48, 79: 49, 80: 49, 81: 50, 82: 50, 83: 51, 84: 51, 85: 51, 86: 52,
    87: 52, 88: 51, 90: 53, 91: 53, 92: 54, 93: 54, 94: 55, 95: 55, 96: 56, 97: 57,
    98: 58, 99: 58, 100: 59, 101: 59, 102: 59, 103: 59, 104: 60, 105: 60, 106: 61,
    108: 62, 109: 62, 110: 63, 111: 63, 112: 63, 113: 64, 114: 64, 115: 64,
    116: 65, 117: 65, 118: 65, 119: 66, 120: 66, 121: 66, 122: 66, 123: 66,
    125: 67, 126: 67, 127: 68, 128: 69, 129: 70, 130: 70, 131: 71, 132: 71,
    133: 72, 134: 72, 135: 72, 136: 73, 137: 73, 138: 73,
    140: 74, 141: 75, 142: 76, 143: 77, 144: 78, 145: 79, 146: 80, 147: 81, 148: 82, 149: 83, 150: 84,
}

# Renumber internal blocks in merged v2 articles: old_prefix -> new_prefix
SUB_RENUMBER = {
    12: [("14.", "12.")],
    15: [("18.", "15."), ("20.", "15.")],
    16: [("19.", "16.")],
    17: [("20.", "17.")],
    42: [("64.", "42."), ("52.", "42.")],
}


def renumber_in_article(text: str, art_num: int) -> str:
    if art_num not in SUB_RENUMBER:
        return text
    for old, new in SUB_RENUMBER[art_num]:
        text = text.replace(old, new)
    return text


# Solo artículos cuya numeración cambió entre v1 y v2
CHANGED_V1_TO_V2 = {k: v for k, v in V1_TO_V2.items() if k != v}


def replace_art_reference(match):
    prefix = match.group(1)
    num = int(match.group(2))
    sub = match.group(3) or ""
    degree = match.group(4) or ""
    if num not in CHANGED_V1_TO_V2:
        return match.group(0)
    new = CHANGED_V1_TO_V2[num]
    return f"{prefix}{new}{sub}{degree}"


ART_PATTERN = re.compile(
    r"((?:artículos?|Artículos?|arts?\.?)\s*)(\d+)(\.\d+)?(°)?",
    re.IGNORECASE,
)


def replace_text(text: str) -> str:
    text = ART_PATTERN.sub(replace_art_reference, text)
    # bare subsection refs in annexes like (art. 19.5)
    def sub_repl(m):
        major = int(m.group(1))
        minor = m.group(2)
        if major not in CHANGED_V1_TO_V2:
            return m.group(0)
        new_major = CHANGED_V1_TO_V2[major]
        return f"art. {new_major}{minor}"
    text = re.sub(r"\bart\.\s*(\d+)(\.\d+)", sub_repl, text, flags=re.I)
    return text


def replace_in_paragraph(p, annex_only=False):
    full = p.text
    if not full.strip():
        return False
    m = re.match(r"Artículo\s+(\d+)\s*[°.]?", full.strip())
    art_num = int(m.group(1)) if m and len(full.strip()) < 100 else None
    new = full
    if annex_only:
        new = replace_text(full)
    if art_num:
        new = renumber_in_article(new, art_num)
    if new == full:
        return False
    if p.runs:
        p.runs[0].text = new
        for r in p.runs[1:]:
            r.text = ""
    else:
        p.text = new
    return True


def insert_anexo_k(doc):
    """Insert Anexo K before BIBLIOGRAFÍA (once)."""
    for p in doc.paragraphs:
        if p.text.strip().startswith("Anexo K. Cuadro de equivalencias"):
            return
    bib_idx = None
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip() == "BIBLIOGRAFÍA":
            bib_idx = i
            break
    if bib_idx is None:
        return
    anchor = doc.paragraphs[bib_idx - 1]

    def ins(text, bold=False):
        nonlocal anchor
        el = OxmlElement("w:p")
        anchor._p.addnext(el)
        np = Paragraph(el, anchor._parent)
        r = np.add_run(text)
        r.bold = bold
        np.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        anchor = np
        return np

    rows = [
        ("1-6", "1-5", "Disposiciones generales; referentes unificados en Art. 4°"),
        ("7-12", "6-10", "Principios e identidad consolidados"),
        ("13-23", "11-20", "Gobernanza; órganos, matriz/mapa/registro integrados"),
        ("24-33", "21-27", "Usos permitidos, condicionados y prohibidos"),
        ("34-44", "28-32", "Docencia; reglas generales + Anexos E-J"),
        ("45-57", "33-38", "Investigación, extensión y transferencia"),
        ("58-71 / 52+64", "39-45", "Gestión; contratación unificada (Art. 42°)"),
        ("72-89", "46-52", "Responsabilidades por actor"),
        ("90-107", "53-61", "Derechos y garantías"),
        ("108-124", "62-66", "Formación y cultura"),
        ("125-139", "67-73", "Seguimiento; indicadores flexibles (Art. 69°)"),
        ("140-150", "74-84", "Disposiciones finales"),
    ]
    ins("Anexo K. Cuadro de equivalencias entre la versión 1.0 y la versión 2.0 depurada", bold=True)
    ins("")
    ins(
        "A efectos de lectura e implementación, las referencias de guías, resoluciones o actas "
        "que aludan a la numeración anterior del Marco podrán consultarse en este cuadro. "
        "Salvo indicación expresa, prevalece la numeración de la versión 2.0."
    )
    ins("")
    ins("Equivalencias principales (v1 → v2):", bold=True)
    for v1, v2, note in rows:
        ins(f"• Arts. {v1} → Arts. {v2}: {note}.")
    ins("")
    ins("Referencias operativas frecuentes:", bold=True)
    ins("• Definiciones: Art. 4° (v1) → Art. 3° (v2).")
    ins("• Evaluación de riesgo: Art. 19° (v1) → Art. 16° (v2).")
    ins("• Usos prohibidos: Art. 28° (v1) → Art. 24° (v2).")
    ins("• Contratación y proveedores: Arts. 52° y 64° (v1) → Art. 42° (v2).")
    ins("• Privacidad y datos: Art. 97° (v1) → Art. 57° (v2).")
    ins("• Órganos competentes: Arts. 14°, 74° y 75° (v1) → Art. 12° (v2).")
    ins("")


def process(path):
    doc = Document(path)
    n_annex = n_body = 0
    in_annex = False
    in_v2_body = False
    current_art = None
    for p in doc.paragraphs:
        t = p.text.strip()
        if "VERSIÓN 2.0" in t:
            in_v2_body = True
        if t == "ANEXOS":
            in_annex = True
            in_v2_body = False
        if t == "BIBLIOGRAFÍA":
            break
        m = re.match(r"Artículo\s+(\d+)\s*[°.]?", t)
        if m and len(t) < 100 and "\t" not in t:
            current_art = int(m.group(1))
        full = p.text
        if not full.strip():
            continue
        new = full
        if in_annex:
            new = replace_text(full)
        if in_v2_body and current_art:
            new = renumber_in_article(new, current_art)
        if new != full:
            if p.runs:
                p.runs[0].text = new
                for r in p.runs[1:]:
                    r.text = ""
            else:
                p.text = new
            if in_annex:
                n_annex += 1
            elif in_v2_body:
                n_body += 1
    insert_anexo_k(doc)
    doc.save(path)
    print(f"{path}: annex={n_annex}, body={n_body}, Anexo K added")


if __name__ == "__main__":
    for p in PATHS:
        process(p)
