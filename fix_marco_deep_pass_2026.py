#!/usr/bin/env python3
"""Deep editorial pass on Marco 2026: dedupe, citation style, leftover refs, boilers."""
from __future__ import annotations

import re
import shutil
from datetime import date
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph

from rewrite_unique_intros import INTROS

TARGET = Path(
    "/Users/claudiolarrea/Library/CloudStorage/OneDrive-Personal/"
    "16 Secretaría de Investigación/60 Observatorio de Inteligencia Artificial/"
    "MARCO DE GOBERNANZA DE LA IA/Documento FINAL/"
    "Marco de Gobernanza IA - UCCuyo 2026.docx"
)
BACKUP_DIR = TARGET.parent / "Versiones previas"
WORKSPACE_COPY = Path(
    "/Users/claudiolarrea/Documents/Observatorio/"
    "Marco de Gobernanza IA - UCCuyo FINAL.docx"
)

# Curated intros for remaining boilerplate heads (not always in INTROS).
EXTRA_INTROS = {
    "24.2. Prohibiciones en clave de derechos de las personas": (
        "Quedan vedados los usos de IA que lesionen dignidad, privacidad, igualdad o "
        "autonomía de las personas, o que instrumenten vigilancia, manipulación o "
        "exclusión ilegítima. Las enumeraciones que siguen concretan ese límite en "
        "clave de derechos: no son un catálogo cerrado, pero sí un estándar mínimo "
        "de inadmisibilidad en el ámbito universitario."
    ),
    "Artículo 45°. Roles, prohibiciones e implementación gradual": (
        "La puesta en marcha de sistemas de IA en la gestión exige roles claros, "
        "prohibiciones operativas e implementación gradual. No basta con adquirir "
        "una herramienta: hace falta responsable humano, controles y ritmo "
        "proporcional al riesgo. Las disposiciones siguientes ordenan esos deberes "
        "para evitar un despliegue de hecho, opaco o irreversible."
    ),
    "Artículo 77°. Interpretación y lagunas": INTROS.get(
        "Artículo 77°. Interpretación y lagunas",
        "El Marco se interpreta de forma armónica, privilegiando dignidad, integridad, "
        "transparencia, equidad y responsabilidad humana. Ante silencio, rige el criterio "
        "pro persona y de prevención del riesgo.",
    ),
}

EXACT = [
    # --- Art. 17.3: quitar meta-comentario ---
    (
        "Esta repetición operativa respecto del mapa no es redundancia vacía: asegura que ningún sistema entre en producción por vía de hecho, sin responsable ni posibilidad ulterior de suspensión o baja.",
        "Así se asegura que ningún sistema entre en producción por vía de hecho, sin responsable ni posibilidad ulterior de suspensión o baja.",
    ),
    # --- Glosario / anexos: refs residuales ---
    (
        "conforme al artículo 48° y a la cláusula general de los anexos.",
        "conforme al Artículo 78° y a la cláusula general de los anexos.",
    ),
    (
        "(principio 9.11)",
        "(Art. 9° y deber de contrastación)",
    ),
    (
        "(principio 9.12)",
        "(Art. 9°, lit. f)",
    ),
    (
        "(principio 9.17)",
        "(Arts. 22° y 42°)",
    ),
    (
        "(9.12)",
        "(Art. 9°, lit. f)",
    ),
    (
        "(9.12 y art. 8.f)",
        "(Art. 9°, lit. f)",
    ),
    (
        "cuando el cambio sea de alcance institucional (art. 12).",
        "cuando el cambio sea de alcance institucional (Art. 14°).",
    ),
    (
        "(art. 19.3 y 39.4)",
        "(Arts. 30° y 37°)",
    ),
    (
        "(Capítulo IV y art. 10.5)",
        "(Capítulo IV y Art. 24°)",
    ),
    (
        "herramienta autorizada (art. 22).",
        "herramienta autorizada (Arts. 21° y 22°).",
    ),
    (
        "incumplimiento (art. 25).",
        "incumplimiento (Arts. 30° y 56°).",
    ),
    (
        "(art. 26.f)",
        "(Arts. 32° y 72°)",
    ),
    (
        "definición operativa en el artículo 3°, inciso j).",
        "definición operativa en el Artículo 3°, inciso j).",
    ),
    (
        "Véase Delegación de decisiones y artículo 3°, inciso e).",
        "Véase Delegación de decisiones y Artículo 3°, inciso e).",
    ),
    (
        "Véanse las definiciones del artículo 3°, incisos a) y b).",
        "Véanse las definiciones del Artículo 3°, incisos a) y b).",
    ),
    (
        "Véase artículo 3°, inciso h).",
        "Véase Artículo 3°, inciso h).",
    ),
    (
        "definiciones del artículo 3° del Marco",
        "definiciones del Artículo 3° del Marco",
    ),
    (
        "definido en el artículo 3° y también aquí",
        "definido en el Artículo 3° y también aquí",
    ),
    (
        "definiciones del artículo 3° salvo reforma",
        "definiciones del Artículo 3° salvo reforma",
    ),
    (
        "ámbito del artículo 3° y por las reglas",
        "ámbito del Artículo 3° y por las reglas",
    ),
]


def set_text(p: Paragraph, text: str) -> None:
    el = p._element
    for child in list(el):
        if child.tag in (qn("w:r"), qn("w:hyperlink")):
            el.remove(child)
    p.add_run(text)


def delete_paragraph(p: Paragraph) -> None:
    el = p._element
    parent = el.getparent()
    if parent is not None:
        parent.remove(el)


def normalize_citations(text: str) -> str:
    """Unify short citations to Art./Arts. N°; keep 'Artículo' for full formal refs."""
    # arts. 3, 16 y 17 → already mostly fixed; normalize remaining lowercase art.
    def repl_art(m: re.Match) -> str:
        prefix = m.group(1)  # art. / arts.
        nums = m.group(2)
        # Keep subsection form art. 25.4 → Art. 25°, numeral 25.4 when clear?
        if re.match(r"^\d+\.\d+", nums):
            major = nums.split(".")[0]
            return f"Art. {major}°, numeral {nums}"
        # arts. X y Y / arts. X, Y y Z
        # Ensure ° on each standalone number
        def add_deg(n: str) -> str:
            n = n.strip()
            if n.endswith("°"):
                return n
            if re.match(r"^\d+$", n):
                return n + "°"
            return n

        # split on commas and y
        parts = re.split(r"\s*,\s*|\s+y\s+", nums)
        if len(parts) == 1:
            label = "Art." if prefix.lower().startswith("art.") and not prefix.lower().startswith("arts.") else "Art."
            # arts. → Arts.
            if prefix.lower().startswith("arts"):
                label = "Arts."
            return f"{label} {add_deg(parts[0])}"
        # multiple
        deg_parts = [add_deg(p) for p in parts]
        if len(deg_parts) == 2:
            return f"Arts. {deg_parts[0]} y {deg_parts[1]}"
        return "Arts. " + ", ".join(deg_parts[:-1]) + f" y {deg_parts[-1]}"

    # Only touch lowercase art./arts. (not Art./Arts. already normalized, not Artículo)
    text = re.sub(
        r"(?<![A-Za-zÁÉÍÓÚÑáéíóúñ])(arts?\.)\s+(\d+(?:\.\d+)?(?:\s*,\s*\d+(?:\.\d+)?)*(?:\s+y\s+\d+(?:\.\d+)?)?)",
        repl_art,
        text,
    )
    # Lowercase artículo N° → Artículo N°
    text = re.sub(r"\bartículo\s+(\d+)\s*°", r"Artículo \1°", text)
    return text


def apply_exact(doc: Document) -> int:
    n = 0
    for p in doc.paragraphs:
        st = p.style.name if p.style else ""
        if st.startswith("toc"):
            continue
        original = p.text
        text = original
        for old, new in EXACT:
            if old in text:
                text = text.replace(old, new)
        if text != original:
            set_text(p, text)
            n += 1
    return n


def dedupe_and_trim(doc: Document) -> dict:
    stats = {"art17_removed": 0, "incident_removed": 0, "boilers": 0}

    # Remove Art. 17 redundant second intro (exact thematic duplicate of first)
    for i, p in enumerate(list(doc.paragraphs)):
        t = p.text.strip()
        if t.startswith(
            "Todo sistema de IA institucional seguirá el ciclo de identificación, "
            "evaluación, autorización, registro, monitoreo y baja previsto en este Capítulo"
        ):
            # Keep only if no prior similar intro in Art 17 — delete this one
            delete_paragraph(p)
            stats["art17_removed"] += 1
            break

    # Remove duplicate incident paragraph (second identical copy)
    seen_incident = False
    for p in list(doc.paragraphs):
        t = p.text.strip()
        if t.startswith(
            "Todo incidente de seguridad vinculado a sistemas de IA deberá reportarse "
            "de inmediato por los canales institucionales"
        ):
            if seen_incident:
                delete_paragraph(p)
                stats["incident_removed"] += 1
            else:
                seen_incident = True

    # Replace leftover formulaic intros
    for i, p in enumerate(list(doc.paragraphs)):
        if p._element.getparent() is None:
            continue
        t = p.text.strip()
        if "pauta institucional para su aplicación concreta" not in t and (
            "Su objetivo es precisar el sentido del punto antes de las enumeraciones" not in t
        ):
            continue
        head = ""
        for k in range(i - 1, max(i - 10, -1), -1):
            pk = doc.paragraphs[k]
            if pk.style and pk.style.name in ("Heading 2", "Heading 3"):
                head = pk.text.strip().split("\t")[0]
                break
        replacement = EXTRA_INTROS.get(head) or INTROS.get(head)
        if not replacement:
            # generic clean from heading
            title = re.sub(r"^(Artículo\s+\d+\s*[°º]\.?\s*|\d+\.\d+\.?\s*)", "", head).strip()
            title_l = (title[:1].lower() + title[1:]) if title else "este punto"
            replacement = (
                f"Este apartado precisa el alcance de {title_l}. "
                f"Las enumeraciones posteriores son estándares mínimos y se aplican "
                f"de modo integrado, proporcional al riesgo y coherente con la "
                f"protección de las personas."
            )
        set_text(p, replacement)
        stats["boilers"] += 1

    return stats


def unify_citation_style(doc: Document) -> int:
    n = 0
    for p in doc.paragraphs:
        st = p.style.name if p.style else ""
        if st.startswith("toc") or st.startswith("Heading"):
            continue
        original = p.text
        if not original.strip():
            continue
        # Only process paras that still have lowercase art. or artículo
        if not re.search(r"\barts?\.\s+\d|\bartículo\s+\d", original):
            continue
        new = normalize_citations(original)
        if new != original:
            set_text(p, new)
            n += 1
    return n


def light_orthography(doc: Document) -> int:
    n = 0
    fixes = [
        (r"  +", " "),
        (r"\s+([.,;:)])", r"\1"),
        (r"\(Art\. (\d+)°\)", r"(Art. \1°)"),  # noop anchor
    ]
    for p in doc.paragraphs:
        st = p.style.name if p.style else ""
        if st.startswith("toc"):
            continue
        t = p.text
        new = t
        new = re.sub(r"  +", " ", new)
        new = re.sub(r"\s+([.,;:)])", r"\1", new)
        # anexos lowercase → Anexo when citing letter
        new = re.sub(r"\banexo ([A-K])\b", r"Anexo \1", new)
        if new != t:
            set_text(p, new)
            n += 1
    return n


def verify(doc: Document) -> list[str]:
    probs = []
    joined = "\n".join(p.text for p in doc.paragraphs)
    for bad in [
        "Esta repetición operativa",
        "pauta institucional para su aplicación concreta",
        "principio 9.",
        "art. 19.3",
        "art. 10.5",
        "artículo 48°",
        "101°",
        "144)",
    ]:
        if bad in joined:
            probs.append(bad)
    # duplicate incident count
    c = joined.count(
        "Todo incidente de seguridad vinculado a sistemas de IA deberá reportarse "
        "de inmediato por los canales institucionales"
    )
    if c > 1:
        probs.append(f"incident dup x{c}")
    # Art17 double intro
    c2 = joined.count(
        "Todo sistema de IA institucional seguirá el ciclo de identificación"
    )
    if c2 > 1:
        probs.append(f"art17 cycle intro x{c2}")
    return probs


def main() -> None:
    BACKUP_DIR.mkdir(parents=True, exist_ok=True)
    backup = (
        BACKUP_DIR
        / f"Marco de Gobernanza IA - UCCuyo 2026.pre-deep-pass-{date.today().isoformat()}.docx"
    )
    shutil.copy2(TARGET, backup)
    print("Backup:", backup)

    doc = Document(str(TARGET))

    print("Exact:", apply_exact(doc))
    print("Dedupe/boilers:", dedupe_and_trim(doc))
    print("Citation unify:", unify_citation_style(doc))
    print("Orthography:", light_orthography(doc))

    # Second citation pass after exact (some art. may remain)
    print("Citation unify 2:", unify_citation_style(doc))

    doc.save(str(TARGET))
    shutil.copy2(TARGET, WORKSPACE_COPY)

    doc2 = Document(str(TARGET))
    probs = verify(doc2)
    print("VERIFY:", probs or "OK")

    # Stats
    body = [
        p.text
        for p in doc2.paragraphs
        if p.style and not p.style.name.startswith("toc")
    ]
    joined = "\n".join(body)
    print(
        "Counts — Artículo°:",
        len(re.findall(r"Artículo\s+\d+\s*°", joined)),
        "| Art./Arts.°:",
        len(re.findall(r"Arts?\.\s+\d+\s*°", joined)),
        "| art. lowercase leftover:",
        len(re.findall(r"\barts?\.\s+\d+", joined)),
    )


if __name__ == "__main__":
    main()
