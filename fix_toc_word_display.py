#!/usr/bin/env python3
"""Rebuild Word TOC with dot leaders, page numbers, and clickable hyperlinks."""
from docx import Document
from docx.text.paragraph import Paragraph
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re
import subprocess
import os
import shutil

PATHS = [
    "/Users/claudiolarrea/Desktop/Marco de Gobernanza IA - UCCuyo FINAL.docx",
    "/Users/claudiolarrea/Documents/Observatorio/Marco de Gobernanza IA - UCCuyo FINAL.docx",
    "/Users/claudiolarrea/Library/CloudStorage/OneDrive-Personal/16 Secretaría de Investigación/60 Observatorio de Inteligencia Artificial/MARCO DE GOBERNANZA DE LA IA/Marco de Gobernanza IA - UCCuyo FINAL.docx",
    "/Users/claudiolarrea/Library/CloudStorage/OneDrive-Personal/16 Secretaría de Investigación/60 Observatorio de Inteligencia Artificial/MARCO DE GOBERNANZA DE LA IA/Documento FINAL/Marco de Gobernanza IA - UCCuyo FINAL.docx",
]


def delete_paragraph(p):
    p._element.getparent().remove(p._element)


def set_right_tab(p, pos=9062):
    pPr = p._element.get_or_add_pPr()
    tabs = pPr.find(qn("w:tabs"))
    if tabs is None:
        tabs = OxmlElement("w:tabs")
        pPr.append(tabs)
    else:
        for t in list(tabs):
            tabs.remove(t)
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "right")
    tab.set(qn("w:leader"), "dot")
    tab.set(qn("w:pos"), str(pos))
    tabs.append(tab)


def max_bookmark_id(doc):
    highest = 0
    for el in doc.element.body.iter():
        if el.tag.endswith("bookmarkStart") or el.tag.endswith("bookmarkEnd"):
            bid = el.get(qn("w:id"))
            if bid and bid.isdigit():
                highest = max(highest, int(bid))
    return highest


def remove_bookmarks(p):
    el = p._element
    for child in list(el):
        if child.tag.endswith("bookmarkStart") or child.tag.endswith("bookmarkEnd"):
            el.remove(child)


def bookmark_name(title, idx):
    if title.startswith("CAPÍTULO"):
        m = re.search(r"CAPÍTULO\s+([IVXLC]+)", title)
        return f"_Cap_{m.group(1) if m else idx}"
    if title.startswith("Artículo"):
        m = re.search(r"Artículo\s+(\d+)", title)
        return f"_Art_{m.group(1) if m else idx}"
    if title.startswith("Anexo"):
        m = re.search(r"Anexo\s+([A-K])", title)
        return f"_Anx_{m.group(1) if m else idx}"
    if title == "ANEXOS":
        return "_Anexos"
    if title == "BIBLIOGRAFÍA":
        return "_Bibliografia"
    return f"_Bm_{idx}"


def add_bookmark(p, name, bid):
    remove_bookmarks(p)
    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:id"), str(bid))
    start.set(qn("w:name"), name)
    end = OxmlElement("w:bookmarkEnd")
    end.set(qn("w:id"), str(bid))
    p._element.insert(0, start)
    p._element.append(end)


def add_text_line(anchor, text, center=False, bold=False):
    new_el = OxmlElement("w:p")
    anchor._p.addnext(new_el)
    p = Paragraph(new_el, anchor._parent)
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(text)
    r.font.name = "Times New Roman"
    r.bold = bold
    return p


def add_toc_line(anchor, title, page, style="toc 1", bookmark=None):
    new_el = OxmlElement("w:p")
    anchor._p.addnext(new_el)
    p = Paragraph(new_el, anchor._parent)
    try:
        p.style = style
    except Exception:
        pass
    set_right_tab(p)

    if bookmark:
        hyperlink = OxmlElement("w:hyperlink")
        hyperlink.set(qn("w:anchor"), bookmark)
        hyperlink.set(qn("w:history"), "1")

        r1 = OxmlElement("w:r")
        r_pr = OxmlElement("w:rPr")
        r_style = OxmlElement("w:rStyle")
        r_style.set(qn("w:val"), "Hyperlink")
        r_pr.append(r_style)
        r_fonts = OxmlElement("w:rFonts")
        r_fonts.set(qn("w:ascii"), "Times New Roman")
        r_fonts.set(qn("w:hAnsi"), "Times New Roman")
        r_pr.append(r_fonts)
        r1.append(r_pr)
        t1 = OxmlElement("w:t")
        t1.text = title
        r1.append(t1)
        hyperlink.append(r1)
        p._element.append(hyperlink)
    else:
        r1 = p.add_run(title)
        r1.font.name = "Times New Roman"

    rt = OxmlElement("w:r")
    rt.append(OxmlElement("w:tab"))
    p._element.append(rt)

    r2 = OxmlElement("w:r")
    r2_pr = OxmlElement("w:rPr")
    r2_fonts = OxmlElement("w:rFonts")
    r2_fonts.set(qn("w:ascii"), "Times New Roman")
    r2_fonts.set(qn("w:hAnsi"), "Times New Roman")
    r2_pr.append(r2_fonts)
    r2.append(r2_pr)
    t2 = OxmlElement("w:t")
    t2.text = str(page)
    r2.append(t2)
    p._element.append(r2)
    return p


def set_body_format(p):
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.line_spacing = 1.5


def is_heading_text(t):
    if t.startswith("CAPÍTULO ") and "\t" not in t:
        return True
    if re.match(r"Artículo\s+\d+\s*[°.]", t) and "\t" not in t and len(t) < 120:
        return True
    if re.match(r"Anexo [A-K]\.", t) and "\t" not in t and len(t) < 120:
        return True
    if t in ("ANEXOS", "BIBLIOGRAFÍA"):
        return True
    return False


def collect_headings(doc):
    headings = []
    in_body = False
    for p in doc.paragraphs:
        t = p.text.strip().split("\t")[0].strip()
        if "VERSIÓN 2.0" in t:
            in_body = True
            continue
        if in_body and t == "ANEXOS":
            headings.append(("cap1", "ANEXOS"))
            break
        if not in_body:
            continue
        if t.startswith("CAPÍTULO ") and "\t" not in t:
            headings.append(("cap1", t))
        elif re.match(r"Artículo\s+\d+\s*[°.]", t) and "\t" not in t and len(t) < 120:
            headings.append(("toc2", t))

    in_annex = False
    seen = set()
    for p in doc.paragraphs:
        t = p.text.strip().split("\t")[0].strip()
        if t == "ANEXOS":
            if in_annex:
                continue
            in_annex = True
            continue
        if in_annex and t == "BIBLIOGRAFÍA":
            headings.append(("cap1", "BIBLIOGRAFÍA"))
            break
        if in_annex and re.match(r"Anexo [A-K]\.", t) and "\t" not in t and len(t) < 120:
            if t in seen:
                continue
            seen.add(t)
            headings.append(("toc2", t))
    return headings


def add_body_bookmarks(doc):
    bookmark_map = {}
    bid = max_bookmark_id(doc) + 1
    in_body = in_annex = False
    seen_titles = set()

    for p in doc.paragraphs:
        t = p.text.strip().split("\t")[0].strip()
        if "VERSIÓN 2.0" in t:
            in_body = True
            continue
        if in_body and t == "ANEXOS":
            in_body = False
            in_annex = True
        if in_annex and t == "BIBLIOGRAFÍA":
            if t not in seen_titles:
                name = bookmark_name(t, bid)
                add_bookmark(p, name, bid)
                bookmark_map[t] = name
                bid += 1
                seen_titles.add(t)
            break

        if not (in_body or in_annex):
            continue
        if not is_heading_text(t):
            continue
        if t in seen_titles:
            continue
        name = bookmark_name(t, bid)
        add_bookmark(p, name, bid)
        bookmark_map[t] = name
        seen_titles.add(t)
        bid += 1

    return bookmark_map


def map_pages(path, headings):
    pdf = path.replace(".docx", ".pdf")
    if os.path.exists(pdf):
        os.remove(pdf)
    subprocess.run(
        ["soffice", "--headless", "--convert-to", "pdf", "--outdir", os.path.dirname(path), path],
        capture_output=True,
    )
    page_map = {}
    try:
        from pypdf import PdfReader

        pages = [(i + 1, (pg.extract_text() or "").replace("\n", " ")) for i, pg in enumerate(PdfReader(pdf).pages)]
        last = 4
        current_cap = None
        for kind, title in headings:
            if kind == "cap1":
                current_cap = title
                if title in ("ANEXOS", "BIBLIOGRAFÍA"):
                    for pg, txt in pages:
                        if pg >= last and title in txt:
                            page_map[title] = pg
                            last = pg
                            break
                continue
            needle = (title.split(".")[0] + ".") if title.startswith("Artículo") else title[:30]
            for pg, txt in pages:
                if pg >= last and needle in txt:
                    page_map[title] = pg
                    last = pg
                    if current_cap and current_cap.startswith("CAPÍTULO") and current_cap not in page_map:
                        page_map[current_cap] = pg
                    break
        caps = [t for k, t in headings if k == "cap1" and t.startswith("CAPÍTULO")]
        arts = [t for k, t in headings if k == "toc2" and t.startswith("Artículo")]
        ai = 0
        for cap in caps:
            if cap not in page_map:
                while ai < len(arts) and arts[ai] not in page_map:
                    ai += 1
                if ai < len(arts):
                    page_map[cap] = page_map[arts[ai]]
    except Exception as e:
        print("PDF map error:", e)
    return page_map


def style_body(doc):
    in_body = in_annex = False
    for p in doc.paragraphs:
        t = p.text.strip().split("\t")[0].strip()
        if "VERSIÓN 2.0" in t:
            in_body = True
            set_body_format(p)
            continue
        if in_body and t == "ANEXOS":
            in_body = False
            in_annex = True
            continue
        if in_annex and t == "BIBLIOGRAFÍA":
            break
        if in_body:
            if t.startswith("CAPÍTULO ") and "\t" not in t:
                p.style = "Heading 1"
            elif re.match(r"Artículo\s+\d+\s*[°.]", t) and len(t) < 120:
                p.style = "Heading 2"
            if t:
                set_body_format(p)
        elif in_annex and t:
            set_body_format(p)


def rebuild_toc(doc, headings, page_map, bookmark_map):
    res_i = body_i = None
    for i, p in enumerate(doc.paragraphs):
        t = p.text.strip()
        if t == "ÍNDICE RESUMIDO":
            res_i = i
        if "VERSIÓN 2.0" in t:
            body_i = i
            break
    for idx in range(body_i - 1, res_i, -1):
        delete_paragraph(doc.paragraphs[idx])

    for p in doc.paragraphs:
        if p.text.strip() == "ÍNDICE RESUMIDO":
            cur = p
            break

    body_caps = [t for k, t in headings if k == "cap1" and t.startswith("CAPÍTULO")]
    for cap in body_caps:
        cur = add_toc_line(cur, cap, page_map.get(cap, ""), "toc 1", bookmark_map.get(cap))
    cur = add_toc_line(cur, "ANEXOS", page_map.get("ANEXOS", ""), "toc 1", bookmark_map.get("ANEXOS"))
    cur = add_text_line(cur, "")

    cur = add_text_line(cur, "ÍNDICE GENERAL", center=True, bold=True)
    cur = add_text_line(
        cur,
        "Índice detallado del documento (capítulos y artículos). Al hacer clic en un apartado se desplaza a la sección correspondiente.",
    )
    cur = add_text_line(cur, "")

    for kind, title in headings:
        pg = page_map.get(title, "")
        style = "toc 1" if kind == "cap1" else "toc 2"
        cur = add_toc_line(cur, title, pg, style, bookmark_map.get(title))
    add_text_line(cur, "")


def main():
    for path in PATHS:
        if not os.path.exists(path):
            continue
        backup = path.replace(".docx", " - backup.docx")
        shutil.copy2(path, backup)
        doc = Document(path)
        headings = collect_headings(doc)
        style_body(doc)
        bookmark_map = add_body_bookmarks(doc)
        doc.save(path)
        page_map = map_pages(path, headings)
        doc = Document(path)
        rebuild_toc(doc, headings, page_map, bookmark_map)
        style_body(doc)
        doc.save(path)

        doc = Document(path)
        errs = sum(1 for p in doc.paragraphs if "Error" in p.text)
        toc = sum(1 for p in doc.paragraphs if p.style.name.startswith("toc"))
        toc2 = sum(1 for p in doc.paragraphs if p.style.name == "toc 2")
        links = sum(1 for el in doc.element.body.iter() if el.tag.endswith("hyperlink") and el.get(qn("w:anchor")))
        bms = sum(1 for el in doc.element.body.iter() if el.tag.endswith("bookmarkStart"))
        print(
            f"OK {os.path.basename(path)}: toc={toc}, articles={toc2}, "
            f"bookmarks={bms}, toc_links={links}, errors={errs}"
        )


if __name__ == "__main__":
    main()
