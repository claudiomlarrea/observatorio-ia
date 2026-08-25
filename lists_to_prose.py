#!/usr/bin/env python3
"""Convert short non-enumerative lettered lists into prose (Cap. body only)."""
from __future__ import annotations

import re
import shutil
from datetime import date

from docx import Document
from docx.text.paragraph import Paragraph

TARGET = (
    "/Users/claudiolarrea/Library/CloudStorage/OneDrive-Personal/"
    "16 Secretaría de Investigación/60 Observatorio de Inteligencia Artificial/"
    "MARCO DE GOBERNANZA DE LA IA/Documento FINAL/"
    "Marco de Gobernanza IA - UCCuyo 2026.docx"
)

# intro prefix -> replacement prose (intro + following a)-d) become one paragraph)
CONVERSIONS = [
    (
        "Toda delegación condicionada exige:",
        "Toda delegación condicionada exige finalidad explícita, supervisión humana efectiva, "
        "declaración o registro cuando corresponda, y posibilidad de revisión o impugnación "
        "cuando la decisión tenga efecto sobre personas.",
    ),
    (
        "Toda decisión humano-algorítmica de impacto relevante deberá tener un responsable humano identificable, con competencia para:",
        "Toda decisión humano-algorítmica de impacto relevante deberá tener un responsable "
        "humano identificable, con competencia para definir la finalidad del uso, validar, "
        "ajustar o rechazar la salida del sistema, explicar la decisión cuando sea requerido "
        "y responder institucionalmente por sus efectos.",
    ),
    (
        "La generación de datos sintéticos mediante IA es admisible con fines metodológicos o de prueba, siempre que:",
        "La generación de datos sintéticos mediante IA es admisible con fines metodológicos o "
        "de prueba, siempre que se declare su carácter sintético, no se presenten como datos "
        "reales, no permitan reidentificación indebida de personas y no se utilicen para eludir "
        "requisitos éticos o legales aplicables a datos reales.",
    ),
    (
        "El docente deberá indicar en la consigna:",
        "El docente deberá indicar en la consigna si la IA está permitida, condicionada o "
        "prohibida; qué tipo de asistencia se admite; si se exige declaración de uso; y qué "
        "se considerará falta a la integridad académica.",
    ),
    (
        "La Universidad procurará preservar control razonable sobre:",
        "La Universidad procurará preservar control razonable sobre sus datos institucionales "
        "y personales bajo custodia, la finalidad de su tratamiento, la posibilidad de auditar "
        "y discontinuar servicios, y la continuidad de procesos críticos frente a cambios de proveedor.",
    ),
    (
        "La analítica descriptiva sobre datos agregados y debidamente protegidos",
        "La analítica descriptiva sobre datos agregados y debidamente protegidos es admisible "
        "cuando no permita identificación indebida de personas si ello no es necesario, tenga "
        "finalidad institucional explícita, esté bajo responsabilidad de un área competente y "
        "no se utilice para estigmatizar a estudiantes, docentes o no docentes.",
    ),
    (
        "La Universidad llevará un registro agregado de incidentes relevantes",
        "La Universidad llevará un registro agregado de incidentes relevantes, consultas "
        "frecuentes y conflictividades recurrentes vinculadas a IA, a fin de detectar fallas "
        "sistémicas, mejorar guías y matrices de delegación, ajustar formación y controles, "
        "e informar al Rectorado y, cuando corresponda, al Consejo Superior.",
    ),
    (
        "Sin perjuicio de la confidencialidad legítima, la Universidad comunica",
        "Sin perjuicio de la confidencialidad legítima, la Universidad comunicará de modo "
        "periódico y comprensible los avances principales de implementación, las nuevas guías "
        "o criterios, los aprendizajes institucionales relevantes y los cambios normativos u "
        "operativos de interés general.",
    ),
    (
        "Podrán emitirse informes extraordinarios cuando:",
        "Podrán emitirse informes extraordinarios cuando surja un riesgo alto o sistémico, "
        "ocurra un incidente grave de seguridad, sesgo o integridad, un cambio tecnológico o "
        "normativo externo lo requiera, o el Rectorado o el Consejo Superior lo soliciten.",
    ),
    (
        "Frente a incumplimientos leves, de buena fe o derivados de desconocimi",
        "Frente a incumplimientos leves, de buena fe o derivados de desconocimiento, se "
        "priorizará la orientación correctiva, la capacitación específica, la aclaración de "
        "consignas y procedimientos, y el acompañamiento para la adecuación de prácticas.",
    ),
    (
        "Existe deber de cuidado cuando un actor universitario:",
        "Existe deber de cuidado cuando un actor universitario introduce una herramienta de IA "
        "en un proceso, autoriza su uso, supervisa a personas que la utilizan, o conoce o debió "
        "conocer un riesgo relevante y puede mitigarlo.",
    ),
    (
        "Dicho modelo se estructura sobre cuatro pilares:",
        "Dicho modelo se estructura sobre cuatro pilares: autoridad institucional, para decidir "
        "y aprobar; articulación técnica y ética, para orientar, evaluar y proponer; ejecución "
        "académica y administrativa, para aplicar criterios en cada unidad; y aprendizaje "
        "institucional, para registrar, revisar y mejorar continuamente las prácticas.",
    ),
    (
        "El presente Marco expresa, en el campo de la IA, el compromiso institucional de:",
        "El presente Marco expresa, en el campo de la IA, el compromiso institucional de formar "
        "profesionales con sólidos fundamentos científicos y valores éticos; promover el diálogo "
        "entre ciencia, tecnología, ética, derecho, educación y teología; integrar la innovación "
        "con responsabilidad social, integridad académica y cuidado de las personas; y evitar "
        "que la eficiencia técnica desplace el juicio humano, la libertad responsable y el "
        "sentido comunitario de la Universidad.",
    ),
    (
        "Los usos individuales de IA por docentes, estudiantes o personal, cuando no configuren un sistema institucional, se rigen por los",
        "Los usos individuales de IA por docentes, estudiantes o personal, cuando no configuren "
        "un sistema institucional, se rigen por los Capítulos IV y siguientes. No obstante, las "
        "prácticas extendidas “de hecho”, aunque no estén formalizadas, deberán identificarse en "
        "el mapa institucional; las zonas grises no equivalen a autorización implícita; y el "
        "Observatorio propondrá criterios de clarificación normativa cuando una práctica "
        "recurrente genere riesgos académicos, éticos o legales.",
    ),
    (
        "Ante duda sobre la admisibilidad de un uso de IA en una actividad de e",
        "Ante duda sobre la admisibilidad de un uso de IA en una actividad de enseñanza o "
        "evaluación, primará la protección de la integridad del aprendizaje y de la equidad; "
        "se preferirá la opción que preserve supervisión humana y transparencia; el docente "
        "podrá establecer una regla más restrictiva para su consigna, siempre que sea clara, "
        "razonable y comunicada oportunamente; y las dudas institucionales o recurrentes se "
        "elevarán al Observatorio para un criterio orientador de alcance general.",
    ),
    (
        "Los reportes se regirán por",
        "Los reportes se regirán por buena fe, confidencialidad razonable del trámite, "
        "protección frente a represalias, trámite oportuno y remisión al área competente "
        "según la materia (académica, laboral, de seguridad, ética o jurídica).",
    ),
    (
        "En la resolución de conflictos se ponderarán",
        "En la resolución de conflictos se ponderarán, de modo conjunto, la dignidad y los "
        "derechos de la persona, la integridad académica y científica, la proporcionalidad y "
        "la buena fe, la existencia o no de reglas claras previas, el nivel de riesgo y el daño "
        "causado o evitado, y la necesidad de aprendizaje institucional y de prevención futura.",
    ),
    (
        "A los fines de este Marco:",
        "A los fines de este Marco, la salida de un sistema de IA se considera un insumo y no "
        "un acto institucional autónomo; el acto institucional existe cuando una persona o "
        "autoridad lo adopta, confirma, ejecuta o tolera mediando un deber de control; no es "
        "argumento eximente sostener que “lo decidió el algoritmo”; y la complejidad técnica "
        "de un sistema aumenta, en lugar de disminuir, el deber de prudencia de quien lo despliega.",
    ),
    (
        "El uso de datos en investigaciones o proyectos con IA se rige por:",
        "El uso de datos en investigaciones o proyectos con IA se rige por la licitud y la "
        "finalidad legítima, la minimización y la proporcionalidad, la calidad y la seguridad, "
        "la confidencialidad y la limitación de acceso, y la responsabilidad demostrable del "
        "responsable del tratamiento.",
    ),
]


def delete_paragraph(p: Paragraph):
    p._element.getparent().remove(p._element)


def set_text(p: Paragraph, text: str):
    if p.runs:
        p.runs[0].text = text
        for r in p.runs[1:]:
            r.text = ""
    else:
        p.text = text


def body_start(doc: Document) -> int:
    for i, p in enumerate(doc.paragraphs):
        if "VERSIÓN 2026" in p.text or "VERSIÓN 2.0 DEPURADA" in p.text:
            return i
    return 0


def convert(doc: Document) -> int:
    n = 0
    start = body_start(doc)
    letter = re.compile(r"^[a-z]\)\s", re.I)
    paras = list(doc.paragraphs)
    i = start
    while i < len(paras):
        p = paras[i]
        if p.style.name.startswith("toc"):
            i += 1
            continue
        t = p.text.strip()
        if t.startswith("ANEXOS") and p.style.name == "Heading 1":
            break
        matched = None
        for prefix, new in CONVERSIONS:
            if t.startswith(prefix):
                matched = new
                break
        if not matched:
            i += 1
            continue
        # collect following lettered items
        j = i + 1
        items = []
        while j < len(paras):
            tj = paras[j].text.strip()
            if letter.match(tj):
                items.append(paras[j])
                j += 1
                continue
            break
        if len(items) < 2:
            i += 1
            continue
        set_text(p, matched)
        for item in items:
            delete_paragraph(item)
        n += 1
        # refresh list after deletes
        paras = list(doc.paragraphs)
        i += 1
    return n


def main():
    backup = TARGET.replace(".docx", f" - backup pre-prosa {date.today().isoformat()}.docx")
    shutil.copy2(TARGET, backup)
    doc = Document(TARGET)
    n = convert(doc)
    doc.save(TARGET)
    print(f"Backup: {backup}")
    print(f"Saved: {TARGET}")
    print(f"lists_to_prose: {n}")


if __name__ == "__main__":
    main()
