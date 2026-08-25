#!/usr/bin/env python3
"""Apply valid Maluf observations to Marco de Gobernanza IA - UCCuyo.docx"""
from docx import Document
from docx.text.paragraph import Paragraph
from docx.oxml import OxmlElement
import re
import shutil

SRC = (
    "/Users/claudiolarrea/Library/CloudStorage/OneDrive-Personal/"
    "16 Secretaría de Investigación/60 Observatorio de Inteligencia Artificial/"
    "MARCO DE GOBERNANZA DE LA IA/"
    "Marco de Gobernanza, Ética y Uso Responsable de la IA - UCCuyo.docx"
)
LOCAL = (
    "/Users/claudiolarrea/Documents/Observatorio/"
    "Marco de Gobernanza, Ética y Uso Responsable de la IA - UCCuyo.docx"
)
BACKUP = SRC.replace(".docx", " - backup pre-Maluf.docx")


def delete_paragraph(paragraph):
    el = paragraph._element
    el.getparent().remove(el)


def replace_para_text(p, new_text):
    if p.runs:
        p.runs[0].text = new_text
        for r in p.runs[1:]:
            r.text = ""
    else:
        p.text = new_text


def insert_after(paragraph, text):
    new_el = OxmlElement("w:p")
    paragraph._p.addnext(new_el)
    new_p = Paragraph(new_el, paragraph._parent)
    new_p.add_run(text)
    return new_p


def find_body_art(doc, num):
    pat = re.compile(rf"Artículo\s+{num}\s*[°.]")
    for i, p in enumerate(doc.paragraphs):
        t = p.text.strip()
        if pat.match(t) and "\t" not in t:
            return i
    return None


def art_range(doc, num):
    start = find_body_art(doc, num)
    if start is None:
        return None, None
    for j in range(start + 1, len(doc.paragraphs)):
        t = doc.paragraphs[j].text.strip()
        if (re.match(r"Artículo\s+\d", t) and "\t" not in t and len(t) < 100) or (
            t.startswith("CAPÍTULO ") and "\t" not in t
        ):
            return start, j
    return start, len(doc.paragraphs)


def fix_art19(doc):
    paras = doc.paragraphs
    dup_indices = []
    for i, p in enumerate(paras):
        t = p.text.strip()
        if t == "La evaluación considerará, entre otros factores:":
            prev = paras[i - 1].text if i else ""
            if "19.2." not in prev:
                dup_indices.append(i)
                for j in range(i + 1, min(i + 12, len(paras))):
                    tj = paras[j].text.strip()
                    if tj.startswith("19.3."):
                        break
                    if re.match(r"^[a-i]\) ", tj):
                        dup_indices.append(j)
        if t.startswith("19.4. Clasificación orientativa") and i > 0:
            prev = paras[i - 1].text.strip()
            if not prev.startswith("19.3.") and "19.3." not in prev:
                dup_indices.append(i)

    for i, p in enumerate(paras):
        t = p.text
        if t.startswith("19.3. Supuestos") and "19.4. Clasificación" in t:
            parts = t.split("\n19.4. Clasificación orientativa de riesgo\n", 1)
            replace_para_text(p, parts[0])
            break

    for idx in sorted(set(dup_indices), reverse=True):
        delete_paragraph(doc.paragraphs[idx])


def fix_art64(doc):
    paras = doc.paragraphs
    dup_indices = []
    for i, p in enumerate(paras):
        if p.text.strip() == "Los contratos o convenios procurarán incluir:":
            prev = paras[i - 1].text if i else ""
            if "64.3." in prev:
                continue
            dup_indices.append(i)
            for j in range(i + 1, min(i + 10, len(paras))):
                tj = paras[j].text.strip()
                if tj.startswith("64.4.") or tj.startswith("Artículo"):
                    break
                dup_indices.append(j)
    for idx in sorted(set(dup_indices), reverse=True):
        delete_paragraph(doc.paragraphs[idx])


def fix_art2(doc):
    for p in doc.paragraphs:
        if p.text.strip().startswith("d) Se aplica tanto a herramientas de uso individual"):
            insert_after(
                p,
                "e) Las obligaciones, limitaciones y criterios vinculantes de este Marco se aplican "
                "en la medida de las competencias de cada actor; las sanciones y medidas correctoras se "
                "articulan con los regímenes disciplinarios, académicos y laborales vigentes (Capítulo XII), "
                "pudiendo desarrollarse mediante reglamentaciones, guías y anexos operativos del Observatorio.",
            )
            return


def fix_art5(doc):
    for p in doc.paragraphs:
        if p.text.strip().startswith("Artículo 6"):
            p.insert_paragraph_before(
                "5.4. Jerarquía de referentes\n"
                "Para interpretar este Marco, se distinguirá: (i) normativa institucional aplicable "
                "con efectos vinculantes; (ii) referentes externos orientadores en materia de ética, "
                "derechos y gobernanza (Artículo 5.2 y Artículo 11°); y (iii) antecedentes comparados "
                "de otras universidades e instituciones, sin recepción automática ni alteración de "
                "la autonomía universitaria."
            )
            return


def fix_art11(doc):
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip() == "Artículo 11°. Alineación con marcos externos de referencia":
            if i + 1 < len(doc.paragraphs):
                replace_para_text(
                    doc.paragraphs[i + 1],
                    "Complementa y desarrolla el Artículo 5.2, ordenando los referentes externos conforme "
                    "a la jerarquía del Artículo 5.4. El presente Capítulo y el Marco en su conjunto se "
                    "alinean, de modo orientador y no excluyente, con los siguientes referentes:",
                )
            break
    for p in doc.paragraphs:
        t = p.text.strip()
        if t.startswith(
            "a) La Recomendación sobre la Ética de la Inteligencia Artificial de la UNESCO, en particular"
        ):
            replace_para_text(
                p,
                "a) Los referentes generales del Artículo 5.2, especialmente la Recomendación sobre la "
                "Ética de la Inteligencia Artificial de la UNESCO;",
            )
            break


def fix_art9_intro(doc):
    for p in doc.paragraphs:
        t = p.text.strip()
        if t.startswith(
            "El diseño, la adquisición, el uso, la evaluación y la supervisión de sistemas de IA"
        ):
            replace_para_text(
                p,
                t
                + " Los numerales 9.15 a 9.18 desarrollan criterios operativos derivados de los principios "
                "de transparencia, privacidad y protección institucional, sin alterar su carácter vinculante.",
            )
            return


def fix_instruments_17_18_20(doc):
    for p in doc.paragraphs:
        if p.text.strip().startswith("Dicha matriz formará parte de los anexos operativos"):
            replace_para_text(
                p,
                p.text.strip()
                + " Los datos consignados en la matriz serán la fuente primaria para evitar registros "
                "paralelos con el mapa (Artículo 18°) y el Registro institucional (Artículo 20.4).",
            )
            break
    for p in doc.paragraphs:
        if p.text.strip().startswith(
            "La Universidad contará con un Mapa institucional de decisiones humano-algorítmicas"
        ):
            replace_para_text(
                p,
                p.text.strip()
                + " El mapa complementa la Matriz del Artículo 17°: no duplicará los campos ya "
                "registrados allí, sino que aportará el relevamiento estratégico de decisiones críticas, "
                "impactos y zonas grises.",
            )
            break
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip().startswith("20.4. Registro institucional"):
            if i + 1 < len(doc.paragraphs):
                nxt = doc.paragraphs[i + 1]
                if "Registro se alimentará" not in nxt.text:
                    replace_para_text(
                        nxt,
                        nxt.text.strip()
                        + " El Registro se alimentará de la evaluación del Artículo 19° y de la Matriz "
                        "del Artículo 17°, sin exigir cargas duplicadas de información.",
                    )
            break


def fix_art32(doc):
    for p in doc.paragraphs:
        if p.text.strip().startswith("La autorización excepcional es de interpretación restrictiva"):
            replace_para_text(
                p,
                p.text.strip()
                + " En ningún caso alcanzará a los usos expresamente prohibidos en el Artículo 28°. "
                "Los usos inadmisibles de referencia del Artículo 19.5 solo podrán considerarse bajo "
                "esta vía cuando exista habilitación excepcional fundada, evaluación reforzada y base "
                "jurídica suficiente, sin contravenir las prohibiciones del Capítulo IV.",
            )
            return


def fix_art19_5(doc):
    for p in doc.paragraphs:
        t = p.text.strip()
        if t.startswith("Sin perjuicio de las prohibiciones del Capítulo IV") and "inadmisibles" in t:
            if "Artículo 28" not in t:
                replace_para_text(
                    p,
                    t
                    + " Ninguna habilitación excepcional podrá validar usos expresamente prohibidos "
                    "en el Artículo 28°.",
                )
            return


def fix_art52(doc):
    start, end = art_range(doc, 52)
    if start is None:
        return
    paras = doc.paragraphs
    for i in range(start, end):
        if paras[i].text.strip().startswith("Todo convenio, contrato o proyecto"):
            t = paras[i].text.strip()
            if not t.lower().startswith("sin perjuicio"):
                replace_para_text(paras[i], "Sin perjuicio de lo dispuesto en el Artículo 64°, " + t[0].lower() + t[1:])
            break
    to_delete = []
    in_522 = False
    for i in range(start, end):
        t = paras[i].text.strip()
        if t.startswith("52.2. Proveedores"):
            in_522 = True
            continue
        if in_522:
            if t.startswith("52.3."):
                break
            to_delete.append(i)
    for idx in sorted(to_delete, reverse=True):
        delete_paragraph(doc.paragraphs[idx])
    for i in range(start, end):
        if doc.paragraphs[i].text.strip().startswith("52.2. Proveedores"):
            insert_after(
                doc.paragraphs[i],
                "En materia de evaluación de proveedores, se aplicarán supletoriamente los requisitos "
                "de los Artículos 64.2 y 64.3.",
            )
            break


def fix_art74_75(doc):
    start74, end74 = art_range(doc, 74)
    if start74 is None:
        return
    for idx in range(end74 - 1, start74, -1):
        delete_paragraph(doc.paragraphs[idx])
    p = doc.paragraphs[start74]
    replace_para_text(
        p,
        "Artículo 74°. Responsabilidades del Consejo Superior y del Rectorado",
    )
    insert_after(
        p,
        "Las competencias del Consejo Superior y del Rectorado en materia de IA son las establecidas "
        "en el Artículo 14°. Las responsabilidades institucionales se ejercerán con carácter de "
        "rendición de cuentas, coordinación interáreas y respuesta oportuna ante riesgos relevantes, "
        "sin perjuicio de las atribuciones estatutarias de cada órgano.",
    )

    start75, end75 = art_range(doc, 75)
    if start75 is None:
        return
    for idx in range(end75 - 1, start75, -1):
        delete_paragraph(doc.paragraphs[idx])
    p = doc.paragraphs[start75]
    replace_para_text(p, "Artículo 75°. Responsabilidades del Observatorio de Inteligencia Artificial")
    insert_after(
        p,
        "Las funciones del Observatorio son las previstas en el Artículo 14.3. El Observatorio actúa "
        "con independencia técnica y ética dentro de su ámbito de competencia, sin absorber las "
        "responsabilidades de las unidades ejecutoras ni reemplazar a Decanos, docentes, investigadores "
        "o áreas de gestión en sus deberes propios.",
    )


def replace_sintesis(doc, num, replacement):
    start, end = art_range(doc, num)
    if start is None:
        return
    for idx in range(end - 1, start, -1):
        delete_paragraph(doc.paragraphs[idx])
    insert_after(doc.paragraphs[start], replacement)


def fix_sintesis(doc):
    replace_sintesis(
        doc,
        89,
        "Las reglas de responsabilidad de este Capítulo se desarrollan operativamente en los Anexos "
        "y guías del Observatorio de Inteligencia Artificial.",
    )
    replace_sintesis(
        doc,
        107,
        "Los derechos y garantías de este Capítulo se detallan en sus artículos precedentes y en los "
        "protocolos operativos del Observatorio.",
    )
    replace_sintesis(
        doc,
        124,
        "La formación y cultura institucional en IA se implementarán mediante programas, guías y "
        "Anexos E y F, en coherencia con las reglas generales de este Capítulo.",
    )
    replace_sintesis(
        doc,
        139,
        "El seguimiento y la mejora continua se ejecutarán conforme a los artículos 125° a 138° y a "
        "los informes que elabore el Observatorio.",
    )


def fix_art128(doc):
    start, end = art_range(doc, 128)
    if start is None:
        return
    for idx in range(end - 1, start, -1):
        delete_paragraph(doc.paragraphs[idx])
    p = doc.paragraphs[start]
    insert_after(
        p,
        "128.1. Propuesta y actualización\n"
        "El Observatorio propondrá, al menos anualmente o cuando la madurez institucional lo requiera, "
        "un conjunto de indicadores de cumplimiento e impacto, sujetos a aprobación del Rectorado. "
        "Dichos indicadores podrán actualizarse sin modificar el presente Marco.\n"
        "128.2. Indicadores orientativos de cumplimiento\n"
        "A título ejemplificativo, podrán considerarse:\n"
        "a) Porcentaje de unidades académicas con política de IA comunicada en asignaturas o carreras;\n"
        "b) Número de sistemas de IA institucionales registrados y clasificados por riesgo;\n"
        "c) Proporción de sistemas de riesgo medio/alto con evaluación de impacto vigente;\n"
        "d) Cantidad de docentes, estudiantes y gestores alcanzados por formación;\n"
        "e) Número de guías operativas publicadas y actualizadas;\n"
        "f) Tiempo de respuesta ante consultas e incidentes reportados;\n"
        "g) Grado de adecuación de contratos y proveedores a los requisitos del Marco.\n"
        "128.3. Indicadores orientativos de impacto\n"
        "A título ejemplificativo, podrán considerarse:\n"
        "a) Evolución de conflictos de integridad académica vinculados a IA;\n"
        "b) Percepción de claridad y utilidad de las reglas institucionales;\n"
        "c) Evidencia de mejora en prácticas de evaluación auténtica;\n"
        "d) Detección y corrección de sesgos o efectos diferenciales adversos;\n"
        "e) Reducción de usos informales de alto riesgo;\n"
        "f) Calidad del acompañamiento derivado de alertas tempranas, cuando existan;\n"
        "g) Producción de buenas prácticas, investigaciones y transferencias responsables.\n"
        "128.4. Uso de los indicadores\n"
        "Los indicadores no se emplearán para estigmatizar personas ni para ranking punitivo de unidades. "
        "Su finalidad es diagnóstica, formativa y de gobierno institucional.",
    )


def apply_all(path):
    doc = Document(path)
    fix_art19(doc)
    fix_art64(doc)
    fix_art2(doc)
    fix_art5(doc)
    fix_art11(doc)
    fix_art9_intro(doc)
    fix_instruments_17_18_20(doc)
    fix_art32(doc)
    fix_art19_5(doc)
    fix_art52(doc)
    fix_art74_75(doc)
    fix_sintesis(doc)
    fix_art128(doc)
    doc.save(path)


def verify(path):
    doc = Document(path)
    text = "\n".join(p.text for p in doc.paragraphs)
    checks = {
        "art19_dup_intro": text.count("La evaluación considerará, entre otros factores:"),
        "art19_dup_19_4": len(re.findall(r"19\.4\. Clasificación orientativa", text)),
        "art32_prohibid": "Artículo 28°" in text and "autorización excepcional" in text.lower(),
        "art74_remision": "Artículo 14°" in text,
        "art5_4": "5.4. Jerarquía de referentes" in text,
        "art128_propuesta": "128.1. Propuesta y actualización" in text,
    }
    return checks


if __name__ == "__main__":
    shutil.copy2(SRC, BACKUP)
    apply_all(SRC)
    apply_all(LOCAL)
    print("Backup:", BACKUP)
    print("Verify OneDrive:", verify(SRC))
    print("Verify Local:", verify(LOCAL))
