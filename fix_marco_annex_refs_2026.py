#!/usr/bin/env python3
"""Second-pass fix: glossary + annex cross-refs still pointing to v1 numbers
or to the wrong 2026 article with the same number."""
from __future__ import annotations

import re
import shutil
from datetime import date
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph

TARGET = Path(
    "/Users/claudiolarrea/Library/CloudStorage/OneDrive-Personal/"
    "16 Secretaría de Investigación/60 Observatorio de Inteligencia Artificial/"
    "MARCO DE GOBERNANZA DE LA IA/Documento FINAL/"
    "Marco de Gobernanza IA - UCCuyo 2026.docx"
)
BACKUP_DIR = TARGET.parent / "Versiones previas"
WORKSPACE_COPY = Path(
    "/Users/claudiolarrea/Documents/Observatorio/"
    "Marco de Gobernanza IA - UCCuyo FINAL.docx"
)

# Longest / most specific first.
EXACT = [
    # --- OOR packs first ---
    (
        "los arts. 42, 101 y 105 y por el Anexo H",
        "los Arts. 52° y 60° y por el Anexo H",
    ),
    (
        "Este Anexo opera los artículos 42°, 101°, 103°, 104° y 105°.",
        "Este Anexo opera los Arts. 52°, 59° y 60°.",
    ),
    (
        "(arts. 8, 16, 28, 62 y 144)",
        "(Arts. 9°, 14°, 24°, 41° y 78°)",
    ),
    (
        "(arts. 8, 28, 62 y 95)",
        "(Arts. 9°, 24°, 41° y 55°)",
    ),
    (
        "(arts. 7.17, 64 y 97)",
        "(Arts. 42°, 44° y 57°)",
    ),
    (
        "(arts. 14, 28 y 98)",
        "(Arts. 16°, 24° y 58°)",
    ),
    (
        "(arts. 30 y 98)",
        "(Arts. 40° y 41°)",
    ),
    (
        "(arts. 7.14, 64 y 6)",
        "(Arts. 42° y 44°)",
    ),
    (
        "(arts. 3, 64 y 82)",
        "(Arts. 3° y 42°)",
    ),
    (
        "(arts. 3, 38 y 39)",
        "(Arts. 3°, 30°, 34° y 37°)",
    ),
    (
        "(arts. 3, 16 y 17; Anexo B)",
        "(Arts. 3° y 14°; Anexo B)",
    ),
    (
        "(arts. 14.5 y 28)",
        "(Art. 24°)",
    ),
    (
        "(arts. 14.6, 64 y 65)",
        "(Arts. 16° y 42°)",
    ),
    # --- Glosario A ---
    (
        "Analítica de aprendizaje. Uso de datos sobre procesos educativos para describir, predecir o alertar sobre trayectorias, desempeño o riesgo de desvinculación. En la Universidad solo es admisible como instrumento de cuidado y acompañamiento, no como mecanismo automático de sanción, exclusión o etiquetado estigmatizante (arts. 30).",
        "Analítica de aprendizaje. Uso de datos sobre procesos educativos para describir, predecir o alertar sobre trayectorias, desempeño o riesgo de desvinculación. En la Universidad solo es admisible como instrumento de cuidado y acompañamiento, no como mecanismo automático de sanción, exclusión o etiquetado estigmatizante (Arts. 40° y 41°).",
    ),
    (
        "no constituye base jurídica (art. 38).",
        "no constituye base jurídica (Arts. 35° y 57°).",
    ),
    (
        "El modelo institucional se desarrolla en el Anexo C (art. 25).",
        "El modelo institucional se desarrolla en el Anexo C (Art. 30°).",
    ),
    (
        "debido proceso (art. 19.5).",
        "debido proceso (Arts. 31° y 56°).",
    ),
    (
        "(art. 3, inc. k, y art. 14; Anexo D)",
        "(art. 3, inc. k, y Art. 16°; Anexo D)",
    ),
    (
        "(art. 3, inc. m, y art. 37)",
        "(art. 3, inc. m, y Art. 54°)",
    ),
    (
        "El plagio algorítmico —presentar como propio un producto sustancialmente generado por IAG sin declaración ni elaboración suficiente— vulnera la integridad (art. 25).",
        "El plagio algorítmico —presentar como propio un producto sustancialmente generado por IAG sin declaración ni elaboración suficiente— vulnera la integridad (Arts. 30° y 56°).",
    ),
    (
        "por las reglas diferenciadas del Artículo 24°.",
        "por las reglas diferenciadas del Capítulo V y del Anexo J.",
    ),
    (
        "supervisión (art. 46).",
        "supervisión (Art. 73°).",
    ),
    (
        "(art. 25.2)",
        "(Arts. 25° y 57°)",
    ),
    (
        "pueden integrarse a la declaración de uso (art. 25).",
        "pueden integrarse a la declaración de uso (Art. 30°).",
    ),
    (
        "controles específicos (art. 14).",
        "controles específicos (Art. 16°).",
    ),
    (
        "proporcionalidad (art. 15).",
        "proporcionalidad (Art. 18°).",
    ),
    (
        "J.4. Circuito de información a representantes legales (Artículo 29°)",
        "J.4. Circuito de información a representantes legales (Art. 57° y Anexo J)",
    ),
    (
        "Representantes legales. Padres, madres, tutores u otros sujetos que, conforme a la normativa aplicable, ejercen la representación de un NNA. Deben ser informados, en lenguaje comprensible, cuando un sistema institucional trate datos de menores en niveles preuniversitarios (Artículo 29°).",
        "Representantes legales. Padres, madres, tutores u otros sujetos que, conforme a la normativa aplicable, ejercen la representación de un NNA. Deben ser informados, en lenguaje comprensible, cuando un sistema institucional trate datos de menores en niveles preuniversitarios (Art. 57° y Anexo J).",
    ),
    # --- Anexo B ---
    (
        "los artículos 12°, 17° y 18° del Marco",
        "los Arts. 14°, 15° y 17° del Marco",
    ),
    (
        "automatización plena (art. 31),",
        "automatización plena (Art. 41°),",
    ),
    (
        "el mapa de decisiones críticas del artículo 13°",
        "el mapa de decisiones críticas del Artículo 15°",
    ),
    (
        "relación con los artículos 12° a 18°, conforme al artículo 46°.",
        "relación con los Arts. 14° a 17°, conforme al Artículo 72°.",
    ),
    # --- Anexo C ---
    (
        "declaración previsto en el Artículo 25°.",
        "declaración previsto en el Artículo 30°.",
    ),
    (
        "cuando el cambio sea de mera forma (art. 19.3).",
        "cuando el cambio sea de mera forma.",
    ),
    # --- Anexo D ---
    (
        "D.3. Criterios de valoración (Artículo 12°)",
        "D.3. Criterios de valoración (Artículo 16°)",
    ),
    (
        "D.4. Evaluación reforzada (Artículo 12°)",
        "D.4. Evaluación reforzada (Artículo 16°)",
    ),
    (
        "D.5. Pantalla de usos inadmisibles (Artículo 21° y Cap. IV)",
        "D.5. Pantalla de usos inadmisibles (Artículo 24° y Cap. IV)",
    ),
    (
        "D.6. Modelos fundacionales o IA de propósito general (Artículo 12°)",
        "D.6. Modelos fundacionales o IA de propósito general (Artículos 16° y 42°)",
    ),
    (
        "D.7. Calificación de riesgo y vía de autorización (arts. 14.4 y 20.3)",
        "D.7. Calificación de riesgo y vía de autorización (Arts. 16° y 17°)",
    ),
    (
        "hasta cumplir el ciclo del artículo 14° y, cuando corresponda, la inscripción en el registro del Anexo I. Las actualizaciones se versionan conforme al artículo 46°.",
        "hasta cumplir el ciclo del Artículo 17° y, cuando corresponda, la inscripción en el registro del Anexo I. Las actualizaciones se versionan conforme al Artículo 72°.",
    ),
    # --- Anexo E ---
    (
        "el Capítulo V y los artículos 8, 31, 36, 41 y 42 a una pauta de bolsillo",
        "el Capítulo V (Arts. 28° a 32°) y, en lo pertinente, los Arts. 9°, 24° y 41° a una pauta de bolsillo",
    ),
    (
        "se consulta al Observatorio (art. 22.f).",
        "se consulta al Observatorio (Art. 12°).",
    ),
    (
        "E.4. En el aula y en el programa de la asignatura (art. 25)",
        "E.4. En el aula y en el programa de la asignatura (Arts. 28° a 31°)",
    ),
    (
        "crítica de salidas de IA (art. 25).",
        "crítica de salidas de IA (Art. 31°).",
    ),
    (
        "Nada de legajos, notas, trabajos inéditos, salud o datos de menores en IAG pública (Artículo 25°). En niveles preuniversitarios: acompañamiento docente y reglas más estrictas (art. 18.2 y Anexo J).",
        "Nada de legajos, notas, trabajos inéditos, salud o datos de menores en IAG pública (Art. 25°). En niveles preuniversitarios: acompañamiento docente y reglas más estrictas (Anexo J).",
    ),
    # --- Anexo G ---
    (
        "Este Anexo desarrolla el Artículo 25° y la escala del numeral 31.6.",
        "Este Anexo desarrolla el Artículo 30° y la escala del numeral 31.6.",
    ),
    (
        "la IA no puede sustituir la evidencia de aprendizaje evaluada (Artículo 19°).",
        "la IA no puede sustituir la evidencia de aprendizaje evaluada (Art. 31°).",
    ),
    (
        "las dudas, por el artículo 26° y el Observatorio.",
        "las dudas, por el Artículo 32° y el Observatorio.",
    ),
    # --- Anexo H ---
    (
        "(art. 25.1)",
        "(Arts. 24° y 60°)",
    ),
    (
        "puede generar responsabilidad (art. 42.c).",
        "puede generar responsabilidad (Art. 52°).",
    ),
    (
        "Este protocolo se versiona conforme al artículo 46°.",
        "Este protocolo se versiona conforme al Artículo 72°.",
    ),
    # --- Anexo I ---
    (
        "Este Anexo opera los artículos 14.4 y 65°.",
        "Este Anexo opera los Arts. 15° y 17°.",
    ),
    (
        "vía de autorización del artículo 14.3.",
        "vía de autorización del Artículo 17°.",
    ),
    (
        "puede ordenar revisión, suspensión o baja (Artículo 12°).",
        "puede ordenar revisión, suspensión o baja (Art. 17°).",
    ),
    (
        "sin ciclo del artículo 14°, Anexo D y esta ficha. Los pilotos del artículo 46° también se registran",
        "sin ciclo del Artículo 17°, Anexo D y esta ficha. Los pilotos del Artículo 73° también se registran",
    ),
    # --- Anexo J ---
    (
        "Este Anexo desarrolla el Artículo 24° y el Artículo 38°.",
        "Este Anexo desarrolla el Capítulo V (Arts. 28° y 32°) y las garantías de protección de NNA.",
    ),
    (
        "sin relajar el Artículo 24°.",
        "sin relajar el Capítulo V ni este Anexo.",
    ),
    (
        "se versionan conforme al artículo 46°, sin relajar",
        "se versionan conforme al Artículo 72°, sin relajar",
    ),
]


def set_text(p: Paragraph, text: str) -> None:
    p_el = p._element
    for child in list(p_el):
        if child.tag in (qn("w:r"), qn("w:hyperlink")):
            p_el.remove(child)
    p.add_run(text)


def apply_exact(doc: Document) -> tuple[int, list[str]]:
    n = 0
    unused = []
    # Track which patterns matched
    matched = set()
    for p in doc.paragraphs:
        st = p.style.name if p.style else ""
        if st.startswith("toc"):
            continue
        original = p.text
        text = original
        for old, new in EXACT:
            if old in text:
                text = text.replace(old, new)
                matched.add(old)
        if text != original:
            set_text(p, text)
            n += 1
    for old, _ in EXACT:
        if old not in matched:
            unused.append(old[:80])
    return n, unused


def fuzzy_glossary_fixes(doc: Document) -> int:
    """Fallback regex for entries whose exact wording drifted."""
    n = 0
    rules = [
        # Asistente de reunión — any remaining art. 31 at end of that entry
        (
            r"(Asistente de reunión o .note taker\..{0,400}?)\(art\. 31\)\.",
            r"\1(Art. 41°).",
        ),
        (
            r"(Asistente virtual o chatbot institucional\..{0,400}?)\(art\. 31\)\.",
            r"\1(Art. 41°).",
        ),
        (
            r"(Deepfake o contenido sintético engañoso\..{0,400}?)\(art\. 31\)\.",
            r"\1(Arts. 41° y 60°).",
        ),
        (
            r"(Analítica de aprendizaje\..{0,350}?)\(arts?\. 30\)\.",
            r"\1(Arts. 40° y 41°).",
        ),
        (
            r"(Declaración de uso de IA\..{0,400}?)\(art\. 25\)\.",
            r"\1(Art. 30°).",
        ),
        (
            r"arts?\. 14\.6[^\)]{0,40}",
            "Arts. 16° y 42°",
        ),
        # Representantes — avoid double-fixing
        (
            r"\(Artículo 29°\)",
            "(Art. 57° y Anexo J)",
        ),
        # art. 40 leftover in H if any
        (
            r"\(art\. 40\)",
            "(Art. 60°)",
        ),
        # residual OOR
        (
            r"arts?\.\s*42,\s*101\s*y\s*105",
            "Arts. 52° y 60°",
        ),
        (
            r"artículos\s+42°,\s*101°,\s*103°,\s*104°\s*y\s*105°",
            "Arts. 52°, 59° y 60°",
            re.I,
        ),
        (
            r"arts?\.\s*8,\s*16,\s*28,\s*62\s*y\s*144",
            "Arts. 9°, 14°, 24°, 41° y 78°",
        ),
        (
            r"arts?\.\s*30\s*y\s*98",
            "Arts. 40° y 41°",
        ),
        (
            r"arts?\.\s*14,\s*28\s*y\s*98",
            "Arts. 16°, 24° y 58°",
        ),
        (
            r"arts?\.\s*8,\s*28,\s*62\s*y\s*95",
            "Arts. 9°, 24°, 41° y 55°",
        ),
        (
            r"arts?\.\s*7\.17,\s*64\s*y\s*97",
            "Arts. 42°, 44° y 57°",
        ),
    ]
    for p in doc.paragraphs:
        st = p.style.name if p.style else ""
        if st.startswith("toc"):
            continue
        text = p.text
        new = text
        for rule in rules:
            if len(rule) == 3:
                pat, repl, flags = rule
                new = re.sub(pat, repl, new, flags=flags)
            else:
                pat, repl = rule
                new = re.sub(pat, repl, new, flags=re.I | re.S)
        if new != text:
            set_text(p, new)
            n += 1
    return n


def verify(doc: Document) -> list[str]:
    problems = []
    joined = "\n".join(p.text for p in doc.paragraphs)
    for bad in [
        "101°",
        "103°",
        "104°",
        "105°",
        " y 144)",
        " y 95)",
        " y 97)",
        " y 98)",
        "arts. 42, 101",
        "artículo 13° inventaría",
        "artículo 14° y, cuando corresponda, la inscripción",
        "artículos 42°, 101°",
    ]:
        if bad in joined:
            problems.append(bad)
    # OOR in art-citation context
    for i, p in enumerate(doc.paragraphs):
        st = p.style.name if p.style else ""
        if st.startswith("toc"):
            continue
        t = p.text
        for m in re.finditer(
            r"(?:arts?\.|artículos?|Artículo|Art\.)[^\n]{0,50}?(\d{2,3})",
            t,
            re.I,
        ):
            n = int(m.group(1))
            if n > 84:
                problems.append(f"[{i}] OOR {n}: {t.strip()[:120]}")
    return problems


def main() -> None:
    BACKUP_DIR.mkdir(parents=True, exist_ok=True)
    backup = BACKUP_DIR / f"Marco de Gobernanza IA - UCCuyo 2026.pre-annex-refs-{date.today().isoformat()}.docx"
    shutil.copy2(TARGET, backup)
    print("Backup:", backup)

    doc = Document(str(TARGET))
    n, unused = apply_exact(doc)
    print(f"Exact replacements in {n} paragraphs")
    if unused:
        print(f"Unused patterns ({len(unused)}):")
        for u in unused[:40]:
            print("  -", u)

    n2 = fuzzy_glossary_fixes(doc)
    print(f"Fuzzy/regex paragraphs: {n2}")

    doc.save(str(TARGET))
    shutil.copy2(TARGET, WORKSPACE_COPY)

    doc2 = Document(str(TARGET))
    probs = verify(doc2)
    print("Remaining problems:", probs or "OK")


if __name__ == "__main__":
    main()
