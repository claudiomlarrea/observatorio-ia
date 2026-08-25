#!/usr/bin/env python3
"""Replace monotonous list-intros with uniquely opened 60–80 word paragraphs."""
from __future__ import annotations

import re
import sys

sys.path.insert(0, "/Users/claudiolarrea/Documents/Observatorio")
from docx import Document
from rewrite_unique_intros import INTROS as OLD, TARGET, set_text, wc

# Unique openings (no “Este apartado / Este Capítulo / Las viñetas”).
OPEN = {
    "3.2. Funciones sustantivas y de gestión": (
        "Docencia, investigación, extensión y administración forman un mismo proyecto "
        "institucional: la IA que entra en cualquiera de ellas queda sujeta a la misión, "
        "a la integridad académica y al cuidado de las personas."
    ),
    "3.3. Sujetos alcanzados": (
        "No alcanza con regular a docentes y estudiantes. Autoridades, personal de gestión, "
        "investigadores, contratistas y quien intervenga en actividades universitarias con "
        "sistemas algorítmicos quedan vinculados por las mismas salvaguardas, en lo pertinente."
    ),
    "3.4. Tecnologías comprendidas": (
        "Lo decisivo no es que la herramienta se llame chatbot. Quedan comprendidos también "
        "sistemas de recomendación, clasificación, predicción, analítica de aprendizaje y "
        "cualquier dispositivo algorítmico con efectos sobre personas o procesos."
    ),
    "4.1. Normativa e instrumentos institucionales": (
        "Estatutos, plan estratégico, ordenanzas de docencia e investigación y la resolución "
        "de creación del Observatorio son el piso sobre el cual se lee este Marco, no un "
        "cuerpo paralelo ni ajeno a la identidad universitaria."
    ),
    "4.2. Referentes externos de orientación": (
        "UNESCO, lineamientos de IA fiable, magisterio de la Iglesia y principios de confianza "
        "tecnológica iluminan el sentido de estas reglas, sin incorporar automáticamente "
        "normas extranjeras ni alterar la jerarquía interna."
    ),
    "4.4. Ámbito internacional": (
        "En el plano comparado se recogen recomendaciones de una IA digna de confianza, "
        "transparente y centrada en la persona. Su recepción es selectiva y compatible con "
        "el Estatuto y el carácter confesional de la Universidad."
    ),
    "4.5. Ámbito nacional": (
        "En Argentina el texto se alinea, en lo que corresponda, con IA fiable, datos "
        "personales, educación superior, propiedad intelectual e integridad académica, para "
        "evitar contradicciones con el derecho interno."
    ),
    "4.6. Ámbito eclesial y de universidades católicas": (
        "Como universidad católica, la institución lee la tecnología a la luz del magisterio "
        "sobre la persona, el bien común y la educación. La algor-ética orienta límites y "
        "prudencia, sin sustituir la normativa universitaria ni reducir el Marco a técnica."
    ),
    "Artículo 5°. Principios de interpretación": (
        "Dudas, tensiones y lagunas no se resuelven por conveniencia técnica. La persona, "
        "la responsabilidad humana, la proporcionalidad, la transparencia, la prevención y "
        "el ideario católico prevalecen cuando la eficiencia erosione dignidad o integridad."
    ),
    "Artículo 9°. Principio rector de no sustitución del juicio humano": (
        "Acreditación, autoría, sanciones y decisiones de impacto no pueden quedar libradas "
        "a una salida algorítmica. En esta Universidad la IA asiste; no reemplaza el juicio "
        "docente, científico, ético ni institucional sobre las personas."
    ),
    "Artículo 10°. Finalidad formativa de la gobernanza de la IA": (
        "Gobernar no es solo controlar. Se trata de formar para un uso crítico, ético y "
        "creativo, fortalecer el discernimiento frente a recomendaciones opacas y evitar que "
        "la eficiencia desplace el vínculo pedagógico."
    ),
    "12.4. Unidades Académicas y autoridades de gestión académica": (
        "Facultades, escuelas, institutos y direcciones de carrera son el lugar donde la "
        "norma se vuelve práctica. Aplican estas reglas, pueden ser más restrictivas —nunca "
        "más permisivas— y no delegan decisiones académicas de impacto a un sistema."
    ),
    "12.5. Áreas de Tecnologías de la Información y sistemas": (
        "Sin TIC no hay seguridad, interoperabilidad ni resguardo. Evalúan lo técnico, "
        "revisan proveedores, implementan controles y alertan vulnerabilidades, en "
        "coordinación con el Observatorio y el Rectorado, sin gobernar solas la IA."
    ),
    "12.6. Secretarías y áreas de gestión": (
        "Becas, trámites y sistemas de información no quedan fuera del mapa. Investigación, "
        "extensión, administración, bienestar y educación a distancia identifican procesos "
        "propios, aplican la matriz de delegación y aseguran un responsable humano."
    ),
    "12.7. Comunidad universitaria": (
        "Quien enseña, estudia, investiga o trabaja aquí no es destinatario pasivo. Debe usar "
        "la IA conforme al Marco, declarar cuando corresponda, no cargar datos sensibles en "
        "herramientas no autorizadas y reportar incidentes o sesgos por los canales "
        "institucionales. El desconocimiento no exime de esos deberes mínimos de cuidado."
    ),
    "14.1. Decisiones o tareas no delegables": (
        "Evaluación final acreditante, juicio pedagógico sustantivo, sanciones y decisiones "
        "laborales de impacto exigen una persona competente que decida y responda. La "
        "herramienta podrá aportar insumos; no sustituye el acto ni diluye la imputación."
    ),
    "14.2. Delegación condicionada": (
        "Entre lo vedado y lo rutinario hay un tramo de asistencia: corrección orientativa, "
        "patrones, sugerencias o priorizaciones preliminares. Solo es admisible con "
        "supervisión efectiva, transparencia, trazabilidad y revisión cuando afecte personas."
    ),
    "14.3. Tareas delegables": (
        "Análisis descriptivo de datos agregados, automatizaciones de soporte o resúmenes "
        "de información no confidencial pueden delegarse con controles proporcionales. Si "
        "la tarea determina derechos, evaluaciones o sanciones, deja de ser de bajo riesgo."
    ),
    "15.2. Contenidos mínimos": (
        "Un inventario informal no sirve. El mapa debe visibilizar, con un mínimo común, "
        "dónde la IA ya interviene o puede intervenir: decisiones críticas, unidades, "
        "impacto, riesgo y zonas grises, sin duplicar los campos de la Matriz."
    ),
    "15.3. Decisiones críticas de referencia": (
        "Aprobar, desaprobar, acreditar, asignar becas, alertar abandono o evaluar desempeño "
        "pueden alterar trayectorias. Si un sistema asiste esos actos, la supervisión humana "
        "y el registro dejan de ser opcionales. El listado es de referencia, no taxativo."
    ),
    "15.7. Decisión de autorización": (
        "El bajo riesgo puede resolverse en el área con registro; el medio exige TIC y "
        "opinión del Observatorio; el alto queda en el Rectorado, con dictamen y controles "
        "reforzados. Ni parálisis en lo inocuo ni despliegue de hecho de lo sensible."
    ),
    "16.2. Criterios de valoración": (
        "Severidad del daño, escala de afectados, opacidad, sensibilidad de datos, sesgos, "
        "supervisión humana, trazabilidad, dependencia de proveedores y efectos sobre menores "
        "o personas vulnerables deben ponderarse por escrito, no por un sí o no intuitivo."
    ),
    "16.3. Supuestos de evaluación reforzada": (
        "Perfiles de personas, datos sensibles, biométricos o de menores, monitoreo intensivo, "
        "trayectorias académicas o laborales, transferencias internacionales e infraestructuras "
        "de terceros con acceso potencial no se resuelven con una evaluación ordinaria."
    ),
    "17.3. Decisión de autorización": (
        "Ningún sistema institucional entra en producción por vía de hecho. El área autoriza "
        "lo de bajo impacto con registro; lo medio requiere conformidad técnica y opinión del "
        "Observatorio; lo alto se reserva al Rectorado, con posibilidad ulterior de baja."
    ),
    "Artículo 22°. Finalidades legítimas y usos permitidos": (
        "Mejorar enseñanza y tutoría, fortalecer investigación y transferencia, apoyar la "
        "gestión sin sustituir decisiones de impacto: hay finalidades que dan certeza. "
        "Permitido no significa irrestricto: siguen la integridad, la privacidad y la supervisión."
    ),
    "22.1. En docencia y aprendizaje": (
        "En el aula la IA puede apoyar la preparación de clases, la explicación de conceptos "
        "y la práctica formativa si no sustituye la evidencia de aprendizaje ni el juicio "
        "docente. Las consignas podrán ser más restrictivas, nunca más permisivas."
    ),
    "22.2. En investigación y producción de conocimiento": (
        "Búsquedas, estilo de textos propios o análisis exploratorios pueden auxiliarse; la "
        "IA puede ser, además, objeto de estudio. No fabrica datos ni reemplaza autoría ni "
        "validación metodológica. Lo relevante debe poder declararse y defenderse."
    ),
    "22.3. En gestión y soporte institucional": (
        "Tareas repetitivas de bajo impacto, organización de información no sensible y "
        "comunicaciones no decisionales con revisión humana pueden aliviarse. Becas, "
        "sanciones, desempeño o exclusiones no se automatizan ni se perfilan de modo lesivo. "
        "Siempre hay un responsable de área identificable."
    ),
    "23.1. Supuestos de condicionamiento": (
        "Asistencia sustancial en trabajos evaluables, sistemas de impacto medio y tratamientos "
        "de datos que exigen base jurídica clara no están prohibidos de antemano, pero tampoco "
        "son libres: dependen de habilitación, declaración, evaluación o controles reforzados."
    ),
    "24.1. Prohibiciones en clave de integridad académica y científica": (
        "Presentar como propio un trabajo sustancialmente generado por IA sin elaboración ni "
        "declaración, falsificar datos, citas o evidencias, o eludir evaluaciones no son "
        "atajos menores: simulan lo no aprendido e inventan lo no investigado."
    ),
    "24.2. Prohibiciones en clave de derechos de las personas": (
        "Vigilar, perfilar o discriminar de modo ilegítimo, tratar datos sensibles en "
        "herramientas no autorizadas o adoptar decisiones de alto impacto sin intervención "
        "humana lesiona dignidad y privacidad. El catálogo ilustra un umbral, no lo agota."
    ),
    "24.3. Prohibiciones en clave institucional y de seguridad": (
        "Desplegar sistemas de riesgo medio o alto sin el ciclo de autorización, desactivar "
        "controles o extraer información masiva hacia plataformas no habilitadas compromete "
        "la soberanía de la información y la trazabilidad de las decisiones institucionales."
    ),
    "Artículo 25°. Limitaciones transversales y proporcionalidad": (
        "Un uso etiquetado como permitido puede desbordarse por intensidad, escala u opacidad. "
        "Rigen, además, proporcionalidad entre finalidad y riesgo, minimización de datos, "
        "transparencia compatible con la seguridad y preservación de la equidad."
    ),
    "Artículo 26°. Autorización excepcional": (
        "Solo supuestos fundados, con evaluación reforzada y sin contravenir las prohibiciones "
        "del Capítulo IV, pueden habilitar un uso excepcional. No crea precedente automático "
        "ni convalida lo expresamente vedado, ni sirve de atajo informal al ciclo ordinario."
    ),
    "Artículo 27°. Deber de adecuación de prácticas existentes": (
        "Lo ya en marcha al entrar en vigencia no queda al margen. Las unidades identifican, "
        "declaran y adecuan usos y sistemas, y suspenden de inmediato lo manifiestamente "
        "prohibido. No hay doble vara entre lo nuevo que se autoriza y lo viejo que se tolera."
    ),
    "28.1. Niveles preuniversitarios": (
        "En colegios y escuelas dependientes el criterio es más restrictivo que en el grado. "
        "El uso en aprendizajes exige acompañamiento de la persona adulta responsable; no se "
        "cargan datos de menores en herramientas no autorizadas; en la duda, prima la restricción."
    ),
    "29.2. Deberes del docente": (
        "Verificar exactitud y pertinencia, informar las reglas de la asignatura, no cargar "
        "datos de estudiantes en plataformas no autorizadas y no delegar la acreditación "
        "pertenecen al juicio pedagógico. La IA en el aula es asistencia revisada, no sustituto."
    ),
    "29.5. Usos estudiantiles generalmente admisibles": (
        "Pedir explicaciones, organizar ideas, practicar en actividades no evaluativas y "
        "mejorar la claridad de un texto propio puede ser lícito, salvo regla más estricta "
        "de la cátedra. El resultado debe evidenciar aprendizaje real, no un producto ajeno."
    ),
    "29.6. Usos estudiantiles condicionados": (
        "Borradores de trabajos evaluables, asistencia sustancial en TFI o tesis, código o "
        "análisis que sean objeto de evaluación, y usos colaborativos que nuben la atribución "
        "requieren habilitación explícita y declaración. Sin ella, el uso se tiene por no autorizado."
    ),
    "29.7. Usos estudiantiles no permitidos": (
        "Entregar como propio un trabajo sustancialmente generado por IA sin elaboración ni "
        "declaración, usarla en exámenes vedados, fabricar citas o datos, suplantar identidad "
        "o compartir datos de terceros y de pacientes afecta la equidad y, a veces, la confidencialidad."
    ),
    "30.3. TFI, tesis y trabajos finales": (
        "La autoría intelectual y la responsabilidad científica pertenecen al estudiante, con "
        "su dirección. Si se admite IA, se declara de modo específico; no reemplaza marco "
        "teórico, análisis crítico ni conclusiones. El título acredita competencias humanas."
    ),
    "30.5. Matriz orientativa de niveles de uso según riesgo académico": (
        "Uso libre e incentivado en tareas de bajo riesgo, uso guiado y declarado en riesgos "
        "medios, uso restringido o prohibido en evaluaciones sincrónicas o de desempeño "
        "auténtico: la matriz traduce operativamente las prohibiciones, no las reemplaza."
    ),
    "30.7. Declaración de uso": (
        "Herramienta, periodo, finalidad, alcance de la asistencia y verificación humana no "
        "son un trámite vacío. Pueden incluirse prompts principales y validaciones, para que "
        "quien evalúa distinga aprendizaje genuino de mera mediación instrumental y el "
        "Observatorio pueda proponer un modelo institucional homogéneo."
    ),
    "31.2. Orientaciones para la evaluación": (
        "Defensas orales, problemas situados, portafolios, crítica de salidas generativas e "
        "instancias de desempeño auténtico miden procesos, no solo productos. El rediseño es "
        "condición de validez y equidad, no un capricho pedagógico frente a la disponibilidad de IA."
    ),
    "31.3. Lo que no se delega": (
        "Aprobación o desaprobación, ponderación última del mérito, apreciación de la formación "
        "integral cuando integre la evaluación, y determinación de sanciones académicas "
        "permanecen en el oficio docente. Calificar es un acto de juicio responsable, no la "
        "confirmación de un puntaje opaco producido por un sistema."
    ),
    "34.1. Usos admisibles": (
        "Apoyo bibliográfico verificado, asistencia de estilo sobre texto propio, exploración "
        "de hipótesis bajo validación científica y desarrollo de modelos o código con "
        "trazabilidad se admiten con control metodológico y declaración cuando corresponda. "
        "Investigar la IA misma también es lícito; cada resultado asistido debe poder explicarse."
    ),
    "34.2. Deberes del investigador responsable": (
        "Quien dirige o ejecuta define finalidad y alcance, verifica calidad y límites, "
        "documenta el uso relevante y asegura interpretación humana final. No se fabrican "
        "datos ni se ocultan limitaciones metodológicas derivadas de la herramienta."
    ),
    "34.3. Usos condicionados o de alto escrutinio": (
        "Datos personales o de salud, sistemas predictivos sobre personas, impacto clínico o "
        "social, entrenamiento con datos institucionales y estudios en comunidades vulnerables "
        "exigen evaluación reforzada y, en su caso, comité de ética, Observatorio o autoridad "
        "competente. La celeridad académica no justifica omitir salvaguardas ni tratar el "
        "consentimiento como cláusula de estilo."
    ),
    "34.5. Autoría": (
        "Solo personas humanas con aporte intelectual significativo firman. Los sistemas no "
        "se consignan como coautores; el uso relevante se reconoce según la revista o "
        "repositorio. La dirección de tesis no se delega. La firma imputa mérito y falta."
    ),
    "35.3. Datos institucionales y de terceros": (
        "Los datos de la Universidad no se cargan en herramientas externas sin autorización "
        "y evaluación de riesgo. Los obtenidos de empresas, escuelas, centros de salud u "
        "organismos se usan solo dentro del convenio. Reutilizarlos para entrenar exige base jurídica."
    ),
    "35.5. Comités de ética y evaluación de riesgo": (
        "Si un proyecto puede afectar derechos, salud, privacidad, equidad o seguridad, no "
        "basta el criterio del equipo. Corresponde el comité competente, consulta al "
        "Observatorio, revisión técnica si hay despliegue y evaluación de riesgo del Capítulo III."
    ),
    "36.2. Usos admisibles": (
        "Alfabetizar comunidades, asistir a instituciones públicas o sociales en una adopción "
        "responsable, impulsar innovación regional y abrir deliberación ética se fomentan "
        "porque sirven al medio, no porque probar tecnología sea un fin en sí mismo."
    ),
    "36.3. Limitaciones": (
        "Introducir sistemas de alto riesgo sin evaluación, generar dependencia opaca, usar "
        "datos comunitarios sin consentimiento ni resguardo, o presentar la IA como solución "
        "automática a problemas sociales o sanitarios complejos no es servicio: transfiere a "
        "terceros riesgos que la Universidad no aceptaría para su propia comunidad."
    ),
    "Artículo 40°. Finalidades admisibles y sistemas de gestión": (
        "Eficiencia e información son admisibles si no sustituyen decisiones de alto impacto "
        "sobre personas. Analítica, alertas y automatizaciones de soporte se leen junto con "
        "la exigencia de responsable humano y el ciclo de autorización cuando el riesgo lo pida."
    ),
    "40.3. Becas, priorizaciones y asignaciones": (
        "Un ranking opaco que excluya o estigmatice no es admisible. Si la IA interviene en "
        "beneficios, solo asiste con criterios transparentes, datos pertinentes y revisión "
        "humana efectiva: puede ordenar información, no decidir quién accede."
    ),
    "41.2. Prohibición de automatización plena": (
        "Nadie puede ser objeto de una decisión plenamente automatizada de alto impacto en "
        "gestión. Debe haber intervención humana significativa y posibilidad de no ser "
        "clasificado de modo opaco. «Lo decidió el sistema» no es fundamento ni eximente."
    ),
    "42.1. Condiciones generales": (
        "Finalidad, responsable humano, régimen de datos, propiedad intelectual, niveles de "
        "riesgo y salvaguardas de discontinuidad unifican lo que antes quedaba disperso entre "
        "compras, convenios y áreas usuarias. Sin ese piso, no hay contratación alineada."
    ),
    "42.5. Requisitos previos a la contratación": (
        "Necesidad, proporcionalidad, evaluación de riesgo, opinión técnica y, cuando el "
        "impacto lo amerite, Observatorio y autorización del nivel competente se verifican "
        "antes de comprar. La urgencia de un proveedor o de un convenio no es excepción permanente."
    ),
    "42.6. Cláusulas mínimas deseables": (
        "Limitaciones de finalidad, prohibición de usos secundarios no autorizados, "
        "notificación de incidentes, auditoría razonable, portabilidad y condiciones de "
        "salida no pretenden un pliego único: son el estándar mínimo de soberanía institucional."
    ),
    "Artículo 43°. Registro, seguridad y resguardo de la información": (
        "Todo sistema institucional debe poder inscribirse, asegurarse y discontinuarse sin "
        "dejar a la Universidad sin sus datos ni sin rastro de decisiones. Registro y "
        "controles técnicos no son formalidades: son la condición de una gobernanza verificable."
    ),
    "43.2. Controles mínimos": (
        "Gestión de accesos, registro de eventos, copias de resguardo, separación de ambientes "
        "y revisión periódica se escalonan según el riesgo. Un asistente de bajo impacto no "
        "exige el mismo rigor que un sistema que trata legajos o alertas sobre personas."
    ),
    "43.4. Prohibiciones de seguridad": (
        "Compartir credenciales, conectar herramientas no autorizadas a bases institucionales, "
        "extraer información masiva hacia plataformas externas o eludir controles de auditoría "
        "no son atajos de trabajo: exponen datos y rompen la trazabilidad."
    ),
    "44.2. Transferencias y almacenamiento": (
        "Subir a la nube no es un acto neutro. Si datos institucionales o personales se tratan "
        "fuera del entorno controlado, se exige base jurídica y contractual, nivel de protección "
        "del destino y evaluación de los riesgos de acceso por terceros."
    ),
    "Artículo 45°. Roles, prohibiciones e implementación gradual": (
        "Rectorado, áreas dueñas del proceso, TIC y Observatorio se reparten autorización, uso "
        "cotidiano, seguridad y asesoramiento ético. Esa distribución evita el cuello de "
        "botella y el vacío de imputación al pasar de la norma a la práctica."
    ),
    "55.2. Impugnación": (
        "Apoyo indebido en una salida algorítmica, indicios de error o sesgo, falta de "
        "información debida o vulneración de privacidad, equidad o integridad habilitan a "
        "impugnar o pedir reconsideración por las vías académicas, administrativas o disciplinarias ya existentes."
    ),
    "57.1. Derechos en materia de datos": (
        "Finalidad legítima, minimización, seguridad, corrección, oposición a entrenamientos "
        "no autorizados y, cuando proceda, acceso, portabilidad o limitación del tratamiento "
        "se leen con la normativa de protección de datos vigente, no en abstracto."
    ),
    "57.2. Deberes correlativos de la Universidad": (
        "Minimizar la carga en herramientas de IA, preferir agregados o seudonimización, "
        "restringir accesos por rol, evitar plataformas no autorizadas para legajos o salud, "
        "documentar bases de tratamiento cuando el riesgo lo amerite e incorporar privacidad "
        "desde el diseño convierte las garantías en prácticas exigibles, no en enunciados."
    ),
    "57.3. Confidencialidad reforzada": (
        "Salud y discapacidad, legajos disciplinarios, becas, acompañamiento psicológico o "
        "pastoral, información de menores, secretos profesionales y, cuando intervengan, "
        "datos genéticos, biométricos o bancarios no son «un dato más» para alimentar un "
        "modelo. Exigen un umbral más alto de resguardo, independiente del criterio aislado de cada área."
    ),
    "58.2. Deberes de prevención": (
        "En riesgo medio o alto se evalúan sesgos previsibles antes del despliegue, se "
        "monitorean impactos diferenciales y se corrigen diseños excluyentes. Las alertas "
        "deben derivar a acompañamiento, no a prejuicio ni a estigma de quien aparece «alertado»."
    ),
    "59.1. Hechos reportables": (
        "Usos prohibidos o de alto riesgo no autorizados, sesgo o exclusión algorítmica, "
        "filtraciones, decisiones de impacto sin supervisión humana, represalias contra "
        "alertantes e incumplimientos de proveedores merecen canal institucional. El listado no es cerrado."
    ),
    "59.3. Vía preferente": (
        "El aula se resuelve en lo académico-pedagógico; la investigación, en ética; las "
        "becas y trámites, en lo administrativo; lo grave, en lo disciplinario o laboral; "
        "los incidentes informáticos, en seguridad. El Observatorio orienta; no sustituye "
        "a Decanos, comités ni áreas de personal."
    ),
    "Artículo 60°. Protección de denunciantes e incidentes": (
        "Quien reporte de buena fe no sufre represalias; quien resulte reportado conserva "
        "defensa y descargo; se evita la difusión prematura; la confidencialidad no encubre "
        "inacción ante riesgos graves. El canal solo es creíble si ambas garantías se sostienen."
    ),
    "63.4. Formación obligatoria": (
        "Docentes de nuevo ingreso o en renovación de funciones, responsables de sistemas "
        "institucionales y roles de alto contacto con IA pueden tener instancias obligatorias, "
        "a propuesta del Observatorio. No es un castigo: es condición de un uso informado."
    ),
    "Artículo 64°. Formación docente, estudiantil y de gestión": (
        "En docencia priorizan consignas transparentes e integración pedagógica; en "
        "estudiantes, autoría e integridad; en gestión, decisiones humano-algorítmicas, no "
        "discriminación y seguridad de la información. Un curso único para todos suele ser inútil."
    ),
    "Artículo 65°. Buenas prácticas, guías y comunidades de práctica": (
        "Un repositorio muerto no forma cultura. El Observatorio mantiene un banco de "
        "experiencias verificadas en docencia, investigación, gestión y extensión, y promueve "
        "comunidades de práctica; las guías se actualizan con esos casos, sin sustituir el "
        "articulado cuando haya conflicto."
    ),
    "66.2. Acciones mínimas de difusión": (
        "Publicación en sitios institucionales, comunicación a autoridades de carrera y "
        "piezas periódicas son un piso de visibilidad, complementario del Capítulo XII. "
        "De poco sirve un Marco que no llega a Decanatos, secretarías y aulas."
    ),
    "66.3. Lenguaje accesible": (
        "Claridad, ejemplos y distinción nítida entre lo permitido, lo condicionado y lo "
        "prohibido, sin perder rigor. Observatorio, Rectorado y unidades se reparten diseño, "
        "aprobación y aplicación, para que el mensaje no se vuelva tecnicismo incomprensible."
    ),
    "Artículo 68°. Responsables del seguimiento": (
        "El Observatorio coordina e informa; el Rectorado recibe y dispone; las unidades "
        "proveen datos y ejecutan adecuaciones; TIC aporta evidencia de seguridad e incidentes. "
        "El seguimiento no es privativo de un organismo ni un reporte cosmético sin dueños."
    ),
    "69.2. Indicadores orientativos de cumplimiento": (
        "Porcentaje de unidades con política comunicada, sistemas registrados por riesgo, "
        "cobertura formativa y canales de consulta ilustran evidencia de aplicación. No son "
        "metas rígidas ni ranking punitivo; se adaptan a la escala de cada sede y función."
    ),
    "69.3. Indicadores orientativos de impacto": (
        "Conflictos de integridad, percepción de claridad de las reglas, reducción de "
        "incidentes de datos y calidad de las autorizaciones de alto riesgo sirven para "
        "aprender, no para un tablero decorativo ni para un juicio automático sobre personas."
    ),
    "Artículo 70°. Evidencia e informes institucionales": (
        "Mapa, registro de sistemas, consultas, incidentes, auditorías e informes del "
        "Observatorio al Rectorado y, cuando corresponda, al Consejo Superior nutren el "
        "seguimiento. Sin esa evidencia, la mejora continua queda en una fórmula vacía y "
        "los informes dependen del recuerdo informal de unos pocos gestores."
    ),
    "71.1. Auditoría ética": (
        "Dignidad, equidad, supervisión humana efectiva, proporcionalidad y coherencia con "
        "el ideario pueden verificarse sobre sistemas, procesos o unidades. No es cacería de "
        "infracciones menores ni duplica, aunque pueda complementar, la revisión técnica de TIC."
    ),
    "Artículo 72°. Adecuación, actualización del Marco y anexos": (
        "Hallazgos, medidas, plazos, indicadores de cierre y apoyo requerido componen el plan "
        "cuando el seguimiento detecte brechas. Actualizar el articulado y actualizar anexos "
        "siguen vías distintas: el texto no debe envejecer inmóvil ni alterarse su núcleo por circular."
    ),
    "72.2. Reformas sustanciales y operativas": (
        "Principios, prohibiciones centrales, derechos o arquitectura de gobernanza se reforman "
        "por el órgano que aprobó el Marco. Anexos, matrices, formularios y guías pueden "
        "actualizarse por Rectorado, con dictamen del Observatorio, si no contradicen el articulado."
    ),
    "Artículo 73°. Pilotos, transparencia y proporcionalidad del control": (
        "Alcance limitado, duración definida, responsable humano, evaluación de riesgo y "
        "posibilidad de corte: el piloto es experimentación controlada, no una puerta trasera "
        "al Capítulo III. No se eterniza ni se expande sin nueva autorización, ni elude la "
        "transparencia y la proporcionalidad del control."
    ),
    "Artículo 75°. Remisión a normativas vigentes": (
        "En lo no regulado aquí siguen vigentes el Estatuto, el plan estratégico, las "
        "ordenanzas de docencia e investigación, la normativa de datos y el régimen "
        "disciplinario y laboral. El Marco no deroga ese ordenamiento: la IA se gobierna "
        "en este texto; las sanciones, en las normas ya existentes."
    ),
    "76.2. Criterios de tipificación orientativa": (
        "Falsificación de datos o citas, autoría fraudulenta, elusión de evaluaciones, carga "
        "de datos sensibles en herramientas no autorizadas y despliegue clandestino de sistemas "
        "de alto riesgo se ponderan como especialmente graves. La guía no es un código penal interno."
    ),
    "Artículo 77°. Interpretación y lagunas": (
        "Dignidad, integridad, transparencia, equidad y responsabilidad humana privilegian "
        "la lectura armónica. Ante silencio rige el criterio pro persona y de prevención del "
        "riesgo. Una práctica social extendida no convalida lo no autorizado."
    ),
    "Artículo 78°. Anexos y desarrollo reglamentario": (
        "Glosario, matriz, declaración, evaluación de riesgo, guías, política de consignas, "
        "incidentes, registro y guía preuniversitaria operan el Marco una vez aprobados. Si "
        "chocan con el articulado, prevalece el cuerpo normativo, no el documento suelto."
    ),
    "79.1. Período de adecuación": (
        "Ciento ochenta días desde la vigencia sirven para identificar usos y sistemas, "
        "declararlos, adecuar prácticas y suspender lo manifiestamente prohibido. No es "
        "moratoria de las prohibiciones graves ni trampa de incumplimiento instantáneo para "
        "quien actuaba de buena fe al momento de la entrada en vigor."
    ),
    "Artículo 81°. Difusión obligatoria": (
        "Rectorado y vicerrectorados, sedes, facultades, escuelas, institutos, secretarías, "
        "docentes, estudiantes y personal de gestión deben recibir el texto por canales "
        "idóneos. La difusión es condición de exigibilidad razonable, no un gesto de cortesía."
    ),
}

CLOSE = {
    "3.2. Funciones sustantivas y de gestión":
        "Invocar que el uso ocurrió «solo en gestión» o «solo en un piloto de cátedra» no exime de las salvaguardas.",
    "3.3. Sujetos alcanzados":
        "El vacío de imputación se cierra: las mismas reglas rigen, en lo pertinente, para todos los destinatarios.",
    "3.4. Tecnologías comprendidas":
        "Una plataforma ya contratada no escapa a la gobernanza por no anunciarse como inteligencia artificial generativa.",
    "4.1. Normativa e instrumentos institucionales":
        "Ninguna disposición de IA se lee desconectada de los fines universitarios ni de las competencias de gobierno ya asignadas.",
    "4.2. Referentes externos de orientación":
        "Sirven para interpretar, no para sustituir, el criterio propio de la Universidad Católica de Cuyo.",
    "4.4. Ámbito internacional":
        "No hay incorporación automática del derecho europeo ni de otros ordenamientos como fuente de obligaciones extrañas al Estatuto.",
    "4.5. Ámbito nacional":
        "Esas leyes no se transcriben aquí: marcan el horizonte con el cual este texto debe interpretarse de modo armónico.",
    "4.6. Ámbito eclesial y de universidades católicas":
        "Impiden una lectura meramente tecnocrática: la persona y el bien común no son un aditamento retórico del preámbulo.",
    "Artículo 5°. Principios de interpretación":
        "Cede la conveniencia administrativa, no el principio, cuando un uso eficiente ponga en riesgo la equidad o la integridad académica.",
    "Artículo 9°. Principio rector de no sustitución del juicio humano":
        "Ningún sistema se invoca como sustituto del discernimiento responsable en los ámbitos que a continuación se precisan.",
    "Artículo 10°. Finalidad formativa de la gobernanza de la IA":
        "La gobernanza se entiende como cultura institucional, no como un régimen que solo prohíbe o sanciona.",
    "12.4. Unidades Académicas y autoridades de gestión académica":
        "Informan al Observatorio lo relevante: la gobernanza no queda reservada al Rectorado ni a un área central.",
    "12.5. Áreas de Tecnologías de la Información y sistemas":
        "Impiden integraciones no autorizadas, documentan configuraciones y contribuyen a la puesta en marcha o al retiro de herramientas.",
    "12.6. Secretarías y áreas de gestión":
        "Documentan usos relevantes para el mapeo institucional, de modo que la gobernanza no se concentre solo en lo académico.",
    "12.7. Comunidad universitaria":
        "La gobernanza se sostiene en esas prácticas cotidianas, no solo en órganos de gobierno.",
    "14.1. Decisiones o tareas no delegables":
        "Si el acto determina una trayectoria, un derecho o una sanción, la imputación permanece en una persona identificable.",
    "14.2. Delegación condicionada":
        "Los supuestos típicos no habilitan, por sí solos, a automatizar la decisión final ni a prescindir del responsable humano.",
    "14.3. Tareas delegables":
        "Siguen vigentes la seguridad de la información y la coherencia con el Marco: delegable no significa irresponsable.",
    "15.2. Contenidos mínimos":
        "También deben aparecer las prácticas «de hecho» que generan riesgo académico, ético o legal, aunque nadie las haya formalizado.",
    "15.3. Decisiones críticas de referencia":
        "La lista orienta el mapeo y la evaluación de riesgo; no autoriza a omitir un acto crítico porque no figure en el ejemplo.",
    "15.7. Decisión de autorización":
        "La gradación impide tanto la burocracia sobre usos inocuos como el despliegue clandestino de sistemas que afectan personas.",
    "16.2. Criterios de valoración":
        "La autorización se funda en un juicio documentado, proporcional y comparable entre unidades y sedes.",
    "16.3. Supuestos de evaluación reforzada":
        "En esos casos se refuerzan controles, dictámenes y, si corresponde, la vía de autorización del Rectorado.",
    "17.3. Decisión de autorización":
        "Sin responsable ni posibilidad de suspensión, no hay sistema institucional en producción, por urgente que parezca el área.",
    "Artículo 22°. Finalidades legítimas y usos permitidos":
        "La mera reducción de esfuerzo cuando simula competencias, opacidad o elusión de responsabilidades no es finalidad legítima.",
    "22.1. En docencia y aprendizaje":
        "El criterio es pedagógico: enriquecer comprensión y diversidad de estrategias, no simular competencias no adquiridas.",
    "22.2. En investigación y producción de conocimiento":
        "Toda incidencia relevante deberá poder declararse y defenderse ante pares, comités o financiadores.",
    "22.3. En gestión y soporte institucional":
        "Cabe analítica descriptiva agregada y asistencia operativa, nunca perfilamiento lesivo de estudiantes o trabajadores.",
    "23.1. Supuestos de condicionamiento":
        "La zona gris no se interpreta como autorización implícita: sin habilitación, el uso condicionado no está permitido.",
    "24.1. Prohibiciones en clave de integridad académica y científica":
        "Su gravedad se ponderará en cada procedimiento, pero el criterio de fondo no admite simulación de título ni de publicación.",
    "24.2. Prohibiciones en clave de derechos de las personas":
        "A partir de ese umbral un uso deja de ser innovación y se convierte en lesión de derechos de quien estudia, enseña o trabaja.",
    "24.3. Prohibiciones en clave institucional y de seguridad":
        "Aun sin dolo académico, esas conductas comprometen el control institucional sobre la información y sobre las decisiones.",
    "Artículo 25°. Limitaciones transversales y proporcionalidad":
        "La excepción no se vuelve regla ni la eficiencia justifica un control desmesurado de las personas.",
    "Artículo 26°. Autorización excepcional":
        "Condiciones, autoridad competente y documentación se exigen siempre: la excepción informal no existe.",
    "Artículo 27°. Deber de adecuación de prácticas existentes":
        "La adecuación es condición de continuidad, no un trámite optativo de cátedras o áreas de gestión.",
    "28.1. Niveles preuniversitarios":
        "La IA no reemplaza el cuidado psicopedagógico, pastoral ni de bienestar de niños, niñas y adolescentes. El Anexo J precisa el circuito con representantes legales.",
    "29.2. Deberes del docente":
        "Tampoco se vuelcan trabajos inéditos o materiales institucionales a herramientas comerciales inadecuadas.",
    "29.5. Usos estudiantiles generalmente admisibles":
        "La admisibilidad general no borra la autoría ni habilita a entregar un producto sustancialmente generado por IA.",
    "29.6. Usos estudiantiles condicionados":
        "El objetivo es reducir la zona en la que el estudiante alega que «no sabía» si podía servirse de la herramienta.",
    "29.7. Usos estudiantiles no permitidos":
        "La consecuencia académica o disciplinaria no puede sorprender a quien incurra en esas conductas: están nombradas con claridad.",
    "30.3. TFI, tesis y trabajos finales":
        "Directores y tribunales podrán exigir defensas orales o evidencias de proceso. Fabricar datos o fuentes es falta grave.",
    "30.5. Matriz orientativa de niveles de uso según riesgo académico":
        "Cátedras y carreras explicitan el grado de asistencia admitido, en lugar de dejarlo a interpretaciones implícitas.",
    "30.7. Declaración de uso":
        "Sin ese contenido mínimo, quien evalúa no puede distinguir aprendizaje genuino de mediación encubierta.",
    "31.2. Orientaciones para la evaluación":
        "La evaluación debe seguir midiendo lo que dice medir, no la habilidad de ocultar el uso de una herramienta.",
    "31.3. Lo que no se delega":
        "Aunque la IA apoye la retroalimentación formativa, el acto de calificar permanece humano e imputable.",
    "34.1. Usos admisibles":
        "La aceleración no se confunde con rigor: cada resultado asistido debe poder reproducirse en lo pertinente y defenderse.",
    "34.2. Deberes del investigador responsable":
        "Antes de desplegar sistemas con impacto sobre personas se evalúan riesgos éticos y sociales del proyecto.",
    "34.3. Usos condicionados o de alto escrutinio":
        "Ningún calendario de convocatoria justifica omitir esa vía ni tratarla como obstáculo meramente administrativo.",
    "34.5. Autoría":
        "La firma científica sigue imputando responsabilidad a personas identificables, no a un modelo.",
    "35.3. Datos institucionales y de terceros":
        "Se protege el patrimonio informacional de la institución y los derechos de quienes confiaron esos datos.",
    "35.5. Comités de ética y evaluación de riesgo":
        "Ningún calendario de publicación justifica omitir esa salvaguarda.",
    "36.2. Usos admisibles":
        "Las limitaciones del numeral siguiente impiden convertir a la comunidad en laboratorio de sistemas opacos o de alto riesgo.",
    "36.3. Limitaciones":
        "El servicio al medio no transfiere a terceros lo que la Universidad no aceptaría para su propia comunidad.",
    "Artículo 40°. Finalidades admisibles y sistemas de gestión":
        "La modernización no se confunde con delegación indebida de actos que afectan personas.",
    "40.3. Becas, priorizaciones y asignaciones":
        "Frente a la persona afectada permanece el deber de explicación: la herramienta no elude esa respuesta.",
    "41.2. Prohibición de automatización plena":
        "En el ámbito administrativo se reitera el principio rector: el acto institucional es humano.",
    "42.1. Condiciones generales":
        "Sin ese piso contractual, compras, convenios y áreas usuarias vuelven a fragmentar la regla.",
    "42.5. Requisitos previos a la contratación":
        "Comprar primero y gobernar después es exactamente lo que queda vedado.",
    "42.6. Cláusulas mínimas deseables":
        "Sin ellas, la Universidad queda expuesta a dependencias opacas y a perder control sobre sus propios datos.",
    "Artículo 43°. Registro, seguridad y resguardo de la información":
        "Qué se registra, qué controles rigen y qué prácticas están prohibidas se precisan en coordinación con TIC.",
    "43.2. Controles mínimos":
        "«Seguridad» no queda en un enunciado genérico ni se reduce a la configuración por defecto de un proveedor.",
    "43.4. Prohibiciones de seguridad":
        "Personal de gestión y terceros deben saber que la infracción de seguridad no es un detalle técnico menor.",
    "44.2. Transferencias y almacenamiento":
        "Localización, ley aplicable y garantías del encargado forman parte de la soberanía de datos que se procura preservar.",
    "Artículo 45°. Roles, prohibiciones e implementación gradual":
        "Los hitos de implementación evitan improvisar sistemas de alto impacto sin dueño ni controles.",
    "55.2. Impugnación":
        "El derecho a revisión humana no queda en un enunciado sin procedimiento ni causal identificable.",
    "57.1. Derechos en materia de datos":
        "La comunidad debe poder saber qué exigir; la Universidad, qué garantizar.",
    "57.2. Deberes correlativos de la Universidad":
        "Sin esos deberes organizativos, las garantías del numeral anterior serían inexigibles en la práctica cotidiana.",
    "57.3. Confidencialidad reforzada":
        "La excepción de confidencialidad reforzada no depende del criterio aislado de quien carga el dato en una herramienta.",
    "58.2. Deberes de prevención":
        "La no discriminación se convierte en prácticas verificables, no en una declaración de intenciones sin evidencia de control.",
    "59.1. Hechos reportables":
        "El reporte deja de percibirse como delación informal y pasa a ser mecanismo de protección y de aprendizaje institucional.",
    "59.3. Vía preferente":
        "Toda controversia sobre IA no se desvía al Observatorio como si este reemplazara los procedimientos ya vigentes.",
    "Artículo 60°. Protección de denunciantes e incidentes":
        "El registro agregado de casos sirve a la mejora institucional sin exponer personas de más.",
    "63.4. Formación obligatoria":
        "La alfabetización no queda librada al voluntarismo de quien ya está familiarizado con las herramientas.",
    "Artículo 64°. Formación docente, estudiantil y de gestión":
        "Cada público recibe lo que su rol exige; la oferta no se diluye en un curso ocasional sin responsables ni seguimiento.",
    "Artículo 65°. Buenas prácticas, guías y comunidades de práctica":
        "Las «buenas prácticas» no se confunden con anécdotas informales ni desplazan el articulado.",
    "66.2. Acciones mínimas de difusión":
        "Nadie puede alegar con razón que las reglas no se conocían por haber quedado en un archivo interno.",
    "66.3. Lenguaje accesible":
        "El rigor no se diluye, pero tampoco se reserva a un tecnicismo incomprensible para estudiantes y personal de apoyo.",
    "Artículo 68°. Responsables del seguimiento":
        "La mejora continua tiene dueños identificables y plazos, no meras intenciones.",
    "69.2. Indicadores orientativos de cumplimiento":
        "Muestran si el Marco se está aplicando; no sustituyen el juicio sobre personas ni unidades.",
    "69.3. Indicadores orientativos de impacto":
        "La lectura es cualitativa y cuantitativa, sin convertir cada cifra en un veredicto automático.",
    "Artículo 70°. Evidencia e informes institucionales":
        "Los informes no dependen del recuerdo informal de unos pocos gestores.",
    "71.1. Auditoría ética":
        "Es un control de sentido institucional, distinguible de la revisión técnica, aunque ambas puedan complementarse en riesgo medio o alto.",
    "Artículo 72°. Adecuación, actualización del Marco y anexos":
        "El área responsable elabora el plan; no se espera a una reforma general para cerrar brechas operativas.",
    "72.2. Reformas sustanciales y operativas":
        "Se evita tanto la rigidización excesiva como la reforma encubierta por vía de anexo.",
    "Artículo 73°. Pilotos, transparencia y proporcionalidad del control":
        "Si el experimento se vuelve práctica permanente, reingresa al ciclo ordinario de autorización.",
    "Artículo 75°. Remisión a normativas vigentes":
        "Contratos, derechos y procedimientos disciplinarios siguen anclados en las normas que ya existen.",
    "76.2. Criterios de tipificación orientativa":
        "Orientan la ponderación de gravedad para que la respuesta sea proporcional y previsible.",
    "Artículo 77°. Interpretación y lagunas":
        "El Observatorio puede orientar; deciden los órganos competentes. La laguna no se llena con analogías laxas.",
    "Artículo 78°. Anexos y desarrollo reglamentario":
        "La implementación no depende de modelos improvisados ajenos a los instrumentos aquí identificados.",
    "79.1. Período de adecuación":
        "La entrada en vigor no es simbólica: las tareas del transitorio están ordenadas y son exigibles.",
    "Artículo 81°. Difusión obligatoria":
        "Ninguna sede o facultad puede considerarse al margen por no haber sido formalmente notificada del texto.",
}

BANNED_STARTS = (
    "Este apartado",
    "Este Capítulo",
    "Este artículo",
    "Este Marco",
    "Las viñetas",
    "Las enumeraciones",
)

CHAPTER_SWAPS = {
    "Este Capítulo fija el objeto, el alcance y las claves de lectura del Marco. Sin esas definiciones comunes, el resto de las normas se fragmentaría en prácticas de cátedra, de sede o de área, y perdería fuerza institucional.": (
        "Objeto, alcance y claves de lectura evitan que el resto de las normas se fragmente "
        "en prácticas de cátedra, de sede o de área. Sin ese piso común, el texto pierde "
        "fuerza institucional y cada unidad interpreta por su cuenta."
    ),
    "Este Capítulo no es un preámbulo decorativo. Orienta la interpretación de todo el Marco cuando haya tensión entre eficiencia tecnológica, integridad académica y cuidado de las personas.": (
        "Más que un preámbulo, aquí se declara el criterio de lectura de todo el Marco: "
        "cuando la eficiencia tecnológica tense la integridad académica o el cuidado de las "
        "personas, prevalece el sentido institucional ya enunciado."
    ),
    "Este Capítulo ofrece reglas comunes para cátedras y estudiantes, y remite a los Anexos E, F, G y J para los modelos operativos. El cuerpo normativo fija el criterio; los anexos muestran cómo aplicarlo en el aula.": (
        "Cátedras y estudiantes reciben aquí reglas comunes; los Anexos E, F, G y J muestran "
        "modelos operativos. El cuerpo normativo fija el criterio; los anexos ilustran cómo "
        "aplicarlo en el aula, sin sustituirlo."
    ),
    "Este Capítulo articula deberes de investigadores, comités de ética y actividades con el medio. La celeridad de publicación o de transferencia no justifica omitir validación humana, evaluación de riesgo ni resguardo de datos.": (
        "Investigadores, comités de ética y actividades con el medio quedan articulados bajo "
        "un mismo criterio: la celeridad de publicación o de transferencia no justifica omitir "
        "validación humana, evaluación de riesgo ni resguardo de datos."
    ),
    "Este Capítulo regula finalidades admisibles, asistencia automatizada, contratación de proveedores y soberanía de datos. Unifica requisitos que antes estaban dispersos, para que ningún convenio o plataforma se incorpore al margen.": (
        "Finalidades admisibles, asistencia automatizada, contratación de proveedores y "
        "soberanía de datos se unifican para que ningún convenio o plataforma se incorpore "
        "al margen de requisitos que antes estaban dispersos."
    ),
    "Este Capítulo no crea un régimen paralelo al disciplinario o laboral vigente: lo precisa para el contexto algorítmico. Su finalidad primaria es formativa y preventiva; la sanción interviene cuando el incumplimiento lo exige.": (
        "No hay un fuero paralelo al disciplinario o laboral vigente: se precisa ese régimen "
        "para el contexto algorítmico. La finalidad primaria es formativa y preventiva; la "
        "sanción interviene cuando el incumplimiento lo exige."
    ),
    "Este Capítulo organiza esas garantías y las vías para hacerlas valer —consulta, reporte, impugnación—, de modo que el Marco no se reduzca a un catálogo de prohibiciones sin tutela efectiva.": (
        "Consulta, reporte e impugnación organizan las vías para hacer valer las garantías, "
        "de modo que el Marco no se reduzca a un catálogo de prohibiciones sin tutela efectiva "
        "para quien resulte afectado por un uso de IA."
    ),
    "Este Capítulo reconoce derechos y garantías a los miembros de la comunidad universitaria frente al uso de Inteligencia Artificial, con el objeto de proteger dignidad, integridad, privacidad, equidad y debido proceso.": (
        "Dignidad, integridad, privacidad, equidad y debido proceso se reconocen como "
        "derechos y garantías de la comunidad universitaria frente al uso de Inteligencia "
        "Artificial, no como meras declaraciones de intención."
    ),
    "Este Capítulo fija objetivos formativos, responsables y criterios de comunicación institucional, para que las reglas no queden libradas al rumor de cátedra ni a la desigualdad de acceso a criterios claros entre sedes y facultades.": (
        "Objetivos formativos, responsables y criterios de comunicación institucional evitan "
        "que las reglas queden libradas al rumor de cátedra o a la desigualdad de acceso a "
        "criterios claros entre sedes y facultades."
    ),
    "Este Capítulo asigna responsables, indicadores flexibles y vías de actualización, con proporcionalidad: más intensidad donde hay mayor riesgo para personas, integridad académica o datos, y menos carga donde el uso es de bajo impacto.": (
        "Responsables, indicadores flexibles y vías de actualización se asignan con "
        "proporcionalidad: más intensidad donde hay mayor riesgo para personas, integridad "
        "académica o datos, y menos carga donde el uso es de bajo impacto."
    ),
    "Este Capítulo regula el uso de IA en:": "Queda comprendido el uso de IA en:",
    "Este Capítulo se aplica a:": "El alcance cubre:",
    "Este artículo unifica los requisitos para convenios, contratación y proveedores de IA.": (
        "Convenios, contratación y proveedores de IA se rigen por un mismo piso de requisitos."
    ),
    "Este Marco no deroga las normas generales de la Universidad, sino que las complementa en lo específico del uso de Inteligencia Artificial.": (
        "Las normas generales de la Universidad siguen vigentes: aquí se las complementa en "
        "lo específico del uso de Inteligencia Artificial."
    ),
    "Este Marco se interpreta en continuidad con el ordenamiento propio de la Universidad, no como un cuerpo ajeno o paralelo. Estatutos, plan estratégico, ordenanzas de docencia e investigación y la resolución de creación del Observatorio conforman el piso normativo sobre el cual se apoyan las reglas de IA. Las viñetas recogen esos instrumentos para que ninguna disposición se lea desconectada de la identidad institucional, de los fines universitarios ni de las competencias ya asignadas a los órganos de gobierno.": None,  # replaced via heading 4.1
}


EXTRA = {
    "3.2. Funciones sustantivas y de gestión":
        "El listado de ámbitos identifica el campo de aplicación, sin pretender un catálogo cerrado de actividades universitarias.",
    "3.3. Sujetos alcanzados":
        "Terceros contratados y visitantes en actividades oficiales quedan comprendidos cuando usen o resulten afectados por sistemas de IA.",
    "3.4. Tecnologías comprendidas":
        "Lo que importa es la función que cumple el sistema, no la etiqueta comercial ni el canal por el que se accede.",
    "4.1. Normativa e instrumentos institucionales":
        "El Observatorio opera sobre ese piso; no crea una legalidad paralela ni desplaza al Consejo Superior o al Rectorado.",
    "4.2. Referentes externos de orientación":
        "La autonomía universitaria permanece intacta: los referentes comparados no alteran la jerarquía de las normas internas.",
    "4.4. Ámbito internacional":
        "UNESCO, directrices de confianza y aportes comparados operan como brújula ética y de buena práctica, no como ley extranjera aplicable.",
    "4.5. Ámbito nacional":
        "Docentes, investigadores y gestores encuentran aquí anclaje de obligaciones ya existentes en el derecho argentino, no un régimen extraño.",
    "4.6. Ámbito eclesial y de universidades católicas":
        "Las universidades católicas se conciben como espacios de discernimiento; esa vocación orienta límites, prudencia y servicio.",
    "Artículo 5°. Principios de interpretación":
        "Operan como claves de lectura de todo el articulado, no como un decálogo ornamental sin efecto práctico.",
    "Artículo 9°. Principio rector de no sustitución del juicio humano":
        "Acompañamiento psicopedagógico, pastoral o de bienestar permanece reservado a personas humanas, aun cuando exista apoyo tecnológico.",
    "Artículo 10°. Finalidad formativa de la gobernanza de la IA":
        "Educar para el uso crítico y creativo de la tecnología es tan parte de la misión como impedir abusos.",
    "12.4. Unidades Académicas y autoridades de gestión académica":
        "Criterios por asignatura o evaluación son legítimos si no diluyen el piso común ni ocultan el uso a estudiantes.",
    "12.5. Áreas de Tecnologías de la Información y sistemas":
        "Sin ese soporte crítico, el Observatorio y el Rectorado decidirían a ciegas sobre riesgos técnicos y de proveedores.",
    "12.6. Secretarías y áreas de gestión":
        "Bienestar, trámites y educación a distancia aplican la misma matriz de delegación que el resto de la gestión académica.",
    "12.7. Comunidad universitaria":
        "Reportar sesgos o incidentes por canales institucionales es un deber de cuidado, no un acto heroico ni una delación.",
    "14.1. Decisiones o tareas no delegables":
        "La evaluación acreditante y las sanciones no se «confirman» con un clic sobre una recomendación opaca.",
    "14.2. Delegación condicionada":
        "Si hay efectos sobre personas, debe existir posibilidad real de revisión, no una supervisión meramente nominal.",
    "14.3. Tareas delegables":
        "El umbral se cruza apenas la tarea deja de ser de soporte y pasa a determinar derechos o calificaciones.",
    "15.2. Contenidos mínimos":
        "El mapa no duplica la Matriz: visibiliza riesgos y responsables que de otro modo permanecerían invisibles.",
    "15.3. Decisiones críticas de referencia":
        "Alertas de abandono, priorizaciones y evaluaciones de desempeño entran en ese núcleo cuando alteran trayectorias.",
    "15.7. Decisión de autorización":
        "El alto riesgo no se autoriza por silencio administrativo ni por reiteración de un uso de hecho.",
    "16.2. Criterios de valoración":
        "Dependencia de un único proveedor y opacidad del modelo pesan tanto como la sensibilidad de los datos tratados.",
    "16.3. Supuestos de evaluación reforzada":
        "Menores, biométricos y monitoreo intensivo no se diluyen en una evaluación genérica de «herramienta de productividad».",
    "17.3. Decisión de autorización":
        "La reiteración operativa respecto del mapa no es redundancia: cierra la puerta al sistema que entra «en prueba eterna».",
    "Artículo 22°. Finalidades legítimas y usos permitidos":
        "Accesibilidad, inclusión y alfabetización crítica se fomentan; la opacidad y la discriminación no.",
    "22.1. En docencia y aprendizaje":
        "Guías, ejemplos y rúbricas pueden apoyarse en IA si el docente revisa y asume la responsabilidad pedagógica.",
    "22.2. En investigación y producción de conocimiento":
        "El rigor científico no se confunde con acelerar la escritura o la búsqueda a costa de la validación humana.",
    "22.3. En gestión y soporte institucional":
        "Las comunicaciones no decisionales también exigen revisión humana cuando puedan generar expectativas o derechos.",
    "23.1. Supuestos de condicionamiento":
        "Sistemas institucionales de impacto medio y datos con base jurídica dudosa entran en este tramo, no en el de uso libre.",
    "24.1. Prohibiciones en clave de integridad académica y científica":
        "El sentido del título universitario y de la publicación científica es el bien jurídico que estas prohibiciones tutelan.",
    "24.2. Prohibiciones en clave de derechos de las personas":
        "Tratar datos sensibles en herramientas no autorizadas es, además de ilicitud, una ruptura de la confianza institucional.",
    "24.3. Prohibiciones en clave institucional y de seguridad":
        "Desactivar controles de supervisión para «ganar agilidad» no es una decisión de gestión: es una infracción.",
    "Artículo 25°. Limitaciones transversales y proporcionalidad":
        "Minimización de datos y equidad rigen también sobre usos que, en abstracto, habrían sido calificados como permitidos.",
    "Artículo 26°. Autorización excepcional":
        "No convalida lo expresamente vedado en el Capítulo IV ni crea un atajo reiterable para la misma unidad.",
    "Artículo 27°. Deber de adecuación de prácticas existentes":
        "Lo manifiestamente prohibido se suspende de inmediato, sin esperar el vencimiento de ningún plazo de inventario.",
    "28.1. Niveles preuniversitarios":
        "En la duda sobre una herramienta o un dato de menores, se opta por no usar hasta que exista habilitación expresa.",
    "29.2. Deberes del docente":
        "Modelar alfabetización crítica forma parte del oficio: no se enseña a eludir consignas ni a ocultar asistencia.",
    "29.5. Usos estudiantiles generalmente admisibles":
        "Practicar en actividades no evaluativas no autoriza a replicar esa asistencia en el examen o en el trabajo final.",
    "29.6. Usos estudiantiles condicionados":
        "Usos colaborativos que nuben quién hizo qué se tratan como no autorizados mientras la cátedra no los habilite.",
    "29.7. Usos estudiantiles no permitidos":
        "Suplantar identidad o compartir datos de pacientes no es «trampa de examen»: es ilicitud con otra escala de gravedad.",
    "30.3. TFI, tesis y trabajos finales":
        "Marco teórico, análisis crítico y conclusiones son aportes humanos; la herramienta no los sustituye ni los firma.",
    "30.5. Matriz orientativa de niveles de uso según riesgo académico":
        "Evaluaciones sincrónicas y de desempeño auténtico se presumen de uso restringido o prohibido, salvo criterio escrito más abierto.",
    "30.7. Declaración de uso":
        "El modelo institucional de declaración busca homogeneidad entre facultades, no un formulario distinto en cada cátedra.",
    "31.2. Orientaciones para la evaluación":
        "Portafolios, problemas situados y crítica de salidas generativas son medios, no un único método obligatorio para todas las asignaturas.",
    "31.3. Lo que no se delega":
        "La formación integral, cuando integra la evaluación, no se reduce a un puntaje emitido por un sistema.",
    "34.1. Usos admisibles":
        "El desarrollo de código o modelos con trazabilidad se admite; la opacidad sobre cómo se obtuvo el resultado, no.",
    "34.2. Deberes del investigador responsable":
        "Documentar el uso relevante es parte de la integridad metodológica, no un trámite burocrático posterior a la publicación.",
    "34.3. Usos condicionados o de alto escrutinio":
        "El consentimiento informado no se degrada a una cláusula de estilo copiada de otro protocolo.",
    "34.5. Autoría":
        "Las normas de la revista o repositorio se respetan, con un piso institucional mínimo de transparencia sobre el uso de IA.",
    "35.3. Datos institucionales y de terceros":
        "Anonimizar o seudonimizar, cuando proceda, no autoriza a cambiar la finalidad del convenio original.",
    "35.5. Comités de ética y evaluación de riesgo":
        "Si hay despliegue sobre personas, la revisión técnica se suma a la ética: ninguna de las dos basta por sí sola.",
    "36.2. Usos admisibles":
        "La deliberación ética con el medio se fomenta; el extractivismo de datos comunitarios, no.",
    "36.3. Limitaciones":
        "Dependencia tecnológica opaca sobre una escuela, un municipio o un centro de salud es incompatible con la extensión universitaria.",
    "Artículo 40°. Finalidades admisibles y sistemas de gestión":
        "Alertas y analítica de gestión exigen responsable humano antes de producir efectos sobre estudiantes o trabajadores.",
    "40.3. Becas, priorizaciones y asignaciones":
        "Criterios transparentes y datos pertinentes son condición de la asistencia; el secreto del puntaje no es eficiencia.",
    "41.2. Prohibición de automatización plena":
        "La persona tiene derecho a no ser clasificada de modo opaco, aun cuando el sistema «haya acertado» en el pasado.",
    "42.1. Condiciones generales":
        "Propiedad intelectual y discontinuidad del servicio se pactan antes, no cuando el proveedor ya es imprescindible.",
    "42.5. Requisitos previos a la contratación":
        "La oferta agresiva de plazos o precios no sustituye evaluación de riesgo ni dictamen cuando el impacto lo amerite.",
    "42.6. Cláusulas mínimas deseables":
        "Portabilidad y condiciones de salida se negocian a la entrada, porque a la salida el desequilibrio de poder es mayor.",
    "Artículo 43°. Registro, seguridad y resguardo de la información":
        "Sin rastro de decisiones y sin posibilidad de baja ordenada, no hay sistema institucional admisible.",
    "43.2. Controles mínimos":
        "Separación de ambientes y revisión periódica se exigen con más rigor cuanto mayor sea el impacto sobre personas.",
    "43.4. Prohibiciones de seguridad":
        "Eludir la auditoría «para no entorpecer el trabajo» es, precisamente, el tipo de atajo que este numeral veda.",
    "44.2. Transferencias y almacenamiento":
        "El acceso potencial de un subencargado en otro país forma parte del riesgo, no un detalle contractual menor.",
    "Artículo 45°. Roles, prohibiciones e implementación gradual":
        "Áreas jurídicas y de cumplimiento, cuando intervengan, revisan bases legales, contratos y protección de datos.",
    "55.2. Impugnación":
        "Falta de información debida o indicios de sesgo bastan para abrir la vía; no se exige prueba plena del algoritmo.",
    "57.1. Derechos en materia de datos":
        "Acceso y portabilidad, cuando procedan, se tramitan por las vías de protección de datos ya vigentes, no por un canal informal de cátedra o de área.",
    "57.2. Deberes correlativos de la Universidad":
        "Legajos y datos de salud no circulan por plataformas de consumo masivo no autorizadas, aunque sean «gratuitas».",
    "57.3. Confidencialidad reforzada":
        "Acompañamiento psicológico o pastoral no se convierte en insumo de un modelo de analítica institucional.",
    "58.2. Deberes de prevención":
        "Corregir un diseño excluyente es parte del deber, no un gesto de buena voluntad posterior al daño.",
    "59.1. Hechos reportables":
        "Incumplimientos de proveedores y represalias contra alertantes entran en el canal, no se diluyen como «conflicto interno».",
    "59.3. Vía preferente":
        "Un incidente informático no se «académiza»; un conflicto de cátedra no se informatiza: cada uno sigue su fuero.",
    "Artículo 60°. Protección de denunciantes e incidentes":
        "Quien resulte reportado conserva el derecho de descargo; la confidencialidad del trámite no se usa para archivar el caso sin actuación.",
    "63.4. Formación obligatoria":
        "Responsables de sistemas institucionales no pueden permanecer al margen de la alfabetización por antigüedad en el cargo.",
    "Artículo 64°. Formación docente, estudiantil y de gestión":
        "Seguridad de la información y no discriminación son contenidos mínimos de la formación de gestión, no optativos.",
    "Artículo 65°. Buenas prácticas, guías y comunidades de práctica":
        "El banco de experiencias cubre docencia, investigación, gestión y extensión, para que el aprendizaje no quede en un silo.",
    "66.2. Acciones mínimas de difusión":
        "La comunicación a autoridades de carrera es tan necesaria como la publicación en el sitio institucional.",
    "66.3. Lenguaje accesible":
        "Distinguir con nitidez permitido, condicionado y prohibido es un deber de comunicación, no un simplismo.",
    "Artículo 68°. Responsables del seguimiento":
        "El Consejo Superior toma conocimiento de informes de relevancia política y decide las reformas sustanciales del Marco.",
    "69.2. Indicadores orientativos de cumplimiento":
        "Canales de consulta existentes y cobertura formativa son evidencias blandas, pero no prescindibles, de que la norma circula.",
    "69.3. Indicadores orientativos de impacto":
        "La calidad de las autorizaciones de alto riesgo dice más que el mero recuento de formularios presentados.",
    "Artículo 70°. Evidencia e informes institucionales":
        "Consultas recurrentes e incidentes son fuente de aprendizaje, no ruido que se archiva sin lectura institucional.",
    "71.1. Auditoría ética":
        "En sistemas de riesgo medio o alto, ética y TIC pueden complementarse; ninguna absorbe el mandato de la otra.",
    "Artículo 72°. Adecuación, actualización del Marco y anexos":
        "Indicadores de cierre y apoyo requerido evitan planes de adecuación que solo enuncian intenciones sin fecha.",
    "72.2. Reformas sustanciales y operativas":
        "Un anexo no puede contradecir principios, prohibiciones centrales, derechos ni la arquitectura de gobernanza.",
    "Artículo 73°. Pilotos, transparencia y proporcionalidad del control":
        "El corte anticipado es una facultad institucional, no un fracaso que deba ocultarse para «salvar» el experimento.",
    "Artículo 75°. Remisión a normativas vigentes":
        "El régimen laboral y disciplinario aplicable no se reescribe aquí: se aplica al incumplimiento en materia de IA.",
    "76.2. Criterios de tipificación orientativa":
        "El despliegue clandestino de alto riesgo y la autoría fraudulenta no se tratan como faltas equivalentes a un olvido formal.",
    "Artículo 77°. Interpretación y lagunas":
        "El criterio pro persona y de prevención del riesgo de los Capítulos II y IX orienta el silencio normativo.",
    "Artículo 78°. Anexos y desarrollo reglamentario":
        "Declaración, evaluación de riesgo y registro no son papelería opcional: operan artículos ya vigentes del cuerpo.",
    "79.1. Período de adecuación":
        "Declarar lo existente es el primer acto de cumplimiento; ocultarlo no se beneficia del plazo de ciento ochenta días.",
    "Artículo 81°. Difusión obligatoria":
        "La notificación formal a cada unidad es presupuesto de exigibilidad; no se presume el conocimiento por el solo hecho de existir el archivo.",
}


def build(key: str) -> str:
    if key not in CLOSE or key not in EXTRA:
        raise KeyError(f"missing CLOSE/EXTRA for {key}")
    parts = [OPEN[key].rstrip("."), CLOSE[key].rstrip(".")]
    text = ". ".join(p.strip() for p in parts if p.strip()) + "."
    if wc(text) < 60:
        extra = EXTRA[key].strip().rstrip(".")
        candidate = text.rstrip(".") + ". " + extra + "."
        while wc(candidate) > 80 and extra.split():
            extra = " ".join(extra.split()[:-1]).rstrip(".")
            candidate = text.rstrip(".") + ". " + extra + "."
        text = " ".join(candidate.split())
    if not (60 <= wc(text) <= 80):
        raise ValueError(f"{key} wc={wc(text)}")
    return text


def main() -> None:
    missing = [k for k in OLD if k not in OPEN or k not in CLOSE or k not in EXTRA]
    extra = [k for k in list(OPEN) + list(CLOSE) if k not in OLD]
    if missing or extra:
        print("KEY MISMATCH missing", missing, "extra", extra)
        raise SystemExit(1)

    built = {k: build(k) for k in OLD}
    bad = [(k, wc(v), v[:80]) for k, v in built.items() if not (60 <= wc(v) <= 80)]
    if bad:
        print("WORD COUNT OUT OF RANGE")
        for k, n, s in bad:
            print(f"  {n:3}  {k}  | {s}")
        raise SystemExit(1)

    starts = [" ".join(v.split()[:5]) for v in built.values()]
    from collections import Counter
    dups = [s for s, n in Counter(starts).items() if n > 1]
    if dups:
        print("DUPLICATE 5-WORD STARTS", dups)
        raise SystemExit(1)
    for k, v in built.items():
        if v.startswith(BANNED_STARTS) or "Este apartado" in v:
            print("BANNED IN", k, v[:80])
            raise SystemExit(1)

    doc = Document(TARGET)
    start = False
    replaced = 0
    missing_h = 0
    for i, p in enumerate(doc.paragraphs):
        t = p.text.strip().split("\t")[0]
        if "VERSIÓN 2026" in t or "VERSIÓN 2.0 DEPURADA" in t:
            start = True
            continue
        if not start or p.style.name.startswith("toc"):
            continue
        if p.style.name not in ("Heading 2", "Heading 3"):
            continue
        if t not in built:
            continue
        # replace first following nonempty Normal paragraph that is not a list item
        for j in range(i + 1, min(i + 5, len(doc.paragraphs))):
            q = doc.paragraphs[j]
            qt = q.text.strip()
            if q.style.name.startswith(("Heading", "toc")):
                break
            if not qt:
                continue
            if qt.startswith(("a)", "b)", "c)", "d)", "e)", "f)")):
                missing_h += 1
                print("LIST DIRECTLY AFTER", t)
                break
            set_text(q, built[t])
            replaced += 1
            break

    chap_n = 0
    regula_n = 0
    prefix_swaps = [
        ("Este Capítulo regula el uso de IA en:", "Queda comprendido el uso de IA en:"),
        ("Este Capítulo se aplica a:", "El alcance cubre:"),
        ("Este artículo unifica los requisitos para convenios, contratación y proveedores de IA.",
         "Convenios, contratación y proveedores de IA se rigen por un mismo piso de requisitos."),
        ("Este Marco no deroga las normas generales de la Universidad, sino que las complementa en lo específico del uso de Inteligencia Artificial.",
         "Las normas generales de la Universidad siguen vigentes: aquí se las complementa en lo específico del uso de Inteligencia Artificial."),
        ("Este Capítulo fija el objeto, el alcance y las claves de lectura del Marco.",
         "Objeto, alcance y claves de lectura evitan que el resto de las normas se fragmente "
         "en prácticas de cátedra, de sede o de área. Sin ese piso común, el texto pierde "
         "fuerza institucional y cada unidad interpreta por su cuenta."),
        ("Este Capítulo no es un preámbulo decorativo.",
         "Más que un preámbulo, aquí se declara el criterio de lectura de todo el Marco: "
         "cuando la eficiencia tecnológica tense la integridad académica o el cuidado de las "
         "personas, prevalece el sentido institucional ya enunciado."),
        ("Este Capítulo ofrece reglas comunes para cátedras y estudiantes,",
         "Cátedras y estudiantes reciben aquí reglas comunes; los Anexos E, F, G y J muestran "
         "modelos operativos. El cuerpo normativo fija el criterio; los anexos ilustran cómo "
         "aplicarlo en el aula, sin sustituirlo."),
        ("Este Capítulo articula deberes de investigadores,",
         "Investigadores, comités de ética y actividades con el medio quedan articulados bajo "
         "un mismo criterio: la celeridad de publicación o de transferencia no justifica omitir "
         "validación humana, evaluación de riesgo ni resguardo de datos."),
        ("Este Capítulo regula finalidades admisibles,",
         "Finalidades admisibles, asistencia automatizada, contratación de proveedores y "
         "soberanía de datos se unifican para que ningún convenio o plataforma se incorpore "
         "al margen de requisitos que antes estaban dispersos."),
        ("Este Capítulo no crea un régimen paralelo",
         "No hay un fuero paralelo al disciplinario o laboral vigente: se precisa ese régimen "
         "para el contexto algorítmico. La finalidad primaria es formativa y preventiva; la "
         "sanción interviene cuando el incumplimiento lo exige."),
        ("Este Capítulo organiza esas garantías",
         "Consulta, reporte e impugnación organizan las vías para hacer valer las garantías, "
         "de modo que el Marco no se reduzca a un catálogo de prohibiciones sin tutela efectiva "
         "para quien resulte afectado por un uso de IA."),
        ("Este Capítulo reconoce derechos y garantías",
         "Dignidad, integridad, privacidad, equidad y debido proceso se reconocen como "
         "derechos y garantías de la comunidad universitaria frente al uso de Inteligencia "
         "Artificial, no como meras declaraciones de intención."),
        ("Este Capítulo fija objetivos formativos,",
         "Objetivos formativos, responsables y criterios de comunicación institucional evitan "
         "que las reglas queden libradas al rumor de cátedra o a la desigualdad de acceso a "
         "criterios claros entre sedes y facultades."),
        ("Este Capítulo asigna responsables,",
         "Responsables, indicadores flexibles y vías de actualización se asignan con "
         "proporcionalidad: más intensidad donde hay mayor riesgo para personas, integridad "
         "académica o datos, y menos carga donde el uso es de bajo impacto."),
    ]
    for p in doc.paragraphs:
        if p.style.name.startswith("toc"):
            continue
        raw = p.text.strip()
        for pref, neu in prefix_swaps:
            if raw.startswith(pref):
                if pref.startswith("Este Capítulo regula el uso"):
                    regula_n += 1
                    neu = (
                        "Queda comprendido el uso de IA en:"
                        if regula_n == 1
                        else "En la gestión universitaria, la regulación cubre:"
                    )
                set_text(p, neu)
                chap_n += 1
                break

    leftover = sum(1 for p in doc.paragraphs if "Este apartado" in p.text)
    doc.save(TARGET)
    print("replaced_by_heading", replaced)
    print("chapter_swaps", chap_n)
    print("list_directly_after", missing_h)
    print("leftover_apartado", leftover)
    print("defined", len(built), "wc", min(map(wc, built.values())), max(map(wc, built.values())))


if __name__ == "__main__":
    main()
