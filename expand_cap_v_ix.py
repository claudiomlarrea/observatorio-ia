#!/usr/bin/env python3
"""Re-expand Capítulos V and IX in Marco 2026 with prose bridges from v1."""
from __future__ import annotations

import re
import shutil
from datetime import date

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph

TARGET = (
    "/Users/claudiolarrea/Library/CloudStorage/OneDrive-Personal/"
    "16 Secretaría de Investigación/60 Observatorio de Inteligencia Artificial/"
    "MARCO DE GOBERNANZA DE LA IA/Documento FINAL/"
    "Marco de Gobernanza IA - UCCuyo 2026.docx"
)

INSERTIONS = [
    (
        "Artículo 29°. Reglas generales para docentes y estudiantes",
        "El presente artículo establece reglas comunes para docentes y estudiantes. Las consignas "
        "de cátedra, los programas de asignatura y las normativas de carrera podrán ser más "
        "restrictivas, pero no más permisivas que este Marco. Su finalidad es orientar la práctica "
        "docente sin sustituir el juicio pedagógico ni la responsabilidad evaluativa de cada cátedra.",
    ),
    (
        "29.1. Usos docentes admisibles",
        "Los docentes conservan la responsabilidad pedagógica sobre el diseño, la conducción y la "
        "evaluación del aprendizaje. La IA puede apoyar tareas preparatorias o de apoyo, siempre "
        "que el contenido final sea revisado críticamente y adaptado al nivel, al perfil de "
        "egreso y a los objetivos formativos de la asignatura.",
    ),
    (
        "29.4. Principio de autoría y aprendizaje real",
        "La formación universitaria exige que el estudiante construya competencias propias. Por "
        "ello, toda producción académica evaluable debe reflejar comprensión, criterio y "
        "apropiación intelectual, aun cuando se haya utilizado IA como herramienta de apoyo.",
    ),
    (
        "Artículo 30°. Producción académica, transparencia e integridad",
        "La producción académica evaluable requiere especial cuidado en materia de autoría, "
        "transparencia e integridad. Este artículo desarrolla criterios diferenciados según el "
        "tipo de actividad, el nivel de riesgo académico y la etapa formativa del estudiante.",
    ),
    (
        "30.6. Deber de transparencia",
        "La transparencia no es un trámite accesorio: permite evaluar con equidad, proteger la "
        "autoría legítima y distinguir el aprendizaje genuino del mero uso instrumental de "
        "herramientas generativas.",
    ),
    (
        "Artículo 31°. Evaluación e información en el aula",
        "La disponibilidad generalizada de la IA obliga a repensar las estrategias de evaluación "
        "sin renunciar a la exigencia académica. Este artículo orienta un rediseño progresivo "
        "hacia evaluaciones válidas, equitativas y significativas en un entorno tecnológicamente "
        "mediado.",
    ),
    (
        "Artículo 32°. Acompañamiento y resolución de dudas",
        "Dado el dinamismo de las herramientas y la diversidad de contextos disciplinares, la "
        "Universidad debe acompañar a docentes y estudiantes con orientación clara, accesible y "
        "actualizada. Este artículo regula el rol del Observatorio y de las áreas académicas en "
        "esa función de apoyo interpretativo.",
    ),
    (
        "Artículo 54°. Derecho a la información y a la explicación",
        "El uso de IA institucional no puede operar como caja negra frente a quienes resultan "
        "afectados por sus efectos. Este artículo garantiza información comprensible y, cuando "
        "corresponda, explicaciones suficientes sobre la intervención algorítmica y la validación humana.",
    ),
    (
        "Artículo 55°. Revisión humana e impugnación",
        "Los derechos de revisión e impugnación concretan el principio rector de no sustitución "
        "del juicio humano. Toda persona afectada por una decisión de impacto debe poder exigir "
        "intervención humana significativa y acceder a vías razonables de reconsideración.",
    ),
    (
        "Artículo 56°. Integridad académica equitativa",
        "La integridad académica no se opone al uso responsable de la IA, pero exige reglas claras, "
        "criterios de evaluación transparentes y procedimientos que no estigmaticen sin fundamento "
        "ni impongan vigilancia desproporcionada.",
    ),
    (
        "Artículo 57°. Privacidad y protección de datos",
        "El tratamiento de datos personales mediante IA debe ajustarse a la normativa vigente, al "
        "principio de minimización y a los estándares institucionales de seguridad. Este artículo "
        "desarrolla derechos y deberes correlativos en coherencia con el Capítulo III.",
    ),
    (
        "Artículo 59°. Consultas, reporte y resolución de conflictos",
        "Los derechos reconocidos en este Capítulo requieren vías accesibles de consulta, reporte "
        "y resolución. Este artículo articula esos mecanismos con los procedimientos académicos, "
        "administrativos y disciplinarios ya existentes en la Universidad.",
    ),
]

DUPLICATE_EXACT = {
    "Minimizar la carga de datos personales en herramientas de IA;",
    "Preferir datos agregados o anonimizados cuando basten para la finalidad;",
    "Restringir accesos según necesidad y rol;",
    "Evitar el uso de plataformas no autorizadas para legajos, historias académicas, datos de salud u otra información sensible;",
    "Documentar bases de tratamiento y responsables cuando el riesgo lo amerite.",
}

REPLACE_LIST_BLOCKS = {
    "Los reportes se regirán por:\nBuena fe;\nConfidencialidad razonable del trámite;\nProtección frente a represalias;\nTrámite oportuno;\nRemisión al área competente según la materia (académica, laboral, de seguridad, ética, jurídica).": (
        "Los reportes se regirán por los siguientes principios:\n"
        "a) Buena fe;\n"
        "b) Confidencialidad razonable del trámite;\n"
        "c) Protección frente a represalias;\n"
        "d) Trámite oportuno;\n"
        "e) Remisión al área competente según la materia (académica, laboral, de seguridad, ética, jurídica)."
    ),
}

BRIDGE_AFTER = (
    "55.2. Impugnación",
    "Las vías de impugnación previstas en este artículo se complementan con los mecanismos de "
    "consulta, reporte y resolución de conflictos del Artículo 59°, que operan como instancias "
    "preferentes de orientación, alerta temprana y tratamiento institucional de controversias.",
)


def delete_paragraph(p: Paragraph):
    p._element.getparent().remove(p._element)


def insert_after(anchor: Paragraph, text: str) -> Paragraph:
    new_el = OxmlElement("w:p")
    anchor._p.addnext(new_el)
    p = Paragraph(new_el, anchor._parent)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.line_spacing = 1.5
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    return p


def split_insert_after(anchor: Paragraph, text: str) -> Paragraph:
    cur = anchor
    for line in text.split("\n"):
        cur = insert_after(cur, line)
    return cur


def body_start_index(doc: Document) -> int:
    for i, p in enumerate(doc.paragraphs):
        if "VERSIÓN 2026" in p.text or "VERSIÓN 2.0 DEPURADA" in p.text:
            return i
    for i, p in enumerate(doc.paragraphs):
        t = p.text.strip().split("\t")[0]
        if t == "CAPÍTULO I. DISPOSICIONES GENERALES" and p.style.name == "Heading 1":
            return i
    return 0


def find_heading(doc: Document, prefix: str, start: int = 0):
    for i, p in enumerate(doc.paragraphs):
        if i < start:
            continue
        if p.style.name.startswith("toc"):
            continue
        t = p.text.strip().split("\t")[0]
        if t == prefix or (t.startswith(prefix) and not t.endswith("\t")):
            if (
                p.style.name.startswith("Heading")
                or t.startswith("Artículo")
                or re.match(r"^\d+\.\d+\.", t)
            ):
                return i, p
    return None, None


def already_has_next(doc: Document, idx: int, snippet: str) -> bool:
    for j in range(idx + 1, min(idx + 4, len(doc.paragraphs))):
        if snippet[:40] in doc.paragraphs[j].text:
            return True
    return False


def apply_insertions(doc: Document) -> int:
    n = 0
    start = body_start_index(doc)
    for anchor_text, prose in INSERTIONS:
        idx, p = find_heading(doc, anchor_text, start)
        if p is None:
            continue
        if already_has_next(doc, idx, prose):
            start = idx + 1
            continue
        if "\n" in prose:
            split_insert_after(p, prose)
        else:
            insert_after(p, prose)
        n += 1
        start = idx + 1
    return n


def remove_duplicates_art57(doc: Document) -> int:
    removed = 0
    in_art57 = False
    seen = set()
    drop_next_deber = False

    for p in list(doc.paragraphs):
        if p.style.name.startswith("toc"):
            continue
        t = p.text.strip()
        if t.startswith("Artículo 57°"):
            in_art57 = True
            continue
        if in_art57 and t.startswith("Artículo 58°"):
            break
        if not in_art57 or not t:
            continue

        key = t
        if t in DUPLICATE_EXACT:
            if key in seen:
                delete_paragraph(p)
                removed += 1
            else:
                seen.add(key)
            continue

        if t == "La Universidad deberá:":
            if "La Universidad deberá:" in seen:
                drop_next_deber = True
                delete_paragraph(p)
                removed += 1
            else:
                seen.add(t)
            continue

        if drop_next_deber and not re.match(r"^[a-g]\)\s", t) and not t.startswith("57."):
            delete_paragraph(p)
            removed += 1
            if t.startswith("57.") or t.startswith("Gozan de protección"):
                drop_next_deber = False
            continue
        if t.startswith("57.") or t.startswith("Gozan de protección"):
            drop_next_deber = False

        if t.startswith("Gozan de protección reforzada"):
            if key in seen:
                # remove header and following duplicate a)-f) until next subsection
                delete_paragraph(p)
                removed += 1
                continue
            seen.add(key)
            continue

        if t.startswith("Queda prohibido el ingreso de datos biométricos"):
            if key in seen:
                delete_paragraph(p)
                removed += 1
            else:
                seen.add(key)
            continue

        if re.match(r"^[a-g]\)\s", t) and key in seen:
            delete_paragraph(p)
            removed += 1
            continue
        if re.match(r"^[a-g]\)\s", t):
            seen.add(key)

    return removed


def fix_list_paragraphs(doc: Document) -> int:
    n = 0
    replacements = {
        "Buena fe;": "a) Buena fe;",
        "Confidencialidad razonable del trámite;": "b) Confidencialidad razonable del trámite;",
        "Protección frente a represalias;": "c) Protección frente a represalias;",
        "Trámite oportuno;": "d) Trámite oportuno;",
        "Remisión al área competente según la materia (académica, laboral, de seguridad, ética, jurídica).": (
            "e) Remisión al área competente según la materia (académica, laboral, de seguridad, ética, jurídica)."
        ),
        "Dignidad y derechos de la persona;": "a) Dignidad y derechos de la persona;",
        "Integridad académica y científica;": "b) Integridad académica y científica;",
        "Proporcionalidad y buena fe;": "c) Proporcionalidad y buena fe;",
        "Existencia o no de reglas claras previas;": "d) Existencia o no de reglas claras previas;",
        "Nivel de riesgo y daño causado o evitado;": "e) Nivel de riesgo y daño causado o evitado;",
        "Necesidad de aprendizaje institucional y prevención futura.": (
            "f) Necesidad de aprendizaje institucional y prevención futura."
        ),
    }
    for p in doc.paragraphs:
        t = p.text.strip()
        if t in replacements:
            new = replacements[t]
            if p.runs:
                p.runs[0].text = new
                for r in p.runs[1:]:
                    r.text = ""
            else:
                p.text = new
            n += 1
    return n


def add_bridge_55_59(doc: Document) -> bool:
    idx, p = find_heading(doc, BRIDGE_AFTER[0])
    if p is None:
        return False
    if already_has_next(doc, idx, BRIDGE_AFTER[1][:40]):
        return False
    # insert after impugnación paragraph block - find 55.3 or end of 55.2 list
    anchor = p
    for j in range(idx + 1, idx + 12):
        t = doc.paragraphs[j].text.strip()
        if t.startswith("55.3") or t.startswith("Nadie podrá"):
            anchor = doc.paragraphs[j - 1]
            break
    insert_after(anchor, BRIDGE_AFTER[1])
    return True


def add_cap_v_closings(doc: Document) -> int:
    n = 0
    start = body_start_index(doc)
    closings = [
        (
            "Los procedimientos, ejemplos y modelos operativos se desarrollan en los Anexos E, F y G.",
            "Los Anexos E, F y G desarrollan procedimientos, ejemplos y modelos operativos. En caso "
            "de discrepancia entre una consigna docente y este Marco, prevalecerá la norma más "
            "protectora de la integridad académica y de la equidad evaluativa.",
        ),
        (
            "En niveles preuniversitarios rige además el Anexo J.",
            "En niveles preuniversitarios rige además el Anexo J, que desarrolla criterios específicos "
            "de acompañamiento adulto, información a representantes legales y restricciones reforzadas "
            "de uso.",
        ),
    ]
    for anchor_text, extra in closings:
        for i, p in enumerate(doc.paragraphs):
            if i < start or p.style.name.startswith("toc"):
                continue
            if p.text.strip().startswith(anchor_text):
                if already_has_next(doc, i, extra[:40]):
                    break
                insert_after(p, extra)
                n += 1
                break
    return n


def add_cap_ix_closing(doc: Document) -> bool:
    idx, p = find_heading(doc, "Artículo 61°. Interpretación pro persona")
    if p is None:
        return False
    text = (
        "La interpretación de este Capítulo debe realizarse de manera sistemática con los "
        "Capítulos III, IV, VII y VIII, de modo que derechos, deberes, responsabilidades y "
        "mecanismos de resguardo se lean como un conjunto coherente y no como disposiciones aisladas."
    )
    if already_has_next(doc, idx, text[:40]):
        return False
    # insert before list a)-e)
    for j in range(idx + 1, idx + 5):
        if doc.paragraphs[j].text.strip().startswith("a)"):
            insert_after(doc.paragraphs[j - 1], text)
            return True
    insert_after(p, text)
    return True


def main():
    backup = TARGET.replace(".docx", f" - backup pre-expansion {date.today().isoformat()}.docx")
    shutil.copy2(TARGET, backup)

    doc = Document(TARGET)
    stats = {
        "insertions": apply_insertions(doc),
        "duplicates_removed": remove_duplicates_art57(doc),
        "list_fixes": fix_list_paragraphs(doc),
        "bridge_55_59": add_bridge_55_59(doc),
        "cap_v_closings": add_cap_v_closings(doc),
        "cap_ix_closing": add_cap_ix_closing(doc),
    }
    doc.save(TARGET)

    print(f"Backup: {backup}")
    print(f"Saved: {TARGET}")
    for k, v in stats.items():
        print(f"  {k}: {v}")


if __name__ == "__main__":
    main()
