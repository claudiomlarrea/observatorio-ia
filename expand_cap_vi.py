#!/usr/bin/env python3
"""Re-expand Capítulo VI (investigación, extensión y transferencia) in Marco 2026."""
from __future__ import annotations

import re
import shutil
from datetime import date

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph

TARGET = (
    "/Users/claudiolarrea/Library/CloudStorage/OneDrive-Personal/"
    "16 Secretaría de Investigación/60 Observatorio de Inteligencia Artificial/"
    "MARCO DE GOBERNANZA DE LA IA/Documento FINAL/"
    "Marco de Gobernanza IA - UCCuyo 2026.docx"
)

INSERTIONS = [
    (
        "Artículo 34°. Investigación, integridad y publicación",
        "La IA puede acelerar búsquedas, análisis y redacción, pero no sustituye el rigor "
        "metodológico ni la autoría intelectual. Este artículo distingue usos admisibles, "
        "deberes del investigador, supuestos de alto escrutinio y reglas de integridad, "
        "autoría y publicación, de modo que la asistencia algorítmica sea trazable y "
        "científicamente justificada.",
    ),
    (
        "34.4. Integridad científica",
        "La integridad científica no se reduce a evitar el plagio: comprende la veracidad "
        "de los datos, la trazabilidad del método y la honestidad en la comunicación de "
        "hallazgos. El uso de IA no atenúa esas exigencias; las vuelve más visibles.",
    ),
    (
        "34.6. Publicación y comunicación científica",
        "La comunicación científica exige poder distinguir, frente a pares, editores o "
        "financiadores, qué hizo la máquina y qué validó el juicio humano.",
    ),
    (
        "Artículo 35°. Datos, consentimiento y comités de ética",
        "La investigación con IA suele implicar datos de personas, instituciones o "
        "comunidades. Este artículo articula principios de tratamiento, consentimiento, "
        "datos sintéticos y la intervención de comités de ética y del Observatorio, de "
        "modo que ningún criterio de celeridad académica justifique omitir salvaguardas.",
    ),
    (
        "35.2. Consentimiento e información",
        "El consentimiento no es una cláusula formal. La información sobre el uso de IA "
        "debe ser comprensible y suficiente para que las personas involucradas puedan "
        "decidir —o ejercer derechos equivalentes— con conocimiento de implicancias y riesgos.",
    ),
    (
        "Artículo 36°. Extensión y vinculación con el medio",
        "La extensión no es un laboratorio de prueba irrestricto. Las actividades que "
        "incorporen IA deben priorizar el servicio a la comunidad, la inclusión y la "
        "educación pública sobre riesgos y oportunidades, en coherencia con la misión "
        "social de la Universidad, sin transferir a terceros riesgos no evaluados.",
    ),
    (
        "Artículo 37°. Propiedad intelectual y transparencia",
        "Los resultados de investigación, el software y los datasets generados en la "
        "Universidad se rigen por la normativa institucional de propiedad intelectual y "
        "por los convenios específicos. Este artículo ordena las restricciones de licencia "
        "y el deber de consignar el uso relevante de IA en informes, rendiciones y "
        "transferencias. Los convenios con terceros se rigen, además, por el Artículo 42°.",
    ),
    (
        "Artículo 38°. Responsabilidades, prohibiciones y prudencia investigativa",
        "Además de las prohibiciones generales del Capítulo IV, la investigación, la "
        "extensión y la transferencia exigen imputación clara de responsabilidades y un "
        "criterio de prudencia cuando un proyecto pueda generar a la vez beneficios "
        "relevantes y daños significativos. La excelencia científica, en la Universidad, "
        "incluye la excelencia ética.",
    ),
]

LIST_FIXES = {
    "Exactitud de datos y citas;": "a) Exactitud de datos y citas;",
    "Coherencia metodológica;": "b) Coherencia metodológica;",
    "Cumplimiento de ética y consentimiento;": "c) Cumplimiento de ética y consentimiento;",
    "Resguardo de información confidencial o propietaria;": "d) Resguardo de información confidencial o propietaria;",
    "Adecuada declaración del uso de IA, cuando corresponda.": (
        "e) Adecuada declaración del uso de IA, cuando corresponda."
    ),
    "Se declare su carácter sintético;": "a) Se declare su carácter sintético;",
    "No se presenten como datos reales;": "b) No se presenten como datos reales;",
    "No permitan reidentificación indebida de personas;": (
        "c) No permitan reidentificación indebida de personas;"
    ),
    "No se utilicen para eludir requisitos éticos o legales aplicables a datos reales.": (
        "d) No se utilicen para eludir requisitos éticos o legales aplicables a datos reales."
    ),
}

TEXT_REPLACES = [
    (
        "pueden incorporan Inteligencia Artificial",
        "pueden incorporar Inteligencia Artificial",
    ),
]


def insert_after(anchor: Paragraph, text: str) -> Paragraph:
    new_el = OxmlElement("w:p")
    anchor._p.addnext(new_el)
    p = Paragraph(new_el, anchor._parent)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.line_spacing = 1.5
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    return p


def body_start_index(doc: Document) -> int:
    for i, p in enumerate(doc.paragraphs):
        if "VERSIÓN 2026" in p.text or "VERSIÓN 2.0 DEPURADA" in p.text:
            return i
    return 0


def already_has_next(doc: Document, idx: int, snippet: str) -> bool:
    for j in range(idx + 1, min(idx + 4, len(doc.paragraphs))):
        if snippet[:40] in doc.paragraphs[j].text:
            return True
    return False


def find_heading(doc: Document, prefix: str, start: int = 0):
    for i, p in enumerate(doc.paragraphs):
        if i < start or p.style.name.startswith("toc"):
            continue
        t = p.text.strip().split("\t")[0]
        if t == prefix or (t.startswith(prefix) and len(t) < 160):
            if (
                p.style.name.startswith("Heading")
                or t.startswith("Artículo")
                or re.match(r"^\d+\.\d+\.", t)
            ):
                return i, p
    return None, None


def apply_insertions(doc: Document) -> int:
    n = 0
    start = body_start_index(doc)
    for anchor_text, prose in INSERTIONS:
        idx, p = find_heading(doc, anchor_text, start)
        if p is None:
            print("MISSING", anchor_text)
            continue
        if already_has_next(doc, idx, prose):
            start = idx + 1
            continue
        insert_after(p, prose)
        n += 1
        start = idx + 1
    return n


def add_ethics_heading(doc: Document) -> int:
    """Give the merged ethics-committee block a subsection title."""
    start = body_start_index(doc)
    in_vi = False
    for i, p in enumerate(doc.paragraphs):
        if i < start or p.style.name.startswith("toc"):
            continue
        t = p.text.strip()
        if t.startswith("CAPÍTULO VI."):
            in_vi = True
        if t.startswith("CAPÍTULO VII."):
            break
        if not in_vi:
            continue
        if t.startswith("Los proyectos de investigación que empleen IA con potencial impacto"):
            prev = doc.paragraphs[i - 1].text.strip() if i else ""
            if prev.startswith("35.5"):
                return 0
            h = insert_after(doc.paragraphs[i - 1], "35.5. Comités de ética y evaluación de riesgo")
            try:
                h.style = "Heading 3"
            except Exception:
                pass
            for r in h.runs:
                r.bold = False
            return 1
    return 0


def expand_art36(doc: Document) -> int:
    start = body_start_index(doc)
    idx, p = find_heading(doc, "36.3. Limitaciones", start)
    if p is None:
        return 0
    extra = (
        "Toda actividad de extensión con IA que trate datos comunitarios, escolares o de "
        "salud, o que implante un sistema con efectos sobre personas, deberá contar con "
        "responsable humano identificable, evaluación de riesgo proporcional y, cuando "
        "corresponda, consentimiento o acuerdo con la organización destinataria. La "
        "Universidad no transferirá a la comunidad herramientas opacas ni pilotos de alto "
        "riesgo sin plan de acompañamiento, reversibilidad y rendición de resultados."
    )
    # insert after the last limitation (d) or after 36.3 block
    last = p
    for j in range(idx + 1, idx + 10):
        t = doc.paragraphs[j].text.strip()
        if t.startswith("Artículo 37°"):
            break
        if t:
            last = doc.paragraphs[j]
    if already_has_next(doc, idx, extra[:40]):
        return 0
    insert_after(last, extra)
    return 1


def expand_art38_prudence(doc: Document) -> int:
    start = body_start_index(doc)
    for i, p in enumerate(doc.paragraphs):
        if i < start or p.style.name.startswith("toc"):
            continue
        t = p.text.strip()
        if t.startswith("Cuando un proyecto con IA pueda generar beneficios relevantes"):
            extra = (
                "Ese criterio de prudencia implica gradualidad: validación previa, implementación "
                "piloto cuando sea posible, registro de incidentes y posibilidad de suspender el "
                "proyecto si el daño potencial a personas o comunidades no puede mitigarse de "
                "modo razonable. La celeridad de publicación o de transferencia no justifica "
                "adelantar resultados no validados."
            )
            if already_has_next(doc, i, extra[:40]):
                return 0
            insert_after(p, extra)
            return 1
    return 0


def in_cap_vi(doc, i, start):
    in_vi = False
    for j, p in enumerate(doc.paragraphs):
        if j < start:
            continue
        t = p.text.strip()
        if t.startswith("CAPÍTULO VI.") and not p.style.name.startswith("toc"):
            in_vi = True
        if t.startswith("CAPÍTULO VII.") and not p.style.name.startswith("toc"):
            in_vi = False
        if j == i:
            return in_vi
    return False


def fix_lists_and_typos(doc: Document) -> int:
    n = 0
    start = body_start_index(doc)
    in_vi = False
    for p in doc.paragraphs:
        if p.style.name.startswith("toc"):
            continue
        t = p.text.strip()
        if t.startswith("CAPÍTULO VI."):
            in_vi = True
        if t.startswith("CAPÍTULO VII."):
            break
        if not in_vi:
            continue
        if t in LIST_FIXES:
            new = LIST_FIXES[t]
            if p.runs:
                p.runs[0].text = new
                for r in p.runs[1:]:
                    r.text = ""
            else:
                p.text = new
            n += 1
        for old, new in TEXT_REPLACES:
            if old in p.text:
                nt = p.text.replace(old, new)
                if p.runs:
                    p.runs[0].text = nt
                    for r in p.runs[1:]:
                        r.text = ""
                else:
                    p.text = nt
                n += 1
    return n


def main():
    backup = TARGET.replace(".docx", f" - backup pre-capVI {date.today().isoformat()}.docx")
    shutil.copy2(TARGET, backup)
    doc = Document(TARGET)
    stats = {
        "insertions": apply_insertions(doc),
        "ethics_heading": add_ethics_heading(doc),
        "art36": expand_art36(doc),
        "art38": expand_art38_prudence(doc),
        "lists_typos": fix_lists_and_typos(doc),
    }
    doc.save(TARGET)
    print(f"Backup: {backup}")
    print(f"Saved: {TARGET}")
    for k, v in stats.items():
        print(f"  {k}: {v}")


if __name__ == "__main__":
    main()
