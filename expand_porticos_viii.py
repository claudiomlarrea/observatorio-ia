#!/usr/bin/env python3
"""Add chapter openings (pórticos) and Cap VIII role prose to Marco 2026."""
from __future__ import annotations

import shutil
from datetime import date

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph

TARGET = (
    "/Users/claudiolarrea/Library/CloudStorage/OneDrive-Personal/"
    "16 Secretaría de Investigación/60 Observatorio de Inteligencia Artificial/"
    "MARCO DE GOBERNANZA DE LA IA/Documento FINAL/"
    "Marco de Gobernanza IA - UCCuyo 2026.docx"
)

PORTICOS = {
    "CAPÍTULO I. DISPOSICIONES GENERALES": [
        "Este Capítulo fija el objeto, el alcance y las claves de lectura del Marco. "
        "Sin esas definiciones comunes, el resto de las normas se fragmentaría en prácticas "
        "de cátedra, de sede o de área, y perdería fuerza institucional.",
        "Las disposiciones que siguen no agotan la vida universitaria: la ordenan. Definen "
        "a quién alcanza la norma, con qué referentes se interpreta y qué principios rigen "
        "cuando una herramienta nueva no está todavía descrita en un artículo posterior.",
    ],
    "CAPÍTULO II. FUNDAMENTOS, PRINCIPIOS E IDENTIDAD INSTITUCIONAL": [
        "Antes de regular usos y procedimientos, la Universidad declara desde dónde piensa "
        "la Inteligencia Artificial: la persona humana, la búsqueda de la verdad, el bien "
        "común y la no sustitución del juicio responsable.",
        "Este Capítulo no es un preámbulo decorativo. Orienta la interpretación de todo el "
        "Marco cuando haya tensión entre eficiencia tecnológica, integridad académica y "
        "cuidado de las personas.",
    ],
    "CAPÍTULO III. GOBERNANZA INSTITUCIONAL DE LA INTELIGENCIA ARTIFICIAL": [
        "La gobernanza de la IA es el modo en que la Universidad decide, autoriza, registra "
        "y corrige el uso de sistemas algorítmicos, con responsabilidad humana explícita y "
        "criterios comunes para todas las sedes.",
        "Los artículos de este Capítulo describen el flujo institucional —quién decide, quién "
        "asesora, quién ejecuta— y los instrumentos que evitan tanto la prohibición absoluta "
        "como la adopción de hecho, sin evaluación ni responsable identificable.",
    ],
    "CAPÍTULO IV. USOS PERMITIDOS, CONDICIONADOS Y PROHIBIDOS": [
        "No todo lo técnicamente posible es institucionalmente admisible. Este Capítulo "
        "traduce los principios anteriores en un criterio de admisibilidad: finalidad "
        "legítima, proporcionalidad, supervisión humana y respeto de derechos.",
        "La enumeración de usos permitidos, condicionados y no permitidos busca dar certeza "
        "a la comunidad universitaria. No reemplaza el juicio del caso concreto, pero impide "
        "que la zona gris se convierta en autorización implícita.",
    ],
    "CAPÍTULO V. INTELIGENCIA ARTIFICIAL EN LA FUNCIÓN DOCENTE Y EL APRENDIZAJE": [
        "En la docencia, la IA puede enriquecer la enseñanza solo si fortalece la comprensión, "
        "el pensamiento crítico y el vínculo pedagógico, y no si simula competencias no "
        "adquiridas o desplaza la evaluación responsable.",
        "Este Capítulo ofrece reglas comunes para cátedras y estudiantes, y remite a los "
        "Anexos E, F, G y J para los modelos operativos. El cuerpo normativo fija el criterio; "
        "los anexos muestran cómo aplicarlo en el aula.",
    ],
    "CAPÍTULO VI. INTELIGENCIA ARTIFICIAL EN INVESTIGACIÓN, EXTENSIÓN Y TRANSFERENCIA": [
        "La investigación y la extensión pueden incorporar IA como herramienta, objeto de "
        "estudio o componente de un proyecto, siempre que se preserven la integridad "
        "científica, el consentimiento y el servicio a las comunidades.",
        "Este Capítulo articula deberes de investigadores, comités de ética y actividades "
        "con el medio. La celeridad de publicación o de transferencia no justifica omitir "
        "validación humana, evaluación de riesgo ni resguardo de datos.",
    ],
    "CAPÍTULO VII. INTELIGENCIA ARTIFICIAL EN LA GESTIÓN UNIVERSITARIA": [
        "La gestión universitaria puede servirse de la IA para mayor eficiencia e información, "
        "pero las decisiones de alto impacto —académicas, laborales, disciplinarias o sobre "
        "beneficios— no se delegan a un sistema.",
        "Este Capítulo regula finalidades admisibles, asistencia automatizada, contratación "
        "de proveedores y soberanía de datos. Unifica requisitos que antes estaban dispersos, "
        "para que ningún convenio o plataforma se incorpore al margen de la gobernanza.",
    ],
    "CAPÍTULO VIII. RESPONSABILIDADES Y DEBERES": [
        "La existencia de una herramienta de IA no diluye la imputación. Cada actor —autoridad, "
        "docente, investigador, estudiante, personal no docente o tercero— responde según su "
        "rol, sus deberes de cuidado y el poder efectivo de control que tiene sobre el uso.",
        "Este Capítulo no crea un régimen paralelo al disciplinario o laboral vigente: lo "
        "precisa para el contexto algorítmico. Su finalidad primaria es formativa y preventiva; "
        "la sanción interviene cuando el incumplimiento lo exige.",
    ],
    "CAPÍTULO IX. DERECHOS, GARANTÍAS Y MECANISMOS DE RESGUARDO": [
        "Frente al uso de IA, la comunidad universitaria no es solo destinataria de deberes: "
        "es titular de derechos a la información, a la explicación, a la revisión humana, a "
        "la privacidad y a un entorno digital no discriminatorio.",
        "Este Capítulo organiza esas garantías y las vías para hacerlas valer —consulta, "
        "reporte, impugnación—, de modo que el Marco no se reduzca a un catálogo de prohibiciones "
        "sin tutela efectiva.",
    ],
    "CAPÍTULO X. FORMACIÓN, CULTURA INSTITUCIONAL Y BUENAS PRÁCTICAS": [
        "Ninguna norma de IA se sostiene si la comunidad no comprende sus alcances. La "
        "alfabetización, la formación por niveles y las comunidades de práctica son parte "
        "de la gobernanza, no un complemento opcional de capacitación.",
        "Este Capítulo fija objetivos formativos, responsables y criterios de comunicación "
        "institucional, para que las reglas no queden libradas al rumor de cátedra ni a la "
        "desigualdad de acceso a criterios claros entre sedes y facultades.",
    ],
    "CAPÍTULO XI. SEGUIMIENTO, EVALUACIÓN Y MEJORA CONTINUA": [
        "Un Marco de IA que no se revisa envejece con la tecnología. El seguimiento no es "
        "control policial: es aprendizaje institucional, evidencia de cumplimiento e "
        "identificación de brechas para adecuar prácticas y anexos.",
        "Este Capítulo asigna responsables, indicadores flexibles y vías de actualización, "
        "con proporcionalidad: más intensidad donde hay mayor riesgo para personas, integridad "
        "académica o datos, y menos carga donde el uso es de bajo impacto.",
    ],
    "CAPÍTULO XII. DISPOSICIONES FINALES": [
        "Las disposiciones finales precisan la naturaleza jurídica del Marco, su relación con "
        "otras normas universitarias, la vigencia, la difusión y el régimen aplicable ante "
        "incumplimientos o lagunas.",
        "No desarrollan usos ni procedimientos nuevos. Cierran el texto para que pueda "
        "aplicarse, actualizarse por anexos y convivir con el Estatuto, la Ordenanza de "
        "Docencia y la normativa de investigación ya vigentes.",
    ],
}

VIII_INTROS = [
    (
        "Artículo 46°. Principio de responsabilidad institucional",
        "La Universidad no puede trasladar a un proveedor, a una cátedra o a un usuario "
        "individual el deber de gobernar la IA. Este artículo fija el piso institucional: "
        "dictar la norma, formar, registrar, proteger datos y preservar supervisión humana "
        "en las decisiones de impacto.",
    ),
    (
        "Artículo 47°. Responsabilidades de autoridades, TIC y gestores",
        "Quienes dirigen unidades y sistemas tienen un deber de vigilancia reforzado: no "
        "basta con no usar mal la IA; corresponde impedir que se consoliden prácticas "
        "informales de alto riesgo y asegurar que cada sistema institucional tenga un "
        "responsable humano y una vía de autorización.",
    ),
    (
        "Artículo 48°. Responsabilidades de docentes",
        "El docente es el primer intérprete del Marco en el aula. Le corresponde fijar "
        "reglas claras de uso, diseñar evaluaciones que sigan midiendo aprendizaje real y "
        "proteger los datos de los estudiantes, sin delegar la acreditación final a un sistema.",
    ),
    (
        "Artículo 49°. Responsabilidades de investigadores y estudiantes",
        "En investigación, la autoría y la veracidad no se delegan. En el aprendizaje, el "
        "estudiante sigue siendo responsable de comprender y poder defender su trabajo. Este "
        "artículo precisa ambos deberes para que la asistencia algorítmica no se confunda "
        "con sustitución de la integridad académica o científica.",
    ),
    (
        "Artículo 50°. Responsabilidades de personal no docente y terceros",
        "Gran parte del riesgo institucional de la IA está en la gestión cotidiana de "
        "legajos, expedientes y plataformas. El personal de apoyo y los proveedores no son "
        "ajenos al Marco: se les exige herramienta autorizada, confidencialidad y, en el "
        "caso de terceros, condiciones contractuales de auditoría y no uso secundario de datos.",
    ),
    (
        "Artículo 51°. Responsabilidad por decisiones, deberes comunes e imputación",
        "Cuando intervienen varios actores, la imputación no se diluye: se reparte según el "
        "deber de cada rol y según quién podía controlar el riesgo. “Lo decidió el algoritmo” "
        "no es eximente. Este artículo formula esa regla común a toda la comunidad universitaria.",
    ),
    (
        "Artículo 52°. Reportes, cooperación y consecuencias del incumplimiento",
        "Sin canales de consulta y reporte, las responsabilidades quedan en el papel. Este "
        "artículo cierra el Capítulo con vías de alerta, cooperación entre áreas y un abanico "
        "de consecuencias —de la orientación correctiva a la sanción— cuya finalidad primera "
        "es proteger personas y procesos, no solo reprimir.",
    ),
]


def insert_after(anchor: Paragraph, text: str) -> Paragraph:
    new_el = OxmlElement("w:p")
    anchor._p.addnext(new_el)
    p = Paragraph(new_el, anchor._parent)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.line_spacing = 1.5
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    return p


def body_start_index(doc: Document) -> int:
    for i, p in enumerate(doc.paragraphs):
        if "VERSIÓN 2026" in p.text or "VERSIÓN 2.0 DEPURADA" in p.text:
            return i
    return 0


def already_has_next(doc: Document, idx: int, snippet: str) -> bool:
    for j in range(idx + 1, min(idx + 5, len(doc.paragraphs))):
        if snippet[:45] in doc.paragraphs[j].text:
            return True
    return False


def add_porticos(doc: Document) -> int:
    n = 0
    start = body_start_index(doc)
    # collect heading indices first (mutating while iterating is messy)
    targets = []
    for i, p in enumerate(doc.paragraphs):
        if i < start or p.style.name.startswith("toc"):
            continue
        t = p.text.strip().split("\t")[0]
        if t in PORTICOS and p.style.name == "Heading 1":
            targets.append((i, t))
    # insert from the end so indices stay valid... actually insert_after only affects XML after
    # current node; later paragraph objects in python-docx list can become stale.
    # Safer: process from last chapter to first.
    for i, title in reversed(targets):
        p1, p2 = PORTICOS[title]
        # re-find live paragraph by text after previous inserts
        live = None
        for q in doc.paragraphs:
            if q.style.name == "Heading 1" and q.text.strip().split("\t")[0] == title:
                # skip TOC: Heading 1 in body after start
                live = q
        if live is None:
            continue
        live_idx = None
        for j, q in enumerate(doc.paragraphs):
            if q._element is live._element:
                live_idx = j
                break
        if live_idx is None or already_has_next(doc, live_idx, p1):
            continue
        # insert p2 first after heading, then p1 after heading so order is p1, p2
        insert_after(live, p2)
        insert_after(live, p1)
        n += 1
    return n


def add_viii_intros(doc: Document) -> int:
    n = 0
    start = body_start_index(doc)
    for title, prose in reversed(VIII_INTROS):
        live = None
        live_idx = None
        for j, q in enumerate(doc.paragraphs):
            if j < start or q.style.name.startswith("toc"):
                continue
            t = q.text.strip().split("\t")[0]
            if t.startswith(title) and q.style.name == "Heading 2":
                live = q
                live_idx = j
        if live is None:
            print("MISSING", title)
            continue
        if already_has_next(doc, live_idx, prose):
            continue
        insert_after(live, prose)
        n += 1
    return n


def main():
    backup = TARGET.replace(".docx", f" - backup pre-porticos {date.today().isoformat()}.docx")
    shutil.copy2(TARGET, backup)
    doc = Document(TARGET)
    stats = {
        "porticos": add_porticos(doc),
        "viii_intros": add_viii_intros(doc),
    }
    doc.save(TARGET)
    print(f"Backup: {backup}")
    print(f"Saved: {TARGET}")
    for k, v in stats.items():
        print(f"  {k}: {v}")


if __name__ == "__main__":
    main()
