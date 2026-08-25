#!/usr/bin/env python3
"""One-step editorial fix for Marco de Gobernanza IA - UCCuyo 2026.docx."""
from __future__ import annotations

import re
import shutil
from datetime import date

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.text.paragraph import Paragraph

from update_v2_references import (
    SUB_RENUMBER,
    replace_text,
    renumber_in_article,
)

TARGET = (
    "/Users/claudiolarrea/Library/CloudStorage/OneDrive-Personal/"
    "16 Secretaría de Investigación/60 Observatorio de Inteligencia Artificial/"
    "MARCO DE GOBERNANZA DE LA IA/Documento FINAL/"
    "Marco de Gobernanza IA - UCCuyo 2026.docx"
)

VERSION_NOTE = (
    "VERSIÓN 2026 — Este texto reorganiza el Marco anterior (v1) conforme a criterios de "
    "técnica normativa, reducción de redundancias y mayor operabilidad, preservando principios, "
    "garantías y anexos operativos. Las referencias internas corresponden a la numeración de esta versión."
)

# Targeted fixes (order matters for overlapping patterns).
EXACT_REPLACEMENTS = [
    (
        "conforme a la jerarquía del Artículo 5.4",
        "conforme a la jerarquía del numeral 4.8",
    ),
    (
        "con el mapa (Artículo 18°) y el Registro institucional (Artículo 20.4)",
        "con el mapa y el Registro institucional (Artículo 15°)",
    ),
    (
        "usos expresamente prohibidos en el Artículo 28°",
        "usos expresamente prohibidos en el Artículo 24°",
    ),
    (
        "Los usos inadmisibles de referencia del Artículo 19.5",
        "Los usos inadmisibles de referencia del Artículo 24°",
    ),
    (
        "actualizarse conforme al Artículo 134,",
        "actualizarse conforme al Artículo 72°,",
    ),
    ("(9.17 y art. 25.4.e)", "(Artículo 29° y numeral 25.4)"),
    ("(9.17)", "(Artículo 29°)"),
    ("(art. 25.4)", "(numeral 25.4)"),
    ("(art. 25.2)", "(Artículo 22°)"),
    ("(art. 25.4.e)", "(numeral 25.4)"),
    ("artículo 41.6", "numeral 31.6"),
    ("artículo 30.2", "Artículo 30°"),
    ("artículo 16°", "Artículo 16°"),
    ("artículo 17.2", "Artículo 17°"),
    ("(art. 14.2)", "(Artículo 14°)"),
    ("(art. 14.3)", "(Artículo 14°)"),
    ("(art. 14.5 y Cap. IV)", "(Artículo 24° y Cap. IV)"),
    ("(art. 14.6)", "(Artículo 14°)"),
    ("artículo 31°", "Artículo 31°"),
    ("artículo 31.6", "numeral 31.6"),
    ("artículo 28.2", "Artículo 28°"),
    ("artículo 57.5", "Artículo 57°"),
    ("(art. 38.5)", "(Artículo 36°)"),
    ("Niveles preuniversitarios. Colegios", "Niveles preuniversitarios. Colegios"),  # anchor
]

SUBSECTION_RE = re.compile(r"^\d+\.\d+\.\s")


def delete_paragraph(p: Paragraph):
    p._element.getparent().remove(p._element)


def replace_in_paragraph(p: Paragraph, old: str, new: str) -> bool:
    if old not in p.text:
        return False
    new_text = p.text.replace(old, new)
    if p.runs:
        p.runs[0].text = new_text
        for r in p.runs[1:]:
            r.text = ""
    else:
        p.text = new_text
    return True


def apply_exact_replacements(doc: Document) -> int:
    n = 0
    for p in doc.paragraphs:
        if p.style.name.startswith("toc"):
            continue
        for old, new in EXACT_REPLACEMENTS:
            if old.endswith("Colegios") and old == new:
                continue
            if replace_in_paragraph(p, old, new):
                n += 1
    return n


def dedupe_version_notes(doc: Document) -> int:
    seen = False
    removed = 0
    to_delete = []
    for p in doc.paragraphs:
        if "VERSIÓN 2.0 DEPURADA" in p.text or "VERSIÓN 2026" in p.text:
            if seen:
                to_delete.append(p)
            else:
                seen = True
                if p.runs:
                    p.runs[0].text = VERSION_NOTE
                    for r in p.runs[1:]:
                        r.text = ""
                else:
                    p.text = VERSION_NOTE
                for r in p.runs:
                    r.bold = False
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for p in to_delete:
        delete_paragraph(p)
        removed += 1
    return removed


def fix_annex_references(doc: Document) -> int:
    n = 0
    in_annex = False
    for p in doc.paragraphs:
        t = p.text.strip()
        if t == "ANEXOS" and p.style.name != "toc 1":
            in_annex = True
            continue
        if in_annex and t == "BIBLIOGRAFÍA" and p.style.name == "Normal":
            break
        if not in_annex or not t or p.style.name.startswith("toc"):
            continue
        new = replace_text(p.text)
        if new != p.text:
            if p.runs:
                p.runs[0].text = new
                for r in p.runs[1:]:
                    r.text = ""
            else:
                p.text = new
            n += 1
    return n


def renumber_merged_body(doc: Document) -> int:
    n = 0
    in_body = False
    current_art = None
    for p in doc.paragraphs:
        t = p.text.strip()
        if "VERSIÓN 2026" in t or "VERSIÓN 2.0" in t:
            in_body = True
            continue
        if t == "ANEXOS" and not p.style.name.startswith("toc"):
            in_body = False
            current_art = None
            continue
        if not in_body or p.style.name.startswith("toc"):
            continue
        m = re.match(r"Artículo\s+(\d+)\s*[°.]?", t)
        if m and len(t) < 120:
            current_art = int(m.group(1))
            continue
        if current_art in SUB_RENUMBER:
            new = renumber_in_article(p.text, current_art)
            if new != p.text:
                if p.runs:
                    p.runs[0].text = new
                    for r in p.runs[1:]:
                        r.text = ""
                else:
                    p.text = new
                n += 1
    return n


def unbold_subsections(doc: Document) -> int:
    n = 0
    for p in doc.paragraphs:
        if p.style.name.startswith("toc"):
            continue
        t = p.text.strip().split("\t")[0]
        if not t or t in ("ÍNDICE RESUMIDO", "ÍNDICE GENERAL", "ANEXOS", "BIBLIOGRAFÍA"):
            continue
        if t.startswith("CAPÍTULO ") or re.match(r"Artículo\s+\d+", t):
            continue
        if SUBSECTION_RE.match(t) or re.match(r"Anexo [A-K]\.", t):
            for r in p.runs:
                r.bold = False
            if SUBSECTION_RE.match(t):
                try:
                    p.style = "Heading 3"
                except Exception:
                    pass
            n += 1
    return n


def set_left_align(doc: Document):
    in_scope = False
    for p in doc.paragraphs:
        t = p.text.strip()
        if "VERSIÓN 2026" in t or "VERSIÓN 2.0" in t:
            in_scope = True
        if not in_scope:
            continue
        if p.style.name.startswith("toc"):
            continue
        if t in ("ÍNDICE RESUMIDO", "ÍNDICE GENERAL"):
            continue
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT


def main():
    backup = TARGET.replace(".docx", f" - backup {date.today().isoformat()}.docx")
    shutil.copy2(TARGET, backup)

    doc = Document(TARGET)
    stats = {
        "exact": apply_exact_replacements(doc),
        "version_removed": dedupe_version_notes(doc),
        "annex_refs": fix_annex_references(doc),
        "body_renumber": renumber_merged_body(doc),
        "unbold": unbold_subsections(doc),
    }
    set_left_align(doc)
    doc.save(TARGET)

    print(f"Backup: {backup}")
    print(f"Saved: {TARGET}")
    for k, v in stats.items():
        print(f"  {k}: {v}")


if __name__ == "__main__":
    main()
