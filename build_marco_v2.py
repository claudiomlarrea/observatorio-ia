#!/usr/bin/env python3
"""Build Marco de Gobernanza IA v2.0 — depuración estructural (criterio Maluf)."""
from __future__ import annotations

import re
import shutil
from copy import deepcopy
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph

SRC = (
    "/Users/claudiolarrea/Library/CloudStorage/OneDrive-Personal/"
    "16 Secretaría de Investigación/60 Observatorio de Inteligencia Artificial/"
    "MARCO DE GOBERNANZA DE LA IA/"
    "Marco de Gobernanza, Ética y Uso Responsable de la IA - UCCuyo.docx"
)
OUT_ONEDRIVE = SRC.replace(".docx", " - v2.0 depurada.docx")
OUT_LOCAL = (
    "/Users/claudiolarrea/Documents/Observatorio/"
    "Marco de Gobernanza, Ética y Uso Responsable de la IA - v2.0 depurada.docx"
)
BODY_START_MARKER = "CAPÍTULO I. DISPOSICIONES GENERALES"
ANNEX_MARKER = "ANEXOS"


def delete_paragraph(paragraph):
    paragraph._element.getparent().remove(paragraph._element)


def add_para(doc, anchor, text, style=None, bold=False, justify=True):
    """Insert paragraph after anchor element."""
    new_el = OxmlElement("w:p")
    anchor._p.addnext(new_el)
    p = Paragraph(new_el, anchor._parent)
    if style:
        p.style = style
    run = p.add_run(text)
    run.bold = bold
    if justify:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 1.5
    return p


def extract_articles(doc: Document) -> dict[int, dict]:
    arts: dict[int, dict] = {}
    current = None
    title = ""
    buf: list[str] = []
    for p in doc.paragraphs:
        t = p.text.strip()
        if "\t" in t:
            continue
        m = re.match(r"Artículo\s+(\d+)\s*[°.]?\s*(.*)", t)
        if m and len(t) < 120:
            if current is not None:
                arts[current] = {"title": title, "body": "\n".join(buf).strip()}
            current = int(m.group(1))
            title = m.group(2).strip().lstrip(". ")
            buf = []
            continue
        if current is not None:
            if t.startswith("CAPÍTULO ") or t == ANNEX_MARKER:
                arts[current] = {"title": title, "body": "\n".join(buf).strip()}
                current = None
                buf = []
            elif t:
                buf.append(t)
    if current is not None:
        arts[current] = {"title": title, "body": "\n".join(buf).strip()}
    return arts


def merge_bodies(arts: dict, nums: list[int], sep="\n\n") -> str:
    parts = []
    for n in nums:
        b = arts.get(n, {}).get("body", "").strip()
        if b:
            parts.append(b)
    return sep.join(parts)


def strip_sections(body: str, starts: list[str]) -> str:
    """Remove numbered sections starting with given prefixes."""
    lines = body.split("\n")
    out = []
    skip = False
    for line in lines:
        if any(re.match(rf"^{re.escape(s)}", line.strip()) for s in starts):
            skip = True
            continue
        if skip and re.match(r"^\d+\.\d+\.", line.strip()):
            skip = False
        if not skip:
            out.append(line)
    return "\n".join(out).strip()


def extract_section(body: str, start_prefix: str) -> str:
    lines = body.split("\n")
    out = []
    capture = False
    for line in lines:
        if line.strip().startswith(start_prefix):
            capture = True
        elif capture and re.match(r"^\d+\.\d+\.", line.strip()) and not line.strip().startswith(start_prefix):
            break
        if capture:
            out.append(line)
    return "\n".join(out).strip()


def condense_principles(body: str) -> str:
    """Keep 9.1-9.14 and 9.16-9.18; drop operational 9.15 and 9.17."""
    intro = ""
    if "principios éticos" in body.lower()[:200]:
        intro = body.split("9.1.")[0].strip()
        intro = re.sub(
            r"Los numerales 9\.15 a 9\.18.*",
            "Los criterios operativos de declaración de uso y de protección de datos confidenciales se regulan en el Capítulo V.",
            intro,
        )
    kept = []
    for m in re.finditer(r"(9\.\d+\..*?)(?=9\.\d+\.|$)", body, re.S):
        block = m.group(1).strip()
        if block.startswith("9.15.") or block.startswith("9.17."):
            continue
        kept.append(block)
    return (intro + "\n" + "\n".join(kept)).strip()


def build_v2_spec(arts: dict) -> list[dict]:
    """Return list of chapters with articles for v2.0."""
    op_915 = extract_section(arts[9]["body"], "9.15.")
    op_917 = extract_section(arts[9]["body"], "9.17.")

    spec: list[dict] = []

    def ch(num_roman, title, articles):
        spec.append({"type": "chapter", "title": f"CAPÍTULO {num_roman}. {title}"})
        for art in articles:
            spec.append({"type": "article", **art})

    # CAP I — 5 arts
    ch(
        "I",
        "DISPOSICIONES GENERALES",
        [
            {"n": 1, "title": "Objeto", "body": arts[1]["body"]},
            {
                "n": 2,
                "title": "Alcance y naturaleza del Marco",
                "body": arts[2]["body"],
            },
            {
                "n": 3,
                "title": "Ámbito de aplicación y definiciones",
                "body": merge_bodies(arts, [3, 4]),
            },
            {
                "n": 4,
                "title": "Marco normativo y referentes externos",
                "body": merge_bodies(arts, [5, 11])
                + "\n\n4.1. Jerarquía de referentes\n"
                "Para interpretar este Marco se distinguirá: (i) normativa institucional aplicable con efectos "
                "vinculantes; (ii) referentes externos orientadores en materia de ética, derechos y gobernanza; "
                "y (iii) antecedentes comparados de otras universidades e instituciones, sin recepción automática "
                "ni alteración de la autonomía universitaria.",
            },
            {"n": 5, "title": "Principios de interpretación", "body": arts[6]["body"]},
        ],
    )

    # CAP II — 5 arts
    ch(
        "II",
        "FUNDAMENTOS, PRINCIPIOS E IDENTIDAD INSTITUCIONAL",
        [
            {"n": 6, "title": "Identidad y sentido institucional", "body": arts[7]["body"]},
            {
                "n": 7,
                "title": "Principios antropológicos, teológicos y humanistas",
                "body": arts[8]["body"],
            },
            {
                "n": 8,
                "title": "Principios éticos orientadores",
                "body": condense_principles(arts[9]["body"]),
            },
            {
                "n": 9,
                "title": "Principio rector de no sustitución del juicio humano",
                "body": arts[10]["body"],
            },
            {
                "n": 10,
                "title": "Finalidad formativa de la gobernanza de la IA",
                "body": arts[12]["body"],
            },
        ],
    )

    # CAP III — 10 arts
    art12_body = (
        arts[14]["body"]
        + "\n\n14.4. Responsabilidades institucionales\n"
        + arts[74]["body"]
        + "\n\n14.5. Funciones del Observatorio\n"
        + arts[75]["body"]
    )
    art15_body = (
        "La Universidad dispondrá de un sistema integrado de trazabilidad humano-algorítmica, "
        "conformado por la Matriz institucional de delegación, el Mapa de decisiones críticas "
        "y el Registro de sistemas de IA institucionales. Los datos comunes se consignarán una "
        "sola vez en la Matriz (Anexo B) y alimentarán el mapa y el registro sin cargas duplicadas.\n\n"
        + arts[18]["body"]
        + "\n\n"
        + arts[20]["body"]
    )
    ch(
        "III",
        "GOBERNANZA INSTITUCIONAL DE LA INTELIGENCIA ARTIFICIAL",
        [
            {"n": 11, "title": "Modelo de gobernanza", "body": arts[13]["body"]},
            {"n": 12, "title": "Órganos competentes, roles y responsabilidades", "body": art12_body},
            {"n": 13, "title": "Responsabilidad humana explícita", "body": arts[15]["body"]},
            {
                "n": 14,
                "title": "Grado de delegación y matriz institucional",
                "body": merge_bodies(arts, [16, 17]),
            },
            {
                "n": 15,
                "title": "Mapa, registro y trazabilidad de sistemas",
                "body": art15_body,
            },
            {"n": 16, "title": "Evaluación de riesgo e impacto", "body": arts[19]["body"]},
            {
                "n": 17,
                "title": "Autorización y ciclo de vida de sistemas institucionales",
                "body": "Todo sistema de IA institucional seguirá el ciclo de identificación, evaluación, "
                "autorización, registro, monitoreo y baja previsto en este Capítulo y en el Anexo D. "
                "Los tramos operativos del registro y seguimiento se integran en el Artículo 15°.\n\n"
                + re.sub(r"20\.4\..*?(?=20\.5\.|$)", "", arts[20]["body"], flags=re.S).strip(),
            },
            {"n": 18, "title": "Usos individuales y zonas grises", "body": arts[21]["body"]},
            {"n": 19, "title": "Coordinación entre sedes", "body": arts[22]["body"]},
            {"n": 20, "title": "Aprendizaje institucional", "body": arts[23]["body"]},
        ],
    )

    # CAP IV — 8 arts
    ch(
        "IV",
        "USOS PERMITIDOS, CONDICIONADOS Y PROHIBIDOS",
        [
            {"n": 21, "title": "Criterio general de admisibilidad", "body": arts[24]["body"]},
            {"n": 22, "title": "Finalidades legítimas y usos permitidos", "body": merge_bodies(arts, [25, 26])},
            {"n": 23, "title": "Usos condicionados", "body": arts[27]["body"]},
            {"n": 24, "title": "Usos no permitidos", "body": arts[28]["body"]},
            {
                "n": 25,
                "title": "Limitaciones transversales y proporcionalidad",
                "body": merge_bodies(arts, [29, 30]),
            },
            {"n": 26, "title": "Autorización excepcional", "body": arts[32]["body"]},
            {"n": 27, "title": "Deber de adecuación de prácticas existentes", "body": arts[33]["body"]},
        ],
    )

    # CAP V — 5 arts
    docencia_ops = ""
    if op_915:
        docencia_ops += op_915 + "\n\n"
    if op_917:
        docencia_ops += op_917 + "\n\n"
    ch(
        "V",
        "INTELIGENCIA ARTIFICIAL EN LA FUNCIÓN DOCENTE Y EL APRENDIZAJE",
        [
            {
                "n": 28,
                "title": "Principio pedagógico y alcance",
                "body": merge_bodies(arts, [34, 35]),
            },
            {
                "n": 29,
                "title": "Reglas generales para docentes y estudiantes",
                "body": merge_bodies(arts, [36, 37])
                + "\n\n"
                + docencia_ops
                + "Los procedimientos, ejemplos y modelos operativos se desarrollan en los Anexos E, F y G.",
            },
            {
                "n": 30,
                "title": "Producción académica, transparencia e integridad",
                "body": merge_bodies(arts, [38, 39, 40]),
            },
            {
                "n": 31,
                "title": "Evaluación e información en el aula",
                "body": merge_bodies(arts, [41, 42]),
            },
            {
                "n": 32,
                "title": "Acompañamiento y resolución de dudas",
                "body": merge_bodies(arts, [43, 44])
                + "\n\nEn niveles preuniversitarios rige además el Anexo J.",
            },
        ],
    )

    # CAP VI — 6 arts
    ch(
        "VI",
        "INTELIGENCIA ARTIFICIAL EN INVESTIGACIÓN, EXTENSIÓN Y TRANSFERENCIA",
        [
            {"n": 33, "title": "Principio rector y alcance", "body": merge_bodies(arts, [45, 46])},
            {
                "n": 34,
                "title": "Investigación, integridad y publicación",
                "body": merge_bodies(arts, [47, 48]),
            },
            {
                "n": 35,
                "title": "Datos, consentimiento y comités de ética",
                "body": merge_bodies(arts, [49, 50]),
            },
            {"n": 36, "title": "Extensión y vinculación con el medio", "body": arts[51]["body"]},
            {
                "n": 37,
                "title": "Propiedad intelectual y transparencia",
                "body": merge_bodies(arts, [53, 54]),
            },
            {
                "n": 38,
                "title": "Responsabilidades, prohibiciones y prudencia investigativa",
                "body": merge_bodies(arts, [55, 56, 57]),
            },
        ],
    )

    # CAP VII — 8 arts (gestión + contratación unificada)
    contratacion = (
        "Este artículo unifica los requisitos para convenios, contratación y proveedores de IA.\n\n"
        + arts[52]["body"]
        + "\n\n"
        + arts[64]["body"]
    )
    ch(
        "VII",
        "INTELIGENCIA ARTIFICIAL EN LA GESTIÓN UNIVERSITARIA",
        [
            {"n": 39, "title": "Principio rector y alcance de la gestión", "body": merge_bodies(arts, [58, 59])},
            {
                "n": 40,
                "title": "Finalidades admisibles y sistemas de gestión",
                "body": merge_bodies(arts, [60, 61]),
            },
            {
                "n": 41,
                "title": "Decisiones de alto impacto y asistencia automatizada",
                "body": merge_bodies(arts, [62, 63]),
            },
            {
                "n": 42,
                "title": "Adquisición, contratación, convenios y proveedores",
                "body": contratacion,
            },
            {
                "n": 43,
                "title": "Registro, seguridad y resguardo de la información",
                "body": merge_bodies(arts, [65, 66]),
            },
            {
                "n": 44,
                "title": "Soberanía de datos y transparencia institucional",
                "body": merge_bodies(arts, [67, 68]),
            },
            {
                "n": 45,
                "title": "Roles, prohibiciones e implementación gradual",
                "body": merge_bodies(arts, [69, 70, 71]),
            },
        ],
    )

    # CAP VIII — 7 arts
    ch(
        "VIII",
        "RESPONSABILIDADES Y DEBERES",
        [
            {
                "n": 46,
                "title": "Principio de responsabilidad institucional",
                "body": merge_bodies(arts, [72, 73]),
            },
            {
                "n": 47,
                "title": "Responsabilidades de autoridades, TIC y gestores",
                "body": merge_bodies(arts, [76, 77]),
            },
            {"n": 48, "title": "Responsabilidades de docentes", "body": arts[78]["body"]},
            {
                "n": 49,
                "title": "Responsabilidades de investigadores y estudiantes",
                "body": merge_bodies(arts, [79, 80]),
            },
            {
                "n": 50,
                "title": "Responsabilidades de personal no docente y terceros",
                "body": merge_bodies(arts, [81, 82]),
            },
            {
                "n": 51,
                "title": "Responsabilidad por decisiones, deberes comunes e imputación",
                "body": merge_bodies(arts, [83, 84, 85, 88]),
            },
            {
                "n": 52,
                "title": "Reportes, cooperación y consecuencias del incumplimiento",
                "body": merge_bodies(arts, [86, 87]),
            },
        ],
    )

    # CAP IX — 10 arts
    ch(
        "IX",
        "DERECHOS, GARANTÍAS Y MECANISMOS DE RESGUARDO",
        [
            {
                "n": 53,
                "title": "Finalidad y titulares de derechos",
                "body": merge_bodies(arts, [90, 91]),
            },
            {
                "n": 54,
                "title": "Derecho a la información y a la explicación",
                "body": merge_bodies(arts, [92, 93]),
            },
            {
                "n": 55,
                "title": "Revisión humana e impugnación",
                "body": merge_bodies(arts, [94, 95]),
            },
            {"n": 56, "title": "Integridad académica equitativa", "body": arts[96]["body"]},
            {"n": 57, "title": "Privacidad y protección de datos", "body": arts[97]["body"]},
            {
                "n": 58,
                "title": "Prevención de sesgos y entorno digital seguro",
                "body": merge_bodies(arts, [98, 99]),
            },
            {
                "n": 59,
                "title": "Consultas, reporte y resolución de conflictos",
                "body": merge_bodies(arts, [100, 101, 102, 103]),
            },
            {
                "n": 60,
                "title": "Protección de denunciantes e incidentes",
                "body": merge_bodies(arts, [104, 105]),
            },
            {"n": 61, "title": "Interpretación pro persona", "body": arts[106]["body"]},
        ],
    )

    # CAP X — 6 arts
    ch(
        "X",
        "FORMACIÓN, CULTURA INSTITUCIONAL Y BUENAS PRÁCTICAS",
        [
            {
                "n": 62,
                "title": "Principio cultural y objetivos formativos",
                "body": merge_bodies(arts, [108, 109]),
            },
            {
                "n": 63,
                "title": "Alfabetización y niveles formativos",
                "body": merge_bodies(arts, [110, 111, 112]),
            },
            {
                "n": 64,
                "title": "Formación docente, estudiantil y de gestión",
                "body": merge_bodies(arts, [113, 114, 115]),
            },
            {
                "n": 65,
                "title": "Buenas prácticas, guías y comunidades de práctica",
                "body": merge_bodies(arts, [116, 117, 118]),
            },
            {
                "n": 66,
                "title": "Comunicación, responsables y cultura institucional",
                "body": merge_bodies(arts, [119, 120, 121, 122, 123]),
            },
        ],
    )

    # CAP XI — 7 arts
    ch(
        "XI",
        "SEGUIMIENTO, EVALUACIÓN Y MEJORA CONTINUA",
        [
            {
                "n": 67,
                "title": "Principio de mejora continua y objeto del seguimiento",
                "body": merge_bodies(arts, [125, 126]),
            },
            {"n": 68, "title": "Responsables del seguimiento", "body": arts[127]["body"]},
            {"n": 69, "title": "Indicadores de cumplimiento e impacto", "body": arts[128]["body"]},
            {
                "n": 70,
                "title": "Evidencia e informes institucionales",
                "body": merge_bodies(arts, [129, 130]),
            },
            {
                "n": 71,
                "title": "Auditoría y revisión de sistemas",
                "body": merge_bodies(arts, [131, 132]),
            },
            {
                "n": 72,
                "title": "Adecuación, actualización del Marco y anexos",
                "body": merge_bodies(arts, [133, 134, 135]),
            },
            {
                "n": 73,
                "title": "Pilotos, transparencia y proporcionalidad del control",
                "body": merge_bodies(arts, [136, 137, 138]),
            },
        ],
    )

    # CAP XII — 11 arts
    ch(
        "XII",
        "DISPOSICIONES FINALES",
        [
            {"n": 74, "title": "Naturaleza jurídica del Marco", "body": arts[140]["body"]},
            {"n": 75, "title": "Remisión a normativas vigentes", "body": arts[141]["body"]},
            {"n": 76, "title": "Régimen sancionatorio", "body": arts[142]["body"]},
            {"n": 77, "title": "Interpretación y lagunas", "body": arts[143]["body"]},
            {"n": 78, "title": "Anexos y desarrollo reglamentario", "body": arts[144]["body"]},
            {"n": 79, "title": "Cláusulas transitorias", "body": arts[145]["body"]},
            {"n": 80, "title": "Entrada en vigencia", "body": arts[146]["body"]},
            {"n": 81, "title": "Difusión obligatoria", "body": arts[147]["body"]},
            {"n": 82, "title": "Disposición de implementación", "body": arts[148]["body"]},
            {"n": 83, "title": "Derogación de criterios incompatibles", "body": arts[149]["body"]},
            {"n": 84, "title": "Cláusula de cierre", "body": arts[150]["body"]},
        ],
    )

    return spec


def find_body_bounds(doc: Document) -> tuple[int, int]:
    start = end = None
    for i, p in enumerate(doc.paragraphs):
        t = p.text.strip()
        if t == BODY_START_MARKER and "\t" not in t:
            start = i
        if start is not None and t == ANNEX_MARKER and "\t" not in t:
            end = i
            break
    if start is None or end is None:
        raise RuntimeError("No se encontraron límites del cuerpo normativo")
    return start, end


def build(output_path: str):
    doc_src = Document(SRC)
    arts = extract_articles(doc_src)
    spec = build_v2_spec(arts)
    art_count = sum(1 for x in spec if x["type"] == "article")
    print(f"v2.0 articles: {art_count}")

    shutil.copy2(SRC, output_path)
    doc = Document(output_path)
    start, end = find_body_bounds(doc)

    for idx in range(end - 1, start - 1, -1):
        delete_paragraph(doc.paragraphs[idx])

    anchor = doc.paragraphs[start - 1]
    note = (
        "VERSIÓN 2.0 DEPURADA — Este texto reorganiza el Marco anterior (v1) conforme a criterios de "
        "técnica normativa, reducción de redundancias y mayor operabilidad, preservando principios, "
        "garantías y anexos operativos. Las referencias internas corresponden a la numeración de esta versión."
    )
    current = add_para(doc, anchor, note, bold=True)
    current = add_para(doc, current, "")

    for item in spec:
        if item["type"] == "chapter":
            current = add_para(doc, current, item["title"], bold=True)
            current = add_para(doc, current, "")
        else:
            heading = f"Artículo {item['n']}°. {item['title']}"
            current = add_para(doc, current, heading, bold=True)
            for block in item["body"].split("\n\n"):
                block = block.strip()
                if block:
                    current = add_para(doc, current, block)
            current = add_para(doc, current, "")

    doc.save(output_path)
    print("Saved:", output_path)


if __name__ == "__main__":
    build(OUT_ONEDRIVE)
    build(OUT_LOCAL)
    print("Done.")
