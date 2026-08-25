#!/usr/bin/env python3
"""Fix broken article cross-refs, AI boilerplate intros, Art. 40 gap, and bibliography
in Marco de Gobernanza IA - UCCuyo 2026.docx."""
from __future__ import annotations

import re
import shutil
from datetime import date
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph

from rewrite_unique_intros import INTROS

TARGET = Path(
    "/Users/claudiolarrea/Library/CloudStorage/OneDrive-Personal/"
    "16 Secretaría de Investigación/60 Observatorio de Inteligencia Artificial/"
    "MARCO DE GOBERNANZA DE LA IA/Documento FINAL/"
    "Marco de Gobernanza IA - UCCuyo 2026.docx"
)
BACKUP_DIR = TARGET.parent / "Versiones previas"
WORKSPACE_COPY = Path(
    "/Users/claudiolarrea/Documents/Observatorio/"
    "Marco de Gobernanza IA - UCCuyo FINAL.docx"
)

# Exact string replacements (order matters).
EXACT = [
    # Art. 15 internal refs (Matriz = 14°, Evaluación = 16°)
    (
        "El mapa complementa la Matriz del Artículo 17°:",
        "El mapa complementa la Matriz del Artículo 14°:",
    ),
    (
        "El Registro se alimentará de la evaluación del Artículo 19° y de la Matriz del Artículo 17°,",
        "El Registro se alimentará de la evaluación del Artículo 16° y de la Matriz del Artículo 14°,",
    ),
    # Glossary / annex — old v1 numbers → 2026 numbering
    (
        "(arts. 41 a 115)",
        "(Arts. 62° a 66°)",
    ),
    (
        "(arts. 28 y 97.4 y 97.5)",
        "(Arts. 35° y 57°)",
    ),
    (
        "(arts. 24.2, 97.3 a 97.5; Anexo J)",
        "(Arts. 28°, 32° y 57°; Anexo J)",
    ),
    (
        "(arts. 25 y 41)",
        "(Arts. 30° y 31°)",
    ),
    (
        "(arts. 25 y 39)",
        "(Arts. 29° y 30°)",
    ),
    (
        "traduce los artículos 25 a 40 a una pauta de bolsillo",
        "traduce los Arts. 28° a 32° a una pauta de bolsillo",
    ),
    (
        "(arts. 25 y 40)",
        "(Arts. 29° y 30°)",
    ),
    (
        "Marco (arts. 25 a 40)",
        "Marco (Arts. 28° a 32°)",
    ),
    (
        "si hay riesgo grave (art. 39):",
        "si hay riesgo grave (Art. 60°):",
    ),
    (
        "medio o alto (art. 39).",
        "medio o alto (Art. 58°).",
    ),
    # Anexo D
    (
        "Este Anexo opera el Artículo 14° y el tramo de evaluación previa del Artículo 14°.",
        "Este Anexo opera el Artículo 14° (matriz) y el tramo de evaluación previa del Artículo 16°.",
    ),
    # Self-referential oddity in Art. 42
    (
        "se aplicarán supletoriamente los requisitos de los Artículos 42.2 y 42.3.",
        "se aplicarán los requisitos de los numerales 42.2 y 42.3.",
    ),
    # Typo / AI fragment leftovers
    (
        "norman tFI, tesis y trabajos finales",
        "normas TFI, tesis y trabajos finales",
    ),
]

# Formulaic tails to strip when a usable lead sentence exists.
FORMULA_STARTS = (
    "No se limita a enunciar criterios aislados",
    "Su objetivo es precisar el sentido del punto antes de las enumeraciones",
    "procura ordenar el sentido institucional del punto",
)

BOILER_MARKERS = (
    "orienta su interpretación práctica dentro del Marco",
    "ofrece una pauta institucional para su aplicación concreta",
    "No se limita a enunciar criterios aislados",
)

# Fallback intros when INTROS has no key (short, grammatical).
FALLBACK_BY_HEAD_PREFIX = {
    "12.3": "Las secretarías y áreas de gestión aplican este Marco en sus procesos, "
    "trámites y sistemas de soporte, con responsable humano identificable y "
    "proporcionalidad al riesgo. Las viñetas precisan deberes mínimos de esas áreas.",
    "12.4": "La comunidad universitaria —docentes, estudiantes, investigadores, personal "
    "no docente y terceros vinculados— queda alcanzada por las reglas de uso, "
    "transparencia y cuidado de datos. Las enumeraciones orientan deberes compartidos.",
    "14.2": "La delegación condicionada admite asistencia algorítmica solo cuando hay "
    "supervisión humana efectiva, trazabilidad y finalidad legítima. Las viñetas "
    "delimitan ese supuesto.",
    "14.3": "Las tareas delegables son de apoyo o bajo riesgo y no sustituyen juicios "
    "académicos, éticos o disciplinarios. Las enumeraciones fijan ese perímetro.",
    "15.3": "Se enumeran decisiones críticas de referencia para el mapa institucional. "
    "No es un catálogo cerrado: sirve para anticipar impactos y zonas grises.",
    "16.2": "Los criterios de valoración del riesgo e impacto se aplican de modo "
    "proporcional e integrado. Las viñetas fijan estándares mínimos de análisis.",
    "22.2": "En docencia y aprendizaje la IA puede apoyar, sin sustituir la evaluación "
    "ni la autoría estudiantil. Las reglas siguientes precisan usos admisibles.",
    "22.4": "En gestión y soporte institucional la IA puede optimizar información y "
    "trámites de bajo riesgo, sin automatizar exclusiones ni sanciones. Las viñetas "
    "delimitan ese alcance.",
    "23.1": "Los supuestos de condicionamiento exigen autorización, transparencia o "
    "controles reforzados según el riesgo. Las enumeraciones ordenan esos casos.",
    "26": "La autorización excepcional es restrictiva, fundada y revisable; no convierte "
    "en regla lo que el Marco prohíbe. Las condiciones siguientes son acumulativas.",
    "30.3": "En TFI, tesis y trabajos finales la IA, si se admite, es asistencia "
    "declarada y subordinada a la autoría y al criterio del director o tribunal. "
    "Las reglas precisan ese umbral de integridad.",
}


def set_text(p: Paragraph, text: str) -> None:
    """Replace paragraph visible text, clearing all runs (incl. hyperlink children)."""
    from docx.oxml import OxmlElement

    # Clear existing runs / hyperlink wrappers
    p_el = p._element
    for child in list(p_el):
        if child.tag == qn("w:r") or child.tag == qn("w:hyperlink"):
            p_el.remove(child)
    run = p.add_run(text)


def delete_paragraph(p: Paragraph) -> None:
    el = p._element
    parent = el.getparent()
    if parent is not None:
        parent.remove(el)


def insert_paragraph_before(paragraph: Paragraph, text: str, style: str | None = None) -> Paragraph:
    from docx.oxml import OxmlElement

    new_el = OxmlElement("w:p")
    paragraph._element.addprevious(new_el)
    np = Paragraph(new_el, paragraph._parent)
    if style:
        try:
            np.style = style
        except KeyError:
            pass
    run = np.add_run(text)
    if style and "Heading" in style:
        run.bold = True
    return np


def prev_heading(doc: Document, idx: int) -> str:
    for k in range(idx - 1, max(idx - 12, -1), -1):
        pk = doc.paragraphs[k]
        name = pk.style.name if pk.style else ""
        if name in ("Heading 2", "Heading 3"):
            return pk.text.strip().split("\t")[0]
    return ""


def is_boiler(text: str) -> bool:
    t = text.strip()
    return any(m in t for m in BOILER_MARKERS)


def strip_formula_tail(text: str) -> str | None:
    """If lead sentence before formula is grammatical enough, keep it."""
    t = text.strip()
    for marker in FORMULA_STARTS:
        if marker in t:
            before = t.split(marker, 1)[0].strip().rstrip(".;:")
            if not before:
                return None
            # Broken AI interpolations
            if re.search(
                r"\b(y orienta|y ofrece|norman|desarrolla en docencia)\b",
                before,
                re.I,
            ):
                return None
            if re.search(
                r"^(Las|La|El|Los|Puntos)\b.{0,90}\s+y\s+(orienta|ofrece)\b",
                before,
                re.I,
            ):
                return None
            if len(before.split()) < 6:
                return None
            if not before.endswith("."):
                before += "."
            return before
    return None


def fallback_intro(head: str) -> str:
    for prefix, text in FALLBACK_BY_HEAD_PREFIX.items():
        if head.startswith(prefix) or head.startswith(f"Artículo {prefix}"):
            return text
    # Generic from heading title
    title = re.sub(r"^(Artículo\s+\d+\s*[°º]\.?\s*|^\d+\.\d+\.?\s*)", "", head).strip()
    if not title:
        title = "este punto"
    title_l = title[0].lower() + title[1:] if title else "este punto"
    return (
        f"Este apartado fija los criterios aplicables a {title_l}. "
        f"Las enumeraciones posteriores constituyen estándares mínimos y deben "
        f"leerse de modo integrado, proporcional al riesgo y coherente con la "
        f"protección de las personas y la finalidad institucional."
    )


def fix_intros(doc: Document) -> dict:
    stats = {"intros_replaced": 0, "tails_stripped": 0, "deleted_empty": 0}
    # Iterate by index snapshot
    for i, p in enumerate(list(doc.paragraphs)):
        if p._element.getparent() is None:
            continue
        style = p.style.name if p.style else ""
        if style.startswith("toc") or style.startswith("Heading"):
            continue
        t = p.text.strip()
        if not t or not is_boiler(t):
            continue

        head = prev_heading(doc, i)
        if head in INTROS:
            set_text(p, INTROS[head])
            stats["intros_replaced"] += 1
            continue

        stripped = strip_formula_tail(t)
        if stripped:
            set_text(p, stripped)
            stats["tails_stripped"] += 1
            continue

        # Still broken / pure formula → fallback
        set_text(p, fallback_intro(head or "este punto"))
        stats["intros_replaced"] += 1
    return stats


def restore_articulo_40(doc: Document) -> bool:
    """Insert missing Artículo 40° heading before its lead content (or 40.1)."""
    for p in doc.paragraphs:
        st = p.style.name if p.style else ""
        if st.startswith("Heading") and p.text.strip().startswith("Artículo 40°"):
            return False

    # Prefer before the lead paragraph of Art. 40 (still under Art. 39 today)
    for p in doc.paragraphs:
        if p.text.strip().startswith(
            "Las disposiciones siguientes establecen los criterios y pautas aplicables "
            "a las finalidades admisibles"
        ):
            insert_paragraph_before(
                p,
                "Artículo 40°. Finalidades admisibles y sistemas de gestión",
                style="Heading 2",
            )
            return True

    for p in doc.paragraphs:
        if p.text.strip().startswith("40.1."):
            insert_paragraph_before(
                p,
                "Artículo 40°. Finalidades admisibles y sistemas de gestión",
                style="Heading 2",
            )
            return True
    return False


def remove_art15_cycle_duplicate(doc: Document) -> int:
    """Remove lifecycle block wrongly embedded in Art. 15 (belongs to Art. 17)."""
    in_art15 = False
    deleting = False
    to_delete: list[Paragraph] = []
    kept_registro = False

    for p in doc.paragraphs:
        t = p.text.strip()
        style = p.style.name if p.style else ""

        if style.startswith("Heading") and t.startswith("Artículo 15°"):
            in_art15 = True
            deleting = False
            continue
        if style.startswith("Heading") and t.startswith("Artículo 16°"):
            break
        if not in_art15:
            continue

        if t.startswith(
            "Todo sistema de IA institucional seguirá, en lo pertinente, el siguiente ciclo"
        ):
            deleting = True
            to_delete.append(p)
            continue

        if re.match(r"^15\.(5|6|7|9|10)\.", t):
            deleting = True
            to_delete.append(p)
            continue

        if t.startswith("15.8."):
            set_text(p, t.replace("15.8.", "15.5.", 1))
            deleting = False
            kept_registro = True
            continue

        if deleting:
            to_delete.append(p)

    for p in to_delete:
        delete_paragraph(p)
    return len(to_delete) + (1 if kept_registro else 0)


def apply_exact(doc: Document) -> int:
    n = 0
    for p in doc.paragraphs:
        style = p.style.name if p.style else ""
        if style.startswith("toc"):
            continue
        original = p.text
        text = original
        for old, new in EXACT:
            if old in text:
                text = text.replace(old, new)
        if text != original:
            set_text(p, text)
            n += 1
    return n


def fix_bibliography(doc: Document) -> int:
    n = 0
    for p in doc.paragraphs:
        t = p.text
        has_local = "file://" in t or "/Users/claudiolarrea" in t

        # Unwrap bad hyperlinks even if visible text was already cleaned
        for rel in list(p._element.xpath(".//w:hyperlink")):
            rId = rel.get(qn("r:id"))
            if not rId or rId not in p.part.rels:
                continue
            target = p.part.rels[rId].target_ref
            if target.startswith("file:") or "/Users/" in target:
                parent = rel.getparent()
                idx = parent.index(rel)
                for child in list(rel):
                    parent.insert(idx, child)
                    idx += 1
                parent.remove(rel)
                n += 1

        if not has_local:
            continue

        if "Cantero Gamito" in t:
            cleaned = (
                "Cantero Gamito, M. (s. f.). Servicios digitales e inteligencia artificial: "
                "Marco normativo de la Unión Europea para la transformación digital "
                "[Material de clase]. Doctorado en Transformación Digital / Proyecto UNI UEAR."
            )
        else:
            cleaned = re.sub(r"\s*file:///Users/\S+", "", t)
            cleaned = re.sub(r"\s*/Users/claudiolarrea/\S+", "", cleaned).strip()
            if cleaned.endswith("."):
                pass
            elif cleaned:
                cleaned += "."
        set_text(p, cleaned)
        n += 1
    return n


def fix_remaining_bad_refs(doc: Document) -> int:
    """Catch leftover out-of-range article numbers in body/annex (not headings/TOC)."""
    n = 0
    # Map common leftover v1-style refs if any remain
    extra = [
        (r"arts?\.\s*97\.4\s*y\s*97\.5", "Arts. 35° y 57°"),
        (r"arts?\.\s*97\.3\s*a\s*97\.5", "Arts. 35° y 57°"),
        (r"arts?\.\s*41\s*a\s*115", "Arts. 62° a 66°"),
        (r"artículos\s+25\s*a\s*40", "Arts. 28° a 32°"),
        (r"arts?\.\s*25\s*a\s*40", "Arts. 28° a 32°"),
    ]
    for p in doc.paragraphs:
        style = p.style.name if p.style else ""
        if style.startswith("toc") or style.startswith("Heading"):
            continue
        text = p.text
        new = text
        for pat, repl in extra:
            new = re.sub(pat, repl, new, flags=re.I)
        if new != text:
            set_text(p, new)
            n += 1
    return n


def orthography_passes(doc: Document) -> int:
    fixes = [
        ("tFI", "TFI"),
        ("  ", " "),
        ("leerce", "leerse"),  # if fallback typo introduced — fix in fallback instead
    ]
    n = 0
    for p in doc.paragraphs:
        t = p.text
        new = t
        for a, b in fixes:
            if a in new:
                new = new.replace(a, b)
        if new != t:
            set_text(p, new)
            n += 1
    return n


def main() -> None:
    if not TARGET.exists():
        raise SystemExit(f"No existe: {TARGET}")

    BACKUP_DIR.mkdir(parents=True, exist_ok=True)
    stamp = date.today().isoformat()
    backup = BACKUP_DIR / f"Marco de Gobernanza IA - UCCuyo 2026.pre-crossrefs-{stamp}.docx"
    shutil.copy2(TARGET, backup)
    print("Backup:", backup)

    doc = Document(str(TARGET))

    art40 = restore_articulo_40(doc)
    print("Artículo 40° restaurado:", art40)

    rem15 = remove_art15_cycle_duplicate(doc)
    print("Art.15 ciclo duplicado — párrafos tocados:", rem15)

    exact_n = apply_exact(doc)
    print("Reemplazos exactos (párrafos):", exact_n)

    intro_stats = fix_intros(doc)
    print("Intros:", intro_stats)

    bib_n = fix_bibliography(doc)
    print("Bibliografía / hipervínculos:", bib_n)

    extra_n = fix_remaining_bad_refs(doc)
    print("Refs residuales:", extra_n)

    ortho_n = orthography_passes(doc)
    print("Ortografía menor:", ortho_n)

    # Fix typo in fallback if any
    for p in doc.paragraphs:
        if "leerce" in p.text:
            set_text(p, p.text.replace("leerce", "leerse"))

    doc.save(str(TARGET))
    shutil.copy2(TARGET, WORKSPACE_COPY)
    print("Guardado:", TARGET)
    print("Copia workspace:", WORKSPACE_COPY)

    # Quick verify
    doc2 = Document(str(TARGET))
    texts = [p.text for p in doc2.paragraphs]
    joined = "\n".join(texts)
    problems = []
    if "arts. 41 a 115" in joined.lower() or "41 a 115" in joined:
        problems.append("still 41 a 115")
    if "97.4" in joined or "97.5" in joined:
        problems.append("still 97.x")
    if "file://" in joined or "/Users/claudiolarrea" in joined:
        problems.append("still local paths")
    if not any(t.strip().startswith("Artículo 40°") for t in texts):
        problems.append("missing Art 40")
    boilers = sum(
        1
        for t in texts
        if "orienta su interpretación práctica dentro del Marco" in t
        or "y orienta su interpretación" in t
    )
    if boilers:
        problems.append(f"boiler leftover {boilers}")
    arts = [
        int(m.group(1))
        for t in texts
        if (m := re.match(r"^Artículo\s+(\d+)\s*[°º]", t.strip()))
        and True
    ]
    # headings only roughly
    heading_arts = []
    for p in doc2.paragraphs:
        st = p.style.name if p.style else ""
        if st.startswith("Heading") and re.match(r"^Artículo\s+(\d+)", p.text.strip()):
            heading_arts.append(int(re.match(r"^Artículo\s+(\d+)", p.text.strip()).group(1)))
    missing = [i for i in range(1, 85) if i not in heading_arts]
    print("Heading articles:", len(heading_arts), "missing:", missing)
    print("VERIFY problems:", problems or "OK")


if __name__ == "__main__":
    main()
