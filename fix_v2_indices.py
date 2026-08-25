#!/usr/bin/env python3
"""Fix v2.0 Marco: manual indices, heading styles, justification."""
from docx import Document
from docx.text.paragraph import Paragraph
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re
import subprocess
import os

PATHS = [
    "/Users/claudiolarrea/Library/CloudStorage/OneDrive-Personal/16 Secretaría de Investigación/60 Observatorio de Inteligencia Artificial/MARCO DE GOBERNANZA DE LA IA/Marco de Gobernanza, Ética y Uso Responsable de la IA - UCCuyo - v2.0 depurada.docx",
    "/Users/claudiolarrea/Documents/Observatorio/Marco de Gobernanza, Ética y Uso Responsable de la IA - v2.0 depurada.docx",
]


def delete_paragraph(p):
    p._element.getparent().remove(p._element)


def insert_after(anchor, text, style_name="toc 1"):
    new_el = OxmlElement("w:p")
    anchor._p.addnext(new_el)
    np = Paragraph(new_el, anchor._parent)
    try:
        np.style = style_name
    except Exception:
        pass
    # tab leader + page number
    if "\t" in text:
        title, page = text.rsplit("\t", 1)
        r1 = np.add_run(title)
        r1.font.name = "Times New Roman"
        # tab
        r_tab = np.add_run()
        r_tab._r.append(OxmlElement("w:tab"))
        r2 = np.add_run(page)
        r2.font.name = "Times New Roman"
    else:
        r = np.add_run(text)
        r.font.name = "Times New Roman"
    return np


def set_justify(p):
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 1.5


def collect_headings(doc):
    headings = []
    in_body = False
    for p in doc.paragraphs:
        t = p.text.strip()
        if "VERSIÓN 2.0" in t:
            in_body = True
            continue
        if in_body and t == "ANEXOS":
            headings.append(("cap1", "ANEXOS"))
            break
        if not in_body:
            continue
        if t.startswith("CAPÍTULO ") and "\t" not in t:
            headings.append(("cap1", t))
        elif re.match(r"Artículo\s+\d+\s*[°.]", t) and "\t" not in t and len(t) < 120:
            headings.append(("toc2", t))

    in_annex = False
    seen_anexo = set()
    for p in doc.paragraphs:
        t = p.text.strip()
        if t == "ANEXOS":
            if in_annex:
                continue
            in_annex = True
            continue
        if in_annex and t == "BIBLIOGRAFÍA":
            headings.append(("cap1", "BIBLIOGRAFÍA"))
            break
        if in_annex and re.match(r"Anexo [A-K]\.", t) and "\t" not in t and len(t) < 120:
            if t in seen_anexo:
                continue
            seen_anexo.add(t)
            headings.append(("toc2", t))
    return headings


def map_pages(path, headings):
    pdf_path = path.replace(".docx", ".pdf")
    if os.path.exists(pdf_path):
        os.remove(pdf_path)
    subprocess.run(
        ["soffice", "--headless", "--convert-to", "pdf", "--outdir", os.path.dirname(path), path],
        capture_output=True,
    )
    page_map = {}
    try:
        from pypdf import PdfReader

        reader = PdfReader(pdf_path)
        pages_text = [
            (i + 1, (pg.extract_text() or "").replace("\n", " "))
            for i, pg in enumerate(reader.pages)
        ]

        def find_page(needle, start_pg=4):
            if not needle:
                return None
            for pg_num, text in pages_text:
                if pg_num < start_pg:
                    continue
                if needle in text:
                    return pg_num
            return None

        last_pg = 4
        current_cap = None
        for kind, title in headings:
            if kind == "cap1":
                current_cap = title
                if title in ("ANEXOS", "BIBLIOGRAFÍA"):
                    pg = find_page(title, last_pg)
                    if pg:
                        page_map[title] = pg
                        last_pg = pg
                continue
            # Artículo / Anexo: use distinctive prefix
            if title.startswith("Artículo"):
                needle = title.split(".")[0] + "."  # e.g. Artículo 16°
                if title.count(".") > 1:
                    needle = title[: min(45, len(title))]
            else:
                needle = title[:30]
            pg = find_page(needle, last_pg)
            if pg:
                page_map[title] = pg
                last_pg = pg
                if current_cap and current_cap not in page_map and current_cap.startswith("CAPÍTULO"):
                    page_map[current_cap] = pg

        # fill chapter pages from first child
        caps = [t for k, t in headings if k == "cap1" and t.startswith("CAPÍTULO")]
        arts = [t for k, t in headings if k == "toc2" and t.startswith("Artículo")]
        ai = 0
        for cap in caps:
            if cap in page_map:
                if ai < len(arts):
                    while ai < len(arts) and page_map.get(arts[ai], 999) < page_map[cap]:
                        ai += 1
                continue
            while ai < len(arts) and arts[ai] in page_map:
                page_map[cap] = page_map[arts[ai]]
                break
            ai += 1
    except Exception as exc:
        print("PDF mapping error:", exc)
    return page_map


def update_toc_pages(doc, page_map):
    for p in doc.paragraphs:
        if p.style.name not in ("toc 1", "toc 2"):
            continue
        t = p.text.strip()
        if not t:
            continue
        title = t.split("\t")[0].strip()
        if title in page_map:
            new_text = f"{title}\t{page_map[title]}"
            if p.runs:
                p.runs[0].text = new_text
                for r in p.runs[1:]:
                    r.text = ""
            else:
                p.text = new_text


def style_body(doc):
    in_body = in_annex = False
    for p in doc.paragraphs:
        t = p.text.strip()
        if "VERSIÓN 2.0" in t:
            in_body = True
            set_justify(p)
            continue
        if in_body and t == "ANEXOS":
            in_body = False
            in_annex = True
            continue
        if in_annex and t == "BIBLIOGRAFÍA":
            break
        if in_body:
            if t.startswith("CAPÍTULO ") and "\t" not in t:
                p.style = "Heading 1"
                set_justify(p)
            elif re.match(r"Artículo\s+\d+\s*[°.]", t) and len(t) < 120:
                p.style = "Heading 2"
                set_justify(p)
            elif t:
                set_justify(p)
        elif in_annex and t:
            set_justify(p)


def rebuild_indices(doc, page_map, headings):
    res_hdr = body_start = None
    for i, p in enumerate(doc.paragraphs):
        t = p.text.strip()
        if t == "ÍNDICE RESUMIDO":
            res_hdr = i
        if "VERSIÓN 2.0" in t:
            body_start = i
            break

    for idx in range(body_start - 1, res_hdr, -1):
        delete_paragraph(doc.paragraphs[idx])

    for i, p in enumerate(doc.paragraphs):
        if p.text.strip() == "ÍNDICE RESUMIDO":
            cur = p
            break

    body_caps = [t for k, t in headings if k == "cap1" and t not in ("ANEXOS", "BIBLIOGRAFÍA")]
    # body_caps from headings may include ANEXOS - filter to only CAPÍTULO
    body_caps = [t for t in body_caps if t.startswith("CAPÍTULO")]

    for cap in body_caps:
        cur = insert_after(cur, f"{cap}\t ", "toc 1")
    cur = insert_after(cur, f"ANEXOS\t ", "toc 1")
    cur = insert_after(cur, "")

    gen_el = OxmlElement("w:p")
    cur._p.addnext(gen_el)
    gen_p = Paragraph(gen_el, cur._parent)
    gen_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rg = gen_p.add_run("ÍNDICE GENERAL")
    rg.bold = True
    rg.font.name = "Times New Roman"
    cur = gen_p

    cur = insert_after(
        cur,
        "Índice detallado del documento (capítulos y artículos). Al hacer clic en un apartado se desplaza a la sección correspondiente.",
        "Normal",
    )
    cur = insert_after(cur, "")

    for kind, title in headings:
        style = "toc 1" if kind == "cap1" else "toc 2"
        cur = insert_after(cur, f"{title}\t ", style)
    insert_after(cur, "")


def fix(path):
    doc = Document(path)
    headings = collect_headings(doc)
    style_body(doc)
    rebuild_indices(doc, {}, headings)  # placeholder pages first
    style_body(doc)
    doc.save(path)
    page_map = map_pages(path, headings)
    doc = Document(path)
    update_toc_pages(doc, page_map)
    style_body(doc)
    doc.save(path)
    print(f"Fixed {path} — {len(headings)} entries, {len(page_map)} pages mapped")


if __name__ == "__main__":
    for p in PATHS:
        if os.path.exists(p):
            fix(p)
