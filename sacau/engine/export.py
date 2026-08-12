"""Exportación CSV / Excel / Word / PDF del plan convertido."""

from __future__ import annotations

from io import BytesIO

import pandas as pd

from .convert import totales_por_anio, totales_por_area
from .models import PlanConvertido
from .validate import ValidationResult


DETAIL_COLUMNS = [
    "codigo",
    "nombre",
    "anio",
    "area",
    "regimen",
    "tipologia",
    "horas_teoricas",
    "horas_practicas",
    "horas_interaccion",
    "horas_autonomas",
    "autonomas_fuente",
    "horas_totales",
    "valor_cre",
    "cre",
    "horas_estimadas",
    "notas",
]


def plan_to_dataframe(plan_conv: PlanConvertido) -> pd.DataFrame:
    rows = []
    for item in plan_conv.items:
        a = item.asignatura
        rows.append(
            {
                "codigo": a.codigo,
                "nombre": a.nombre,
                "anio": a.anio,
                "area": a.area,
                "regimen": a.regimen,
                "tipologia": a.tipologia,
                "horas_teoricas": a.horas_teoricas,
                "horas_practicas": a.horas_practicas,
                "horas_interaccion": item.horas_interaccion,
                "horas_autonomas": item.horas_autonomas,
                "autonomas_fuente": item.autonomas_fuente,
                "horas_totales": item.horas_totales,
                "valor_cre": item.valor_cre,
                "cre": item.cre,
                "horas_estimadas": a.horas_estimadas,
                "notas": a.notas,
            }
        )
    return pd.DataFrame(rows, columns=DETAIL_COLUMNS)


def resumen_anios_df(plan_conv: PlanConvertido) -> pd.DataFrame:
    rows = []
    for anio, tot in totales_por_anio(plan_conv.items).items():
        rows.append(
            {
                "anio": anio,
                "horas_interaccion": tot.horas_interaccion,
                "horas_autonomas": tot.horas_autonomas,
                "horas_totales": tot.horas_totales,
                "cre": tot.cre,
            }
        )
    return pd.DataFrame(rows)


def resumen_areas_df(plan_conv: PlanConvertido) -> pd.DataFrame:
    rows = []
    for area, tot in totales_por_area(plan_conv.items).items():
        rows.append(
            {
                "area": area,
                "horas_interaccion": tot.horas_interaccion,
                "horas_practicas": tot.horas_practicas,
                "horas_autonomas": tot.horas_autonomas,
                "horas_totales": tot.horas_totales,
                "cre": tot.cre,
            }
        )
    return pd.DataFrame(rows)


def validacion_df(resultado: ValidationResult) -> pd.DataFrame:
    return pd.DataFrame([c.__dict__ for c in resultado.checks])


def to_csv_bytes(plan_conv: PlanConvertido) -> bytes:
    return plan_to_dataframe(plan_conv).to_csv(index=False).encode("utf-8-sig")


def to_excel_bytes(
    plan_conv: PlanConvertido,
    validacion: ValidationResult | None = None,
) -> bytes:
    buffer = BytesIO()
    with pd.ExcelWriter(buffer, engine="openpyxl") as writer:
        plan_to_dataframe(plan_conv).to_excel(writer, sheet_name="Asignaturas", index=False)
        resumen_anios_df(plan_conv).to_excel(writer, sheet_name="Por_anio", index=False)
        resumen_areas_df(plan_conv).to_excel(writer, sheet_name="Por_area", index=False)
        totales = plan_conv.totales
        pd.DataFrame(
            [
                {
                    "plan": plan_conv.plan.nombre,
                    "horas_teoricas": totales.horas_teoricas,
                    "horas_practicas": totales.horas_practicas,
                    "horas_interaccion": totales.horas_interaccion,
                    "horas_autonomas": totales.horas_autonomas,
                    "horas_totales": totales.horas_totales,
                    "cre": totales.cre,
                    "anios": totales.anios,
                    "cre_promedio_anual": totales.cre_promedio_anual,
                    "valor_cre_default": plan_conv.opciones.valor_cre_default,
                }
            ]
        ).to_excel(writer, sheet_name="Totales", index=False)
        if validacion is not None:
            validacion_df(validacion).to_excel(writer, sheet_name="Validacion", index=False)
    return buffer.getvalue()


def to_docx_bytes(
    plan_conv: PlanConvertido,
    validacion: ValidationResult | None = None,
) -> bytes:
    from docx import Document
    from docx.shared import Pt

    doc = Document()
    doc.add_heading("Plan de estudios en créditos CRE (SACAU)", level=1)
    doc.add_paragraph(plan_conv.plan.nombre or "Plan de estudios")
    t = plan_conv.totales
    doc.add_paragraph(
        f"{plan_conv.plan.institucion} · CRE default {plan_conv.opciones.valor_cre_default} h"
    )
    doc.add_heading("Totales", level=2)
    doc.add_paragraph(
        f"Interacción: {t.horas_interaccion:.0f} h · Autónomas: {t.horas_autonomas:.0f} h · "
        f"Totales: {t.horas_totales:.0f} h · CRE: {t.cre:.1f} · CRE/año: {t.cre_promedio_anual:.1f}"
    )
    doc.add_heading("Asignaturas", level=2)
    table = doc.add_table(rows=1, cols=8)
    hdr = table.rows[0].cells
    headers = ["Cód.", "Asignatura", "Año", "Área", "Inter.", "Autón.", "Total", "CRE"]
    for i, h in enumerate(headers):
        hdr[i].text = h
    for item in plan_conv.items:
        a = item.asignatura
        row = table.add_row().cells
        row[0].text = str(a.codigo)
        row[1].text = str(a.nombre)
        row[2].text = str(a.anio)
        row[3].text = str(a.area)
        row[4].text = f"{item.horas_interaccion:.0f}"
        row[5].text = f"{item.horas_autonomas:.0f}"
        row[6].text = f"{item.horas_totales:.0f}"
        row[7].text = f"{item.cre:.1f}"

    if validacion is not None:
        doc.add_heading("Cumplimiento SACAU", level=2)
        for c in validacion.checks:
            mark = {"ok": "[OK]", "warning": "[AVISO]", "error": "[REVISAR]"}.get(c.nivel, "[·]")
            p = doc.add_paragraph(f"{mark} {c.mensaje}")
            for run in p.runs:
                run.font.size = Pt(10)

    anexo = (plan_conv.plan.metadata or {}).get("anexo_911")
    if anexo and (plan_conv.plan.metadata or {}).get("incluir_anexo_911", True):
        labels = {
            "perfil_egreso": "Perfil de egreso",
            "competencias_genericas": "Competencias genéricas (transversales)",
            "competencias_especificas": "Competencias específicas",
            "resultados_aprendizaje": "Resultados de aprendizaje",
            "despliegue_horas": "Despliegue de horas y créditos",
            "flexibilidad_curricular": "Flexibilidad curricular",
            "reconocimiento_trayectos": "Reconocimiento de trayectos formativos",
            "movilidad": "Movilidad estudiantil e interinstitucional",
            "matriz_tributacion": "Matriz de tributación (orientación)",
            "notas_unidad_academica": "Notas de la unidad académica",
        }
        doc.add_heading("Anexo curricular (Res. 911-CS-2026 UCCuyo)", level=2)
        doc.add_paragraph(
            "Borrador editable generado por el Convertidor SACAU. Adaptar a la especificidad de la carrera."
        )
        for key, label in labels.items():
            text = anexo.get(key)
            if not text or not str(text).strip():
                continue
            doc.add_heading(label, level=3)
            for line in str(text).splitlines():
                if line.strip():
                    doc.add_paragraph(line.strip())

    doc.add_paragraph(
        "Nota: las horas autónomas son estimaciones editables; no son objeto de verificación "
        "en validez nacional (RESOL-2025-556)."
    )
    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def to_pdf_bytes(
    plan_conv: PlanConvertido,
    validacion: ValidationResult | None = None,
) -> bytes:
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import A4, landscape
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle

    buf = BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=landscape(A4), leftMargin=36, rightMargin=36, topMargin=36, bottomMargin=36)
    styles = getSampleStyleSheet()
    story = []
    t = plan_conv.totales
    story.append(Paragraph("Plan de estudios en créditos CRE (SACAU)", styles["Title"]))
    story.append(Paragraph(plan_conv.plan.nombre or "Plan de estudios", styles["Heading2"]))
    story.append(
        Paragraph(
            f"{plan_conv.plan.institucion} · CRE default {plan_conv.opciones.valor_cre_default} h · "
            f"Interacción {t.horas_interaccion:.0f} h · Autónomas {t.horas_autonomas:.0f} h · "
            f"Totales {t.horas_totales:.0f} h · CRE {t.cre:.1f} · CRE/año {t.cre_promedio_anual:.1f}",
            styles["Normal"],
        )
    )
    story.append(Spacer(1, 12))

    data = [["Cód.", "Asignatura", "Año", "Área", "Inter.", "Autón.", "Total", "CRE"]]
    for item in plan_conv.items:
        a = item.asignatura
        data.append(
            [
                str(a.codigo),
                str(a.nombre)[:60],
                str(a.anio),
                str(a.area),
                f"{item.horas_interaccion:.0f}",
                f"{item.horas_autonomas:.0f}",
                f"{item.horas_totales:.0f}",
                f"{item.cre:.1f}",
            ]
        )
    table = Table(data, repeatRows=1)
    table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), colors.Color(122 / 255, 21 / 255, 50 / 255)),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                ("FONTSIZE", (0, 0), (-1, -1), 8),
                ("GRID", (0, 0), (-1, -1), 0.25, colors.grey),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ]
        )
    )
    story.append(table)

    if validacion is not None:
        story.append(Spacer(1, 14))
        story.append(Paragraph("Cumplimiento SACAU", styles["Heading2"]))
        for c in validacion.checks:
            mark = {"ok": "[OK]", "warning": "[AVISO]", "error": "[REVISAR]"}.get(c.nivel, "[·]")
            story.append(Paragraph(f"{mark} {c.mensaje}", styles["Normal"]))

    doc.build(story)
    return buf.getvalue()
