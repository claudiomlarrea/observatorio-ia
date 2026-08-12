"""
Convertidor SACAU — UCCuyo
Carga un plan en horas (Word/PDF/CSV) y lo transforma a créditos CRE.
"""

from __future__ import annotations

import json
import re
import sys
from io import BytesIO
from pathlib import Path

import pandas as pd
import streamlit as st

ROOT = Path(__file__).resolve().parent
if str(ROOT) not in sys.path:
    sys.path.insert(0, str(ROOT))

from engine.convert import convert_plan, totales_por_anio, totales_por_area
from engine.export import (
    plan_to_dataframe,
    to_csv_bytes,
    to_docx_bytes,
    to_excel_bytes,
    to_pdf_bytes,
)
from engine.models import Asignatura, ConvertOptions, PlanEstudios, Tipologia
from engine.validate import validate_plan

DATA = ROOT / "data"

st.set_page_config(
    page_title="Convertidor SACAU | UCCuyo",
    page_icon="🎓",
    layout="wide",
)


@st.cache_data
def load_json(name: str) -> dict:
    return json.loads((DATA / name).read_text(encoding="utf-8"))


def load_tipologias() -> dict[str, Tipologia]:
    raw = load_json("tipologias.json")
    return {t["id"]: Tipologia.from_dict(t) for t in raw["tipologias"]}


def tipologias_to_editor_df(tips: dict[str, Tipologia]) -> pd.DataFrame:
    return pd.DataFrame(
        [
            {
                "id": t.id,
                "nombre": t.nombre,
                "ratio_autonomo": t.ratio_autonomo,
                "autonomas_fijas": t.autonomas_fijas,
                "descripcion": t.descripcion,
            }
            for t in tips.values()
        ]
    )


def df_to_tipologias(df: pd.DataFrame) -> dict[str, Tipologia]:
    out: dict[str, Tipologia] = {}
    for _, row in df.iterrows():
        tid = str(row["id"]).strip()
        if not tid:
            continue
        out[tid] = Tipologia(
            id=tid,
            nombre=str(row.get("nombre", tid)),
            ratio_autonomo=float(row.get("ratio_autonomo", 1.0) or 0),
            autonomas_fijas=float(row.get("autonomas_fijas", 0) or 0),
            descripcion=str(row.get("descripcion", "") or ""),
        )
    return out


EDITOR_COLUMNS = [
    "codigo",
    "nombre",
    "anio",
    "area",
    "regimen",
    "tipologia",
    "horas_teoricas",
    "horas_practicas",
    "horas_autonomas_override",
    "valor_cre_override",
    "horas_estimadas",
    "notas",
]


def plan_to_editor_df(plan: PlanEstudios) -> pd.DataFrame:
    rows = []
    for a in plan.asignaturas:
        rows.append(
            {
                "codigo": a.codigo,
                "nombre": a.nombre,
                "anio": a.anio,
                "area": a.area,
                "regimen": a.regimen,
                "tipologia": a.tipologia,
                "horas_teoricas": a.horas_teoricas,
                "horas_practicas": a.horas_practicas,
                "horas_autonomas_override": a.horas_autonomas_override,
                "valor_cre_override": a.valor_cre_override,
                "horas_estimadas": a.horas_estimadas,
                "notas": a.notas,
            }
        )
    return pd.DataFrame(rows, columns=EDITOR_COLUMNS)


def editor_df_to_plan(df: pd.DataFrame, base: PlanEstudios) -> PlanEstudios:
    asignaturas: list[Asignatura] = []
    for _, row in df.iterrows():
        nombre = str(row.get("nombre", "") or "").strip()
        if not nombre:
            continue
        override_auto = row.get("horas_autonomas_override")
        override_cre = row.get("valor_cre_override")
        asignaturas.append(
            Asignatura(
                codigo=str(row.get("codigo", "") or ""),
                nombre=nombre,
                anio=int(row.get("anio", 1) or 1),
                area=str(row.get("area", "") or ""),
                horas_teoricas=float(row.get("horas_teoricas", 0) or 0),
                horas_practicas=float(row.get("horas_practicas", 0) or 0),
                regimen=str(row.get("regimen", "S") or "S"),
                tipologia=str(row.get("tipologia", "teorica") or "teorica"),
                horas_autonomas_override=(
                    None
                    if override_auto is None or (isinstance(override_auto, float) and pd.isna(override_auto))
                    else float(override_auto)
                ),
                valor_cre_override=(
                    None
                    if override_cre is None or (isinstance(override_cre, float) and pd.isna(override_cre))
                    else float(override_cre)
                ),
                horas_estimadas=bool(row.get("horas_estimadas", False)),
                notas=str(row.get("notas", "") or ""),
            )
        )
    duracion = base.duracion_anios
    if asignaturas:
        duracion = max(a.anio for a in asignaturas)
    return PlanEstudios(
        id=base.id,
        nombre=base.nombre,
        titulo=base.titulo,
        institucion=base.institucion,
        normativa=base.normativa,
        duracion_anios=duracion,
        carrera_clave="",
        asignaturas=asignaturas,
        metadata=base.metadata,
    )


def empty_plan() -> PlanEstudios:
    return PlanEstudios(
        id="plan-nuevo",
        nombre="Plan de estudios (nuevo)",
        institucion="Universidad Católica de Cuyo",
        asignaturas=[
            Asignatura(
                codigo="01",
                nombre="Asignatura ejemplo",
                anio=1,
                area="FB",
                horas_teoricas=64,
                horas_practicas=0,
                regimen="S",
                tipologia="teorica",
            )
        ],
    )


def guess_tipologia(nombre: str) -> str:
    n = nombre.lower()
    if "práctica profesional" in n or "pps" in n:
        return "pps"
    if "trabajo integrador" in n or "t.i.f" in n or "tesis" in n:
        return "tif"
    if "taller" in n or "seminario" in n:
        return "taller"
    if "optativa" in n:
        return "optativa"
    if "práctica" in n or "exploración" in n:
        return "practica_supervisada"
    return "teorica"


def parse_text_to_asignaturas(text: str) -> list[Asignatura]:
    asignaturas: list[Asignatura] = []
    anio = 1
    for raw in text.splitlines():
        line = re.sub(r"\s+", " ", raw).strip()
        if not line:
            continue
        m_anio = re.search(r"(\d{1,2})\s*[°ºo.]?\s*a[nñ]o\b", line, re.I)
        if m_anio:
            anio = int(m_anio.group(1))
            continue
        m = re.match(
            r"^(\d{1,3})[).\-\s]+(.+?)\s+(\d{1,4})\s+(\d{1,4})\s*$",
            line,
        )
        if m:
            asignaturas.append(
                Asignatura(
                    codigo=m.group(1),
                    nombre=m.group(2).strip(),
                    anio=anio,
                    area="OTRA",
                    horas_teoricas=float(m.group(3)),
                    horas_practicas=float(m.group(4)),
                    tipologia=guess_tipologia(m.group(2)),
                )
            )
            continue
        m = re.match(r"^(.{8,120}?)\s+(\d{2,4})\s+(\d{1,4})\s*$", line)
        if m and not re.match(r"^(total|suma|carga)", m.group(1), re.I):
            asignaturas.append(
                Asignatura(
                    codigo=str(len(asignaturas) + 1).zfill(2),
                    nombre=m.group(1).strip(),
                    anio=anio,
                    area="OTRA",
                    horas_teoricas=float(m.group(2)),
                    horas_practicas=float(m.group(3)),
                    tipologia=guess_tipologia(m.group(1)),
                )
            )
            continue
        m = re.match(r"^(\d{1,3})[).\-\s]+(.{8,120}?)\s+(\d{2,4})\s*$", line)
        if m:
            asignaturas.append(
                Asignatura(
                    codigo=m.group(1),
                    nombre=m.group(2).strip(),
                    anio=anio,
                    area="OTRA",
                    horas_teoricas=float(m.group(3)),
                    horas_practicas=0,
                    tipologia=guess_tipologia(m.group(2)),
                )
            )
    return asignaturas


def load_uploaded_plan(uploaded) -> PlanEstudios:
    name = uploaded.name
    lower = name.lower()
    data = uploaded.read()
    base = PlanEstudios(
        id=f"plan-{Path(name).stem}",
        nombre=f"Plan cargado ({name})",
        institucion="Universidad Católica de Cuyo",
        metadata={"fuente": name},
    )

    if lower.endswith(".csv"):
        df = pd.read_csv(BytesIO(data))
        return import_csv_to_plan(df, base)

    if lower.endswith(".docx"):
        from docx import Document

        doc = Document(BytesIO(data))
        text = "\n".join(p.text for p in doc.paragraphs)
        for table in doc.tables:
            for row in table.rows:
                text += "\n" + " ".join(c.text.strip() for c in row.cells)
        asignaturas = parse_text_to_asignaturas(text)
        base.asignaturas = asignaturas or empty_plan().asignaturas
        base.metadata["advertencia"] = (
            "Revisá las asignaturas detectadas desde Word y completá tipologías/horas si hace falta."
        )
        return base

    if lower.endswith(".pdf"):
        from pypdf import PdfReader

        reader = PdfReader(BytesIO(data))
        text = "\n".join((page.extract_text() or "") for page in reader.pages)
        asignaturas = parse_text_to_asignaturas(text)
        base.asignaturas = asignaturas or empty_plan().asignaturas
        base.metadata["advertencia"] = (
            "Revisá el plan detectado desde PDF. Los PDF escaneados pueden requerir carga manual o CSV."
        )
        return base

    if lower.endswith(".doc"):
        raise ValueError("Los .doc antiguos no están soportados. Guardá como .docx o PDF.")

    raise ValueError("Formato no soportado. Usá PDF, Word (.docx) o CSV.")


def import_csv_to_plan(df: pd.DataFrame, base: PlanEstudios) -> PlanEstudios:
    rename = {
        "código": "codigo",
        "cod": "codigo",
        "asignatura": "nombre",
        "materia": "nombre",
        "año": "anio",
        "teo": "horas_teoricas",
        "practicas": "horas_practicas",
        "prácticas": "horas_practicas",
        "h_teoricas": "horas_teoricas",
        "h_practicas": "horas_practicas",
    }
    cols = {c: rename.get(str(c).strip().lower(), str(c).strip().lower()) for c in df.columns}
    df = df.rename(columns=cols)
    for col in EDITOR_COLUMNS:
        if col not in df.columns:
            df[col] = None if "override" in col else (False if col == "horas_estimadas" else "")
    return editor_df_to_plan(df[EDITOR_COLUMNS], base)


def nivel_emoji(nivel: str) -> str:
    return {"ok": "✅", "warning": "⚠️", "error": "❌"}.get(nivel, "•")


def main() -> None:
    normas_uccuyo = load_json("normas_uccuyo.json")
    tips_default = load_tipologias()

    st.title("Convertidor SACAU → CRE")
    st.caption(
        "Universidad Católica de Cuyo · Res. 788-CS-2026 (CRE = 25 h, hasta 30) · "
        "Cargá un plan en horas (Word/PDF/CSV) y descargá el plan en créditos."
    )

    with st.sidebar:
        st.header("1. Cargar plan")
        uploaded = st.file_uploader(
            "Word (.docx), PDF o CSV",
            type=["docx", "pdf", "csv"],
        )
        if st.button("Usar plan en blanco"):
            st.session_state.plan_base = empty_plan()
            st.session_state.editor_df = plan_to_editor_df(st.session_state.plan_base)
            st.session_state.tips_df = tipologias_to_editor_df(tips_default)
            st.session_state.plan_key = "blank"

        st.divider()
        st.header("2. Parámetros CRE")
        valor_cre = st.number_input(
            "Valor CRE por defecto (horas)",
            min_value=float(normas_uccuyo["sacau"]["cre_rango_min"]),
            max_value=float(normas_uccuyo["sacau"]["cre_rango_max"]),
            value=float(normas_uccuyo.get("cre_default", 25)),
            step=1.0,
        )
        redondeo = st.selectbox(
            "Redondeo de CRE",
            options=[0.5, 0.25, 0.0],
            format_func=lambda x: "Sin redondeo" if x == 0 else f"Múltiplos de {x}",
            index=0,
        )

    if "plan_base" not in st.session_state:
        st.session_state.plan_base = empty_plan()
        st.session_state.editor_df = plan_to_editor_df(st.session_state.plan_base)
        st.session_state.tips_df = tipologias_to_editor_df(tips_default)
        st.session_state.plan_key = "blank"

    if uploaded is not None:
        key = f"{uploaded.name}:{uploaded.size}"
        if st.session_state.get("plan_key") != key:
            try:
                loaded = load_uploaded_plan(uploaded)
                st.session_state.plan_base = loaded
                st.session_state.editor_df = plan_to_editor_df(loaded)
                st.session_state.tips_df = tipologias_to_editor_df(tips_default)
                st.session_state.plan_key = key
                st.success(f"Plan cargado: {len(loaded.asignaturas)} asignaturas detectadas. Revisá y ajustá.")
            except Exception as exc:  # noqa: BLE001
                st.error(str(exc))

    base: PlanEstudios = st.session_state.plan_base
    if base.metadata.get("advertencia"):
        st.info(base.metadata["advertencia"])

    tab_params, tab_plan, tab_result = st.tabs(
        ["Coeficientes autónomos", "Editor del plan", "Resultados y descarga"]
    )

    with tab_params:
        st.write(
            "Las horas autónomas se estiman como `interacción × ratio + horas fijas`, "
            "salvo override por asignatura."
        )
        tips_edited = st.data_editor(
            st.session_state.tips_df,
            num_rows="dynamic",
            use_container_width=True,
            key="tips_editor",
        )
        st.session_state.tips_df = tips_edited

    with tab_plan:
        st.subheader(base.nombre)
        edited = st.data_editor(
            st.session_state.editor_df,
            num_rows="dynamic",
            use_container_width=True,
            key="plan_editor",
            column_config={
                "anio": st.column_config.NumberColumn("Año", min_value=1, max_value=12, step=1),
                "horas_teoricas": st.column_config.NumberColumn("H. teóricas", min_value=0),
                "horas_practicas": st.column_config.NumberColumn("H. prácticas", min_value=0),
                "horas_autonomas_override": st.column_config.NumberColumn(
                    "Autónomas (override)", min_value=0
                ),
                "valor_cre_override": st.column_config.NumberColumn(
                    "CRE h (override)", min_value=25, max_value=30
                ),
                "tipologia": st.column_config.SelectboxColumn(
                    "Tipología", options=list(tips_default.keys())
                ),
                "area": st.column_config.SelectboxColumn(
                    "Área", options=["FB", "FP", "FGC", "FCI", "OTRA"]
                ),
                "regimen": st.column_config.SelectboxColumn("Régimen", options=["A", "S"]),
            },
        )
        st.session_state.editor_df = edited

    tipologias = df_to_tipologias(st.session_state.tips_df)
    plan = editor_df_to_plan(st.session_state.editor_df, base)
    options = ConvertOptions(
        valor_cre_default=float(valor_cre),
        redondeo_cre=float(redondeo),
        tipologias=tipologias,
    )
    convertido = convert_plan(plan, options)
    validacion = validate_plan(convertido, normas_uccuyo, None)

    with tab_result:
        t = convertido.totales
        m1, m2, m3, m4, m5 = st.columns(5)
        m1.metric("Interacción", f"{t.horas_interaccion:,.0f} h")
        m2.metric("Autónomas", f"{t.horas_autonomas:,.0f} h")
        m3.metric("Totales estudiante", f"{t.horas_totales:,.0f} h")
        m4.metric("CRE totales", f"{t.cre:,.1f}")
        m5.metric("CRE / año", f"{t.cre_promedio_anual:,.1f}")

        st.subheader("Cumplimiento SACAU")
        for check in validacion.checks:
            st.markdown(f"{nivel_emoji(check.nivel)} **{check.mensaje}**")

        st.subheader("Detalle por asignatura")
        st.dataframe(plan_to_dataframe(convertido), use_container_width=True, hide_index=True)

        c_a, c_b = st.columns(2)
        with c_a:
            st.markdown("**Por año**")
            st.dataframe(
                pd.DataFrame(
                    [
                        {
                            "Año": anio,
                            "Interacción": tot.horas_interaccion,
                            "Autónomas": tot.horas_autonomas,
                            "CRE": tot.cre,
                        }
                        for anio, tot in totales_por_anio(convertido.items).items()
                    ]
                ),
                hide_index=True,
                use_container_width=True,
            )
        with c_b:
            st.markdown("**Por área**")
            st.dataframe(
                pd.DataFrame(
                    [
                        {
                            "Área": area,
                            "Interacción": tot.horas_interaccion,
                            "Prácticas": tot.horas_practicas,
                            "CRE": tot.cre,
                        }
                        for area, tot in totales_por_area(convertido.items).items()
                    ]
                ),
                hide_index=True,
                use_container_width=True,
            )

        st.subheader("Descargar plan en créditos")
        x1, x2, x3, x4 = st.columns(4)
        with x1:
            st.download_button(
                "Word (.docx)",
                data=to_docx_bytes(convertido, validacion),
                file_name=f"{plan.id}_CRE.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        with x2:
            st.download_button(
                "PDF",
                data=to_pdf_bytes(convertido, validacion),
                file_name=f"{plan.id}_CRE.pdf",
                mime="application/pdf",
            )
        with x3:
            st.download_button(
                "Excel",
                data=to_excel_bytes(convertido, validacion),
                file_name=f"{plan.id}_CRE.xlsx",
                mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            )
        with x4:
            st.download_button(
                "CSV",
                data=to_csv_bytes(convertido),
                file_name=f"{plan.id}_CRE.csv",
                mime="text/csv",
            )


if __name__ == "__main__":
    main()
