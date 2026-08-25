#!/usr/bin/env python3
"""Replace formulaic intros with unique 60–80 word paragraphs in Marco 2026."""
from __future__ import annotations

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

TARGET = (
    "/Users/claudiolarrea/Library/CloudStorage/OneDrive-Personal/"
    "16 Secretaría de Investigación/60 Observatorio de Inteligencia Artificial/"
    "MARCO DE GOBERNANZA DE LA IA/Documento FINAL/"
    "Marco de Gobernanza IA - UCCuyo 2026.docx"
)

INTROS = {
    "3.2. Funciones sustantivas y de gestión": (
        "El Marco no se circunscribe a un área aislada: alcanza las funciones que constituyen "
        "la vida universitaria y también los procesos de gestión que las hacen posible. "
        "Docencia, investigación, extensión y administración se articulan en un mismo proyecto "
        "institucional, de modo que el uso de IA en cualquiera de ellas debe respetar la misión, "
        "la integridad académica y la protección de las personas. Las enumeraciones que siguen "
        "identifican esos ámbitos de aplicación, sin pretender un catálogo cerrado."
    ),
    "3.3. Sujetos alcanzados": (
        "La norma vincula a quienes, por su rol o por su vínculo con la Universidad, pueden "
        "utilizar IA o resultar afectados por ella. No basta con regular a docentes y estudiantes: "
        "también quedan comprendidos autoridades, personal de gestión, investigadores, terceros "
        "contratados y toda persona que intervenga en actividades universitarias con sistemas "
        "algorítmicos. El alcance personal busca evitar vacíos de responsabilidad y garantizar "
        "que las mismas salvaguardas rigen, en lo pertinente, para todos los destinatarios."
    ),
    "3.4. Tecnologías comprendidas": (
        "El Marco no se limita a la inteligencia artificial generativa de uso cotidiano. "
        "Comprende también sistemas de recomendación, clasificación, predicción, analítica de "
        "aprendizaje y cualquier dispositivo algorítmico que intervenga, total o parcialmente, "
        "en decisiones con efectos sobre personas o procesos institucionales. Esa amplitud "
        "evita que una herramienta escape a la gobernanza por no llamarse “chatbot” o por "
        "estar embebida en una plataforma ya contratada. Lo decisivo es la función que cumple."
    ),
    "4.1. Normativa e instrumentos institucionales": (
        "Este Marco se interpreta en continuidad con el ordenamiento propio de la Universidad, "
        "no como un cuerpo ajeno o paralelo. Estatutos, plan estratégico, ordenanzas de docencia "
        "e investigación y la resolución de creación del Observatorio conforman el piso normativo "
        "sobre el cual se apoyan las reglas de IA. Las viñetas recogen esos instrumentos para "
        "que ninguna disposición se lea desconectada de la identidad institucional, de los fines "
        "universitarios ni de las competencias ya asignadas a los órganos de gobierno."
    ),
    "4.2. Referentes externos de orientación": (
        "Sin perjuicio de la autonomía universitaria, el Marco se nutre de referentes éticos y "
        "de gobernanza que iluminan su sentido, sin incorporar automáticamente normas extranjeras. "
        "UNESCO, lineamientos nacionales de IA fiable, magisterio de la Iglesia y principios de "
        "confianza tecnológica operan como orientación comparada. Sirven para interpretar, no "
        "para sustituir, el criterio propio de la Universidad Católica de Cuyo ni para alterar "
        "la jerarquía de sus normas internas o su carácter confesional."
    ),
    "4.4. Ámbito internacional": (
        "En el plano internacional se recogen recomendaciones y principios que informan una IA "
        "digna de confianza, transparente y centrada en la persona. Su recepción es selectiva y "
        "compatible con la identidad de la Universidad: no hay incorporación automática del "
        "derecho europeo ni de otros ordenamientos. Las referencias que siguen —UNESCO, "
        "directrices de confianza, aportes comparados— deben leerse como brújula ética y de "
        "buena práctica, no como fuente de obligaciones extrañas al Estatuto institucional."
    ),
    "4.5. Ámbito nacional": (
        "En el ámbito argentino el Marco se alinea con las recomendaciones de IA fiable y con "
        "la normativa vigente en datos personales, educación superior, propiedad intelectual e "
        "integridad académica, en lo que corresponda. Esa articulación evita contradicciones "
        "con el derecho interno y da anclaje a las obligaciones de docentes, investigadores y "
        "gestores. Las viñetas no reproducen esas leyes: señalan el horizonte normativo con el "
        "cual este texto debe interpretarse de modo armónico y complementario."
    ),
    "4.6. Ámbito eclesial y de universidades católicas": (
        "Como universidad católica, la institución lee la IA a la luz del magisterio sobre la "
        "persona, el bien común, la educación y el desarrollo tecnológico. La llamada “algor-ética” "
        "y el rol de las universidades católicas como espacios de discernimiento no son un "
        "aditamento retórico: orientan la interpretación de límites, prudencia y servicio. "
        "Los referentes eclesiales no sustituyen la normativa universitaria, pero sí impiden "
        "una lectura meramente tecnocrática de este Marco."
    ),
    "Artículo 5°. Principios de interpretación": (
        "Toda disposición del Marco debe leerse a la luz de criterios que no se agotan en una "
        "lista de buenos deseos. La centralidad de la persona, la primacía de la responsabilidad "
        "humana, la proporcionalidad, la transparencia, el enfoque preventivo y la coherencia "
        "con el ideario católico son claves de lectura para resolver dudas, tensiones y lagunas. "
        "Cuando un uso de IA parezca eficiente pero erosione dignidad, equidad o integridad "
        "académica, estos principios prevalecen sobre la mera conveniencia técnica o administrativa."
    ),
    "Artículo 9°. Principio rector de no sustitución del juicio humano": (
        "En la Universidad la IA asiste, pero no reemplaza el juicio docente, científico, ético "
        "ni institucional en las decisiones que afectan de modo relevante a las personas. Ese "
        "principio rector recorre todo el Marco: la acreditación, la autoría, las sanciones y "
        "las decisiones de impacto no pueden quedar libradas a una salida algorítmica. Las "
        "viñetas precisan ámbitos en los que la supervisión humana es ineludible, para que "
        "ningún sistema se invoque como sustituto del discernimiento responsable."
    ),
    "Artículo 10°. Finalidad formativa de la gobernanza de la IA": (
        "Gobernar la IA no se reduce a controlar o sancionar: es también una tarea educativa. "
        "La Universidad busca formar para el uso crítico, ético y creativo de la tecnología, "
        "fortalecer el discernimiento frente a recomendaciones opacas y evitar que la eficiencia "
        "desplace el vínculo pedagógico. Las enumeraciones que siguen traducen esa finalidad "
        "formativa en objetivos concretos, de modo que la gobernanza se entienda como cultura "
        "institucional y no solo como un régimen de prohibiciones."
    ),
    "12.4. Unidades Académicas y autoridades de gestión académica": (
        "Facultades, escuelas, institutos, decanatos y direcciones de carrera son el lugar "
        "donde el Marco se vuelve práctica cotidiana. Les corresponde aplicar estas reglas, "
        "definir criterios específicos por asignatura o evaluación —sin ser más permisivos que "
        "el texto general— y cuidar que las decisiones académicas de impacto no se deleguen "
        "a un sistema. Las viñetas detallan ese deber de aplicación, transparencia e información "
        "al Observatorio, para que la gobernanza no quede solo en el Rectorado."
    ),
    "12.5. Áreas de Tecnologías de la Información y sistemas": (
        "Las áreas de TIC no gobiernan la IA por sí solas, pero sin ellas no hay seguridad, "
        "interoperabilidad ni resguardo de datos. Su función es evaluar aspectos técnicos, "
        "revisar proveedores, implementar controles y alertar vulnerabilidades, en coordinación "
        "con el Observatorio y el Rectorado. Las enumeraciones precisan ese rol de soporte "
        "crítico: impedir integraciones no autorizadas, documentar configuraciones y contribuir "
        "a la puesta en marcha o al retiro de herramientas institucionales."
    ),
    "12.6. Secretarías y áreas de gestión": (
        "Investigación, extensión, administración, bienestar, educación a distancia y demás "
        "secretarías identifican procesos propios en los que ya interviene o puede intervenir "
        "la IA. Deben aplicar la matriz de delegación, asegurar un responsable humano en "
        "decisiones que afecten personas y documentar usos relevantes para el mapeo institucional. "
        "Las viñetas evitan que la gobernanza se concentre solo en lo académico y deje sin "
        "criterio a la gestión cotidiana de becas, trámites y sistemas de información."
    ),
    "12.7. Comunidad universitaria": (
        "Docentes, estudiantes, investigadores y personal no docente no son meros destinatarios "
        "pasivos: tienen deberes de uso conforme al Marco, de declaración cuando corresponda "
        "y de no cargar datos sensibles en herramientas no autorizadas. También deben reportar "
        "incidentes o sesgos por los canales institucionales. Las enumeraciones recuerdan que "
        "la gobernanza se sostiene en prácticas cotidianas, no solo en órganos de gobierno, y "
        "que el desconocimiento no exime de esos deberes mínimos."
    ),
    "14.1. Decisiones o tareas no delegables": (
        "Hay actos que, por su peso sobre trayectorias, derechos o la formación integral, no "
        "pueden transferirse a un sistema de IA. La evaluación final acreditante, el juicio "
        "pedagógico sustantivo, las sanciones y las decisiones laborales de impacto exigen "
        "una persona competente que decida y responda. Las viñetas identifican esos núcleos "
        "no delegables: la IA podrá, a lo sumo, aportar insumos, nunca sustituir el acto "
        "humano ni diluir la imputación institucional."
    ),
    "14.2. Delegación condicionada": (
        "Entre lo prohibido y lo rutinario existe un tramo amplio de asistencia: corrección "
        "orientativa, análisis de patrones, sugerencias de mejora o priorizaciones preliminares. "
        "Esa delegación solo es admisible con supervisión efectiva, transparencia, trazabilidad "
        "y posibilidad de revisión cuando haya efectos sobre personas. Las enumeraciones "
        "ilustran supuestos típicos; no habilitan, por sí solas, a automatizar la decisión "
        "final ni a prescindir del responsable humano identificable."
    ),
    "14.3. Tareas delegables": (
        "Ciertas tareas de bajo riesgo relativo —análisis descriptivo de datos agregados, "
        "automatizaciones administrativas de soporte o resúmenes de información no confidencial— "
        "pueden delegarse con controles proporcionales. Ello no significa ausencia de deberes: "
        "sigue vigente la seguridad de la información y la coherencia con este Marco. Las "
        "viñetas marcan el umbral de lo delegable: si la tarea determina derechos, evaluaciones "
        "o sanciones, deja de ser de bajo riesgo y vuelve al régimen condicionado o prohibido."
    ),
    "15.2. Contenidos mínimos": (
        "El mapa institucional solo cumple su función si visibiliza, con un mínimo común, "
        "dónde la IA ya interviene o puede intervenir. Debe incluir decisiones críticas, "
        "unidades involucradas, tipo de impacto, nivel de riesgo y zonas grises, sin duplicar "
        "los campos de la Matriz. Las enumeraciones fijan ese contenido mínimo para que el "
        "instrumento no se reduzca a un inventario informal ni omita prácticas “de hecho” "
        "que generan riesgo académico, ético o legal."
    ),
    "15.3. Decisiones críticas de referencia": (
        "No toda decisión universitaria tiene el mismo peso. Aprobar o desaprobar, acreditar, "
        "asignar becas, alertar riesgo de abandono, priorizar o evaluar desempeño son actos "
        "que pueden alterar trayectorias y por eso se reputan críticas. Las viñetas ofrecen "
        "un listado de referencia, no taxativo, para orientar el mapeo y la evaluación de "
        "riesgo. Si un sistema asiste esas decisiones, la supervisión humana y el registro "
        "dejan de ser opcionales."
    ),
    "15.7. Decisión de autorización": (
        "La autorización de un sistema institucional no es un trámite único: escala según el "
        "riesgo. El bajo puede resolverse en el área con registro; el medio exige TIC y "
        "opinión del Observatorio; el alto requiere Rectorado, dictamen y controles reforzados. "
        "Las viñetas concretan esa gradación para impedir tanto la parálisis burocrática en "
        "usos inocuos como el despliegue de hecho de sistemas que afectan personas o datos "
        "sensibles sin autoridad competente."
    ),
    "16.2. Criterios de valoración": (
        "Antes de adoptar o ampliar un sistema, la evaluación de riesgo no puede reducirse "
        "a un sí o no intuitivo. Debe ponderar severidad del daño, escala de afectados, "
        "opacidad, sensibilidad de datos, sesgos, supervisión humana, trazabilidad, "
        "dependencia de proveedores y efectos sobre menores o personas vulnerables. Las "
        "viñetas ordenan esos factores para que la autorización se funde en un juicio "
        "documentado, proporcional y comparable entre unidades y sedes."
    ),
    "16.3. Supuestos de evaluación reforzada": (
        "Algunos usos merecen un escrutinio más intenso porque tocan perfiles de personas, "
        "datos sensibles, biométricos o de menores, monitoreo intensivo o trayectorias "
        "académicas y laborales. También lo merecen las transferencias internacionales y "
        "las infraestructuras de terceros con acceso potencial a la información. Las "
        "enumeraciones marcan cuándo no basta una evaluación ordinaria: hay que reforzar "
        "controles, dictámenes y, si corresponde, la vía de autorización del Rectorado."
    ),
    "17.3. Decisión de autorización": (
        "El ciclo de vida de un sistema institucional culmina, en su tramo inicial, en una "
        "decisión de autorización graduada por riesgo. El área puede autorizar lo de bajo "
        "impacto con registro; lo medio exige conformidad técnica y opinión del Observatorio; "
        "lo alto queda reservado al Rectorado. Esta repetición operativa respecto del mapa "
        "no es redundancia vacía: asegura que ningún sistema entre en producción por vía "
        "de hecho, sin responsable ni posibilidad ulterior de suspensión o baja."
    ),
    "Artículo 22°. Finalidades legítimas y usos permitidos": (
        "No todo lo técnicamente posible es admisible, pero tampoco todo uso de IA está "
        "vedado. Hay finalidades legítimas —mejorar enseñanza y tutoría, fortalecer "
        "investigación y transferencia, apoyar la gestión sin sustituir decisiones de "
        "impacto— que este artículo reconoce para dar certeza a la comunidad. Las viñetas "
        "ilustran esas finalidades y, en los numerales siguientes, su traducción por ámbito. "
        "Permitido no significa irrestricto: sigue vigente la integridad, la privacidad y "
        "la supervisión humana."
    ),
    "22.1. En docencia y aprendizaje": (
        "En el aula la IA puede apoyar la preparación de clases, la explicación de conceptos "
        "y la práctica formativa, siempre que no sustituya la evidencia de aprendizaje ni "
        "el juicio docente. El criterio es pedagógico: enriquecer comprensión y diversidad "
        "de estrategias, no simular competencias no adquiridas. Las enumeraciones precisan "
        "usos típicos admisibles; las consignas de cátedra podrán ser más restrictivas, "
        "nunca más permisivas que este Marco, y deberán comunicarse con claridad."
    ),
    "22.2. En investigación y producción de conocimiento": (
        "En investigación la IA puede auxiliar búsquedas, estilo de textos propios o análisis "
        "exploratorios, y ser ella misma objeto de estudio. El límite es nítido: no fabrica "
        "datos ni reemplaza la autoría intelectual ni la validación metodológica. Las viñetas "
        "delimitan esa asistencia lícita para que el rigor científico no se confunda con "
        "aceleración acrítica. Toda incidencia relevante deberá poder declararse y defenderse "
        "ante pares, comités o financiadores."
    ),
    "22.3. En gestión y soporte institucional": (
        "En la gestión, la IA puede aliviar tareas repetitivas de bajo impacto, organizar "
        "información no sensible y apoyar comunicaciones no decisionales con revisión humana. "
        "No habilita a automatizar becas, sanciones, evaluaciones de desempeño o exclusiones. "
        "Las enumeraciones marcan el perímetro de lo permitido en soporte: analítica descriptiva "
        "agregada y asistencia operativa, siempre con responsable de área y sin perfilamiento "
        "lesivo de estudiantes o trabajadores."
    ),
    "23.1. Supuestos de condicionamiento": (
        "Hay usos que no están prohibidos de antemano, pero tampoco son libres: dependen de "
        "habilitación explícita, declaración, evaluación de riesgo o controles reforzados. "
        "Ese tramo condicionado cubre, entre otros, asistencia sustancial en trabajos evaluables, "
        "sistemas institucionales de impacto medio y tratamientos de datos que exigen base "
        "jurídica clara. Las viñetas identifican supuestos típicos para que la “zona gris” "
        "no se interprete como autorización implícita."
    ),
    "24.1. Prohibiciones en clave de integridad académica y científica": (
        "La integridad no admite presentar como propio un trabajo sustancialmente generado "
        "por IA sin elaboración ni declaración, ni falsificar datos, citas o evidencias, ni "
        "eludir evaluaciones. Estas prohibiciones protegen el sentido del título universitario "
        "y de la publicación científica. Las viñetas concretan conductas inadmisibles; su "
        "gravedad se ponderará en cada procedimiento, pero el criterio de fondo es claro: "
        "no se simula lo que no se ha aprendido ni se inventa lo que no se ha investigado."
    ),
    "24.2. Prohibiciones en clave de derechos de las personas": (
        "Queda vedado usar IA para vigilar, perfilar o discriminar de modo ilegítimo, tratar "
        "datos sensibles en herramientas no autorizadas o adoptar decisiones de alto impacto "
        "sin intervención humana. El centro de estas prohibiciones es la dignidad y la "
        "privacidad de quienes estudian, enseñan o trabajan en la Universidad. Las enumeraciones "
        "no agotan el catálogo de ilicitudes: ilustran el umbral a partir del cual un uso "
        "deja de ser innovación y se convierte en lesión de derechos."
    ),
    "24.3. Prohibiciones en clave institucional y de seguridad": (
        "En clave institucional está prohibido desplegar sistemas de riesgo medio o alto sin "
        "el ciclo de autorización, desactivar controles de seguridad o supervisión, o extraer "
        "información masiva hacia plataformas externas no habilitadas. Estas reglas protegen "
        "la continuidad de la Universidad, sus datos y su reputación. Las viñetas precisan "
        "conductas que, aun sin dolo académico, comprometen la soberanía de la información "
        "y la trazabilidad de las decisiones."
    ),
    "Artículo 25°. Limitaciones transversales y proporcionalidad": (
        "Además de las prohibiciones taxativas, rigen límites que atraviesan todos los usos: "
        "proporcionalidad entre finalidad y riesgo, minimización de datos, transparencia "
        "compatible con la seguridad y preservación de la equidad. Este artículo evita que "
        "un uso “permitido” se desborde por intensidad, escala o opacidad. Las enumeraciones "
        "traducen esa proporcionalidad en restricciones concretas, de modo que la excepción "
        "no se vuelva regla ni la eficiencia justifique un control desmesurado de las personas."
    ),
    "Artículo 26°. Autorización excepcional": (
        "Solo en supuestos fundados, con evaluación reforzada y sin contravenir las "
        "prohibiciones del Capítulo IV, podrá habilitarse un uso excepcional. Esa vía es "
        "restrictiva: no crea precedente automático ni convalida lo expresamente vedado. "
        "Las reglas que siguen precisan condiciones, autoridad competente y deber de "
        "documentación, para que la excepción no se convierta en atajo informal frente "
        "al ciclo ordinario de autorización y registro."
    ),
    "Artículo 27°. Deber de adecuación de prácticas existentes": (
        "Los usos y sistemas ya en marcha al entrar en vigencia no quedan al margen del "
        "Marco. Las unidades deben identificarlos, declararlos, adecuarlos a las categorías "
        "de este Capítulo y suspender de inmediato lo manifiestamente prohibido. Las viñetas "
        "ordenan ese deber transitorio para evitar una doble vara: lo nuevo se autoriza y "
        "lo viejo se tolera. La adecuación es condición de continuidad, no un trámite "
        "optativo de las cátedras o áreas de gestión."
    ),
    "28.1. Niveles preuniversitarios": (
        "En colegios y escuelas dependientes el criterio es más restrictivo que en el grado. "
        "El uso en aprendizajes exige acompañamiento docente o de la persona adulta responsable; "
        "no se cargan datos de menores en herramientas no autorizadas; y, en la duda, prima "
        "la restricción. Las viñetas concretan esas salvaguardas y remiten al Anexo J para "
        "el circuito con representantes legales. La IA no reemplaza el cuidado psicopedagógico, "
        "pastoral ni de bienestar de niños, niñas y adolescentes."
    ),
    "29.2. Deberes del docente": (
        "El docente conserva el juicio pedagógico: verifica exactitud y pertinencia, no carga "
        "datos de estudiantes en plataformas no autorizadas, informa las reglas de la asignatura "
        "y no delega la acreditación. También modela alfabetización crítica y se abstiene de "
        "volcar trabajos inéditos o materiales institucionales a herramientas comerciales "
        "inadecuadas. Las enumeraciones traducen ese deber de cuidado, para que la IA en el "
        "aula sea asistencia revisada y no sustituto de la enseñanza."
    ),
    "29.5. Usos estudiantiles generalmente admisibles": (
        "Salvo regla más estricta de la cátedra, el estudiante puede pedir explicaciones, "
        "organizar ideas, practicar en actividades no evaluativas y mejorar claridad de un "
        "texto propio, siempre que el resultado evidencie aprendizaje real. La admisibilidad "
        "general no borra la autoría ni habilita a entregar un producto sustancialmente "
        "generado por IA. Las viñetas delimitan ese perímetro de apoyo formativo, distinto "
        "del tramo condicionado o prohibido que se regula a continuación."
    ),
    "29.6. Usos estudiantiles condicionados": (
        "Borradores de trabajos evaluables, asistencia sustancial en TFI o tesis, generación "
        "de código o análisis que formen parte del objeto de evaluación, y usos colaborativos "
        "que nuben la atribución individual, requieren habilitación explícita y declaración "
        "cuando corresponda. Sin esa habilitación, el uso condicionado se interpreta como "
        "no autorizado. Las enumeraciones buscan reducir la zona gris en la que el estudiante "
        "“no sabía” si podía o no servirse de la herramienta."
    ),
    "29.7. Usos estudiantiles no permitidos": (
        "Está prohibido entregar como propio un trabajo sustancialmente generado por IA sin "
        "elaboración ni declaración, usarla en exámenes vedados, fabricar citas o datos, "
        "suplantar identidad o compartir datos de terceros y de pacientes. Estas conductas "
        "no son “malas prácticas menores”: afectan la equidad evaluativa y, a veces, la "
        "confidencialidad. Las viñetas las nombran con claridad para que la consecuencia "
        "académica o disciplinaria no sorprenda a quien las incurra."
    ),
    "30.3. TFI, tesis y trabajos finales": (
        "En TFI y tesis la autoría intelectual y la responsabilidad científica pertenecen al "
        "estudiante, con su dirección. Si se admite IA, debe declararse de modo específico; "
        "no puede reemplazar marco teórico, análisis crítico ni conclusiones. Directores y "
        "tribunales podrán exigir defensas orales o evidencias de proceso. Las enumeraciones "
        "refuerzan que la fabricación de datos o fuentes es falta grave, y que el título "
        "final acredita competencias humanas, no destreza en el uso de un modelo generativo."
    ),
    "30.5. Matriz orientativa de niveles de uso según riesgo académico": (
        "Para orientar consignas, se distingue un nivel de uso libre e incentivado en tareas "
        "de bajo riesgo, un uso guiado y declarado en riesgos medios, y un uso restringido "
        "o prohibido en evaluaciones sincrónicas o de desempeño auténtico. La matriz no "
        "reemplaza las prohibiciones generales: las traduce operativamente. Las viñetas "
        "describen esos niveles para que cátedras y carreras expliciten el grado de asistencia "
        "admitido, en lugar de dejarlo a interpretaciones implícitas."
    ),
    "30.7. Declaración de uso": (
        "Cuando sea exigible, la declaración no es un trámite vacío: identifica herramienta, "
        "periodo, finalidad, alcance de la asistencia y asunción de responsabilidad por la "
        "verificación humana. Puede incluir prompts principales y las validaciones realizadas. "
        "Las enumeraciones fijan ese contenido mínimo para que quien evalúa pueda distinguir "
        "aprendizaje genuino de mera mediación instrumental, y para que el Observatorio "
        "pueda proponer un modelo institucional homogéneo."
    ),
    "31.2. Orientaciones para la evaluación": (
        "La disponibilidad de IA obliga a evaluar procesos, no solo productos finales: defensas "
        "orales, problemas situados, portafolios, crítica de salidas generativas e instancias "
        "de desempeño auténtico. El rediseño no es capricho pedagógico; es condición de "
        "validez y equidad. Las viñetas recogen orientaciones según la naturaleza de la "
        "asignatura, para que la evaluación siga midiendo lo que dice medir y no la habilidad "
        "de ocultar el uso de una herramienta."
    ),
    "31.3. Lo que no se delega": (
        "Aunque la IA apoye la retroalimentación formativa, no se delega la decisión final "
        "de aprobación o desaprobación, la ponderación última del mérito, la apreciación de "
        "la formación integral cuando integre la evaluación, ni la determinación de sanciones "
        "académicas. Las enumeraciones recuerdan el núcleo no automatizable del oficio "
        "docente: calificar es un acto de juicio responsable, no la confirmación de un "
        "puntaje opaco producido por un sistema."
    ),
    "34.1. Usos admisibles": (
        "En investigación se admite, con control metodológico y declaración cuando corresponda, "
        "el apoyo bibliográfico verificado, la asistencia de estilo sobre texto propio, la "
        "exploración de hipótesis bajo validación científica y el desarrollo de modelos o "
        "código con trazabilidad. También es lícito investigar la IA misma. Las viñetas "
        "delimitan esa asistencia para que la aceleración no se confunda con rigor: cada "
        "resultado asistido debe poder explicarse, reproducirse en lo pertinente y defenderse."
    ),
    "34.2. Deberes del investigador responsable": (
        "Quien dirige o ejecuta una investigación con IA define finalidad y alcance de la "
        "asistencia, verifica calidad y límites, documenta el uso relevante y asegura que "
        "la interpretación final sea humana. Debe evaluar riesgos éticos y sociales antes "
        "de desplegar sistemas con impacto sobre personas. Las enumeraciones concretan ese "
        "deber de integridad: no se fabrican datos ni se ocultan limitaciones metodológicas "
        "derivadas del uso de la herramienta."
    ),
    "34.3. Usos condicionados o de alto escrutinio": (
        "Proyectos con datos personales o de salud, sistemas predictivos sobre personas, "
        "experimentos de impacto clínico o social, entrenamiento con datos institucionales "
        "y estudios en comunidades vulnerables exigen evaluación reforzada y, en su caso, "
        "comité de ética, Observatorio o autoridad competente. Las viñetas marcan ese umbral "
        "de alto escrutinio: la celeridad académica no justifica omitir salvaguardas ni "
        "tratar el consentimiento como una cláusula de estilo."
    ),
    "34.5. Autoría": (
        "Solo las personas humanas que hayan realizado un aporte intelectual significativo "
        "pueden ostentar autoría. Los sistemas de IA no se firman como coautores; el uso "
        "relevante se reconoce según las normas de la revista o repositorio y con transparencia "
        "institucional mínima. La dirección de tesis no se delega a una herramienta. Las "
        "viñetas fijan esa regla para que la firma científica siga imputando responsabilidad, "
        "mérito y, cuando corresponda, falta a personas identificables."
    ),
    "35.3. Datos institucionales y de terceros": (
        "Los datos de la Universidad no se cargan en herramientas externas sin autorización "
        "y evaluación de riesgo. Los obtenidos de empresas, escuelas, centros de salud u "
        "organismos se usan solo dentro del convenio o licencia. Reutilizarlos para entrenar "
        "modelos exige base jurídica, finalidad compatible y anonimización o seudonimización "
        "cuando proceda. Las enumeraciones protegen tanto el patrimonio informacional de la "
        "institución como los derechos de terceros que confiaron esos datos."
    ),
    "35.5. Comités de ética y evaluación de riesgo": (
        "Cuando un proyecto con IA pueda afectar derechos, salud, privacidad, equidad o "
        "seguridad, no basta el criterio del equipo: corresponde evaluación del comité de "
        "ética competente, consulta al Observatorio, revisión técnica si hay despliegue y "
        "evaluación de riesgo conforme al Capítulo III. Las viñetas ordenan esa vía de "
        "salvaguarda. Ningún calendario de publicación o de convocatoria justifica omitirla "
        "ni tratarla como un obstáculo meramente administrativo."
    ),
    "36.2. Usos admisibles": (
        "La extensión puede incorporar IA para alfabetizar comunidades, asistir a instituciones "
        "públicas o sociales en una adopción responsable, impulsar innovación regional y abrir "
        "espacios de deliberación ética. Esos usos se fomentan porque sirven al medio, no "
        "porque “probar tecnología” sea un fin en sí. Las viñetas ilustran intervenciones "
        "admisibles; las limitaciones del numeral siguiente impiden convertir a la comunidad "
        "en laboratorio de sistemas opacos o de alto riesgo sin acompañamiento."
    ),
    "36.3. Limitaciones": (
        "No se admiten intervenciones de extensión que introduzcan sistemas de alto riesgo "
        "sin evaluación, generen dependencia tecnológica opaca, usen datos comunitarios sin "
        "consentimiento ni resguardo, o presenten la IA como solución automática a problemas "
        "sociales, educativos o sanitarios complejos. Las enumeraciones fijan ese límite de "
        "prudencia: el servicio al medio no transfiere a terceros riesgos que la Universidad "
        "no aceptaría para su propia comunidad."
    ),
    "Artículo 40°. Finalidades admisibles y sistemas de gestión": (
        "En la gestión universitaria la IA puede mejorar eficiencia e información, siempre "
        "que no sustituya decisiones de alto impacto sobre personas. Las finalidades admisibles "
        "se leen junto con los sistemas concretos —analítica, alertas, automatizaciones de "
        "soporte— y con la exigencia de responsable humano. Las viñetas y numerales posteriores "
        "delimitan qué cabe en la gestión ordinaria y qué exige el ciclo de autorización, "
        "para que la “modernización” no se confunda con delegación indebida."
    ),
    "40.3. Becas, priorizaciones y asignaciones": (
        "Becas, priorizaciones y asignaciones de beneficios tocan derechos y trayectorias; "
        "por eso la IA, si interviene, solo puede asistir con criterios transparentes, datos "
        "pertinentes y revisión humana efectiva. No se admite un ranking opaco que excluya "
        "o estigmatice. Las reglas que siguen precisan esa cautela: la herramienta puede "
        "ordenar información, no decidir por sí sola quién accede a un beneficio institucional "
        "ni eludir el deber de explicación frente a la persona afectada."
    ),
    "41.2. Prohibición de automatización plena": (
        "Nadie puede ser objeto de una decisión plenamente automatizada de alto impacto en "
        "gestión universitaria. Debe existir intervención humana significativa, posibilidad "
        "de no ser clasificado de modo opaco y garantía de que la IA opere como apoyo. "
        "Las enumeraciones reiteran, en el ámbito administrativo, el principio rector del "
        "Capítulo II: el acto institucional es humano. “Lo decidió el sistema” no constituye "
        "fundamento ni eximente."
    ),
    "42.1. Condiciones generales": (
        "Todo convenio, contrato o proyecto con terceros que involucre desarrollo, provisión "
        "o uso de IA debe contemplar finalidad, responsable humano, régimen de datos, "
        "propiedad intelectual, niveles de riesgo y salvaguardas de discontinuidad. Esas "
        "condiciones generales unifican lo que antes podía quedar disperso entre compras, "
        "convenios y áreas usuarias. Las viñetas marcan el piso contractual: sin él, no hay "
        "contratación alineada con este Marco."
    ),
    "42.5. Requisitos previos a la contratación": (
        "Antes de contratar se verifica necesidad, proporcionalidad, evaluación de riesgo, "
        "opinión técnica y, cuando el impacto lo amerite, intervención del Observatorio y "
        "autorización del nivel competente. Comprar primero y gobernar después es precisamente "
        "lo que este artículo impide. Las enumeraciones ordenan esos requisitos previos para "
        "que la urgencia operativa, la oferta de un proveedor o la presión de plazos de un "
        "convenio no se invoquen como excepción permanente al ciclo de autorización."
    ),
    "42.6. Cláusulas mínimas deseables": (
        "Los contratos deberían incluir, al menos, limitaciones de finalidad, prohibición de "
        "usos secundarios no autorizados, deberes de notificación de incidentes, facultades "
        "de auditoría razonable, portabilidad y condiciones de salida. Las viñetas no pretenden "
        "un modelo único de pliego, pero sí un estándar mínimo de soberanía institucional. "
        "Sin esas cláusulas, la Universidad queda expuesta a dependencias opacas y a la "
        "pérdida de control sobre sus propios datos."
    ),
    "Artículo 43°. Registro, seguridad y resguardo de la información": (
        "Todo sistema institucional de IA debe poder inscribirse, asegurarse y, llegado el "
        "caso, discontinuarse sin dejar a la Universidad sin sus datos ni sin rastro de "
        "decisiones. Registro, controles técnicos y resguardo no son formalidades: son la "
        "condición de una gobernanza verificable. Las disposiciones siguientes precisan "
        "qué se registra, qué controles mínimos rigen y qué prácticas de seguridad están "
        "prohibidas, en coordinación con las áreas de TIC."
    ),
    "43.2. Controles mínimos": (
        "Según el riesgo, los controles incluyen gestión de accesos, registro de eventos "
        "relevantes, copias de resguardo, separación de ambientes y revisión periódica. "
        "No se exige el mismo rigor a un asistente de bajo impacto que a un sistema que "
        "trata legajos o alertas sobre personas. Las viñetas fijan ese mínimo escalonado "
        "para que “seguridad” no quede en un enunciado genérico ni se reduzca a la "
        "configuración por defecto de un proveedor."
    ),
    "43.4. Prohibiciones de seguridad": (
        "Queda prohibido compartir credenciales, conectar herramientas no autorizadas a "
        "bases institucionales, extraer información masiva hacia plataformas externas o "
        "eludir controles de auditoría. Estas conductas, aun cuando se presenten como "
        "atajos de trabajo, exponen datos y rompen la trazabilidad. Las enumeraciones las "
        "nombran para que el personal de gestión y los terceros sepan que la infracción "
        "de seguridad no es un detalle técnico menor."
    ),
    "44.2. Transferencias y almacenamiento": (
        "Si datos institucionales o personales deben tratarse fuera del entorno controlado "
        "de la Universidad, se exige base jurídica y contractual, nivel de protección del "
        "destino y evaluación de los riesgos de acceso por terceros. La transferencia no "
        "es un acto neutro de “subir a la nube”. Las viñetas recuerdan que la localización, "
        "la ley aplicable y las garantías del encargado forman parte de la soberanía de "
        "datos que este Capítulo procura preservar."
    ),
    "Artículo 45°. Roles, prohibiciones e implementación gradual": (
        "La implementación no recae en un único actor: Rectorado, áreas dueñas del proceso, "
        "TIC y Observatorio se reparten autorización, uso cotidiano, seguridad y asesoramiento "
        "ético. Esa distribución evita tanto el cuello de botella como el vacío de imputación. "
        "Las viñetas asignan roles mínimos y, junto con las prohibiciones e hitos de "
        "implementación gradual, ordenan cómo se pasa de la norma a la práctica sin "
        "improvisar sistemas de alto impacto."
    ),
    "55.2. Impugnación": (
        "La persona afectada puede impugnar o pedir reconsideración cuando la decisión se "
        "haya apoyado indebidamente en una salida algorítmica, existan indicios de error o "
        "sesgo, falte información debida o se hayan vulnerado privacidad, equidad o integridad. "
        "El trámite sigue las vías académicas, administrativas o disciplinarias ya existentes. "
        "Las enumeraciones concretan causales para que el derecho a revisión humana no quede "
        "en un enunciado sin procedimiento."
    ),
    "57.1. Derechos en materia de datos": (
        "Quien vea sus datos tratados mediante IA tiene derecho a finalidad legítima, "
        "minimización, seguridad, corrección, oposición a entrenamientos no autorizados y, "
        "cuando proceda, acceso, portabilidad o limitación del tratamiento. Esos derechos "
        "no nacen de este Marco en abstracto: se leen con la normativa de protección de "
        "datos vigente. Las viñetas los enuncian para que la comunidad sepa qué puede "
        "exigir y la Universidad, qué debe garantizar."
    ),
    "57.2. Deberes correlativos de la Universidad": (
        "Correlativamente, la Universidad minimiza la carga de datos en herramientas de IA, "
        "prefiere agregados o seudonimización, restringe accesos por rol, evita plataformas "
        "no autorizadas para legajos o salud, documenta bases de tratamiento cuando el riesgo "
        "lo amerite e incorpora privacidad desde el diseño. Las enumeraciones traducen "
        "derechos en deberes organizativos: sin ellos, las garantías del numeral anterior "
        "serían inexigibles en la práctica cotidiana."
    ),
    "57.3. Confidencialidad reforzada": (
        "Ciertas categorías —salud y discapacidad, legajos disciplinarios, datos de becas, "
        "acompañamiento psicológico o pastoral, información de menores, secretos profesionales "
        "y, cuando intervengan, datos genéticos, biométricos o bancarios— exigen un umbral "
        "más alto de resguardo. No son “un dato más” para alimentar un modelo. Las viñetas "
        "identifican ese núcleo sensible para que la excepción de confidencialidad reforzada "
        "no dependa del criterio aislado de cada área."
    ),
    "58.2. Deberes de prevención": (
        "En sistemas de riesgo medio o alto se evalúan sesgos previsibles antes del despliegue, "
        "se monitorean impactos diferenciales, se corrigen diseños excluyentes y se evita "
        "estigmatizar a quienes aparecen “alertados”. Las alertas deben derivar a acompañamiento, "
        "no a prejuicio. Las enumeraciones convierten la no discriminación en prácticas "
        "verificables, para que el principio no se agote en una declaración de intenciones "
        "sin evidencia de control."
    ),
    "59.1. Hechos reportables": (
        "Pueden reportarse usos prohibidos o de alto riesgo no autorizados, sospechas fundadas "
        "de sesgo o exclusión algorítmica, filtraciones, decisiones de impacto sin supervisión "
        "humana, represalias contra alertantes e incumplimientos de proveedores. El listado "
        "no es cerrado: orienta qué merece canal institucional. Las viñetas buscan que el "
        "reporte deje de percibirse como delación informal y pase a ser un mecanismo de "
        "protección y de aprendizaje."
    ),
    "59.3. Vía preferente": (
        "Los conflictos se resuelven por las vías ya existentes: académico-pedagógica en el "
        "aula, de ética en investigación, administrativa en becas y trámites, disciplinaria "
        "o laboral cuando corresponda, y técnica de seguridad en incidentes informáticos. "
        "Este Marco no crea un fuero paralelo. Las enumeraciones evitan que toda controversia "
        "sobre IA se desvíe al Observatorio como si este sustituyera a Decanos, comités o "
        "áreas de personal."
    ),
    "Artículo 60°. Protección de denunciantes e incidentes": (
        "Quien reporte de buena fe no debe sufrir represalias; quien resulte reportado conserva "
        "defensa y descargo; se evita la difusión prematura; y la confidencialidad no puede "
        "encubrir inacción ante riesgos graves. Ese equilibrio protege tanto la alerta útil "
        "como el debido proceso. Las viñetas fijan esas garantías para que el canal de "
        "incidentes sea creíble, y para que el registro agregado de casos sirva a la mejora "
        "institucional sin exponer personas de más."
    ),
    "63.4. Formación obligatoria": (
        "El Rectorado, a propuesta del Observatorio, puede establecer instancias obligatorias "
        "para docentes de nuevo ingreso o en renovación de funciones, responsables de sistemas "
        "institucionales y otros roles de alto contacto con IA. La obligatoriedad no es un "
        "castigo: es condición de un uso informado. Las viñetas identifican destinatarios "
        "prioritarios, de modo que la alfabetización no quede librada al voluntarismo de "
        "quien ya está más familiarizado con las herramientas."
    ),
    "Artículo 64°. Formación docente, estudiantil y de gestión": (
        "La capacitación no es genérica: en docencia prioriza consignas transparentes e "
        "integración pedagógica; en estudiantes, autoría e integridad; en gestión, decisiones "
        "humano-algorítmicas, no discriminación y seguridad de la información. Formar por "
        "niveles evita tanto el subdiagnóstico como el curso único inútil. Las enumeraciones "
        "orientan contenidos mínimos para que cada público reciba lo que su rol exige, y "
        "para que la formación no se diluya en una oferta ocasional sin responsables ni seguimiento."
    ),
    "Artículo 65°. Buenas prácticas, guías y comunidades de práctica": (
        "El Observatorio mantiene un banco de experiencias verificadas en docencia, investigación, "
        "gestión y extensión, y promueve comunidades de práctica para que el aprendizaje no "
        "quede en un repositorio muerto. Las guías operativas se actualizan con esos casos. "
        "Las viñetas precisan ámbitos del banco, a fin de que las “buenas prácticas” no se "
        "confundan con anécdotas informales ni sustituyan el articulado cuando haya conflicto."
    ),
    "66.2. Acciones mínimas de difusión": (
        "De poco sirve un Marco que no llega a Decanatos, secretarías y aulas. Se procurará "
        "publicación en sitios institucionales, comunicación a autoridades de carrera y "
        "piezas de difusión periódica. Las enumeraciones fijan un piso de visibilidad, "
        "complementario de la difusión obligatoria del Capítulo XII, para que nadie pueda "
        "alegar con razón que las reglas “no se conocían” por haber quedado en un archivo "
        "de OneDrive."
    ),
    "66.3. Lenguaje accesible": (
        "Las piezas de difusión deben ser claras, con ejemplos y distinción nítida entre lo "
        "permitido, lo condicionado y lo prohibido, sin perder rigor. Observatorio, Rectorado "
        "y unidades se reparten diseño, aprobación y aplicación del mensaje. Las viñetas "
        "asignan esos roles de comunicación para que el lenguaje accesible no diluya la norma "
        "ni la deje en un tecnicismo incomprensible para estudiantes y personal de apoyo."
    ),
    "Artículo 68°. Responsables del seguimiento": (
        "El seguimiento no es privativo del Observatorio: este coordina e informa; el Rectorado "
        "recibe y dispone; las unidades proveen datos y ejecutan adecuaciones; TIC aporta "
        "evidencia de seguridad e incidentes. Esa red evita tanto el reporte cosmético como "
        "la sobrecarga de un solo organismo. Las enumeraciones asignan esos roles para que "
        "la mejora continua tenga dueños identificables y plazos, no meras intenciones."
    ),
    "69.2. Indicadores orientativos de cumplimiento": (
        "Pueden considerarse, a título ejemplificativo, el porcentaje de unidades con política "
        "de IA comunicada, el número de sistemas registrados por riesgo, la cobertura formativa "
        "y la existencia de canales de consulta. No son metas rígidas ni ranking punitivo. "
        "Las viñetas ilustran evidencias de que el Marco se está aplicando, con flexibilidad "
        "para adaptar indicadores a la escala de cada sede y función."
    ),
    "69.3. Indicadores orientativos de impacto": (
        "El impacto se mira, entre otros, en la evolución de conflictos de integridad, la "
        "percepción de claridad de las reglas, la reducción de incidentes de datos y la "
        "calidad de las autorizaciones de alto riesgo. El objetivo no es producir un tablero "
        "decorativo, sino aprender. Las enumeraciones orientan esa lectura cualitativa y "
        "cuantitativa, sin convertir cada cifra en un juicio automático sobre personas o "
        "unidades."
    ),
    "Artículo 70°. Evidencia e informes institucionales": (
        "El seguimiento se nutre del mapa, del registro de sistemas, de consultas, incidentes, "
        "auditorías y de los informes que el Observatorio eleva al Rectorado y, cuando "
        "corresponda, al Consejo Superior. Sin esa evidencia, la “mejora continua” es una "
        "fórmula vacía. Las viñetas identifican fuentes mínimas de información para que los "
        "informes no dependan del recuerdo informal de unos pocos gestores."
    ),
    "71.1. Auditoría ética": (
        "La Universidad puede auditar éticamente sistemas, procesos o unidades para verificar "
        "dignidad, equidad, supervisión humana efectiva, proporcionalidad y coherencia con "
        "el ideario. No es una cacería de infracciones menores: es un control de sentido "
        "institucional. Las enumeraciones marcan ejes de esa auditoría, distinguibles de la "
        "revisión técnica de TIC, aunque ambas puedan complementarse cuando el sistema sea "
        "de riesgo medio o alto."
    ),
    "Artículo 72°. Adecuación, actualización del Marco y anexos": (
        "Cuando el seguimiento detecte brechas, el área responsable elabora un plan con "
        "hallazgos, medidas, plazos, indicadores de cierre y apoyo requerido. La actualización "
        "del Marco y de los anexos sigue vías distintas según sea reforma sustancial u "
        "operativa. Las viñetas y numerales posteriores ordenan esa mecánica para que el "
        "texto no envejezca inmóvil ni se altere su núcleo de principios por simple circular "
        "administrativa."
    ),
    "72.2. Reformas sustanciales y operativas": (
        "Los cambios de principios, prohibiciones centrales, derechos o arquitectura de "
        "gobernanza requieren el órgano que aprobó el Marco. En cambio, anexos, matrices, "
        "formularios y guías pueden actualizarse por Rectorado, con dictamen del Observatorio, "
        "si no contradicen el articulado. Esa distinción evita tanto la rigidización excesiva "
        "como la reforma encubierta por vía de anexo. Las viñetas precisan ambos tramos, "
        "para que cada actualización siga el cauce institucional que le corresponde."
    ),
    "Artículo 73°. Pilotos, transparencia y proporcionalidad del control": (
        "Los pilotos de innovación con IA son experimentación controlada: alcance limitado, "
        "duración definida, responsable humano, evaluación de riesgo y posibilidad de corte. "
        "No son una puerta trasera para saltarse el Capítulo III. Las enumeraciones fijan "
        "esas condiciones y, junto con las reglas de transparencia y proporcionalidad del "
        "control, impiden que el piloto se eternice o se expanda sin nueva autorización."
    ),
    "Artículo 75°. Remisión a normativas vigentes": (
        "En lo no regulado expresamente rigen Estatutos, plan estratégico, ordenanzas de "
        "docencia e investigación, normativa de datos, régimen disciplinario y laboral, y "
        "demás resoluciones aplicables. Este artículo evita el mito de que el Marco de IA "
        "deroga o reemplaza el resto del ordenamiento universitario. Las viñetas recuerdan "
        "ese encaje: la IA se gobierna aquí, pero las sanciones, los contratos y los derechos "
        "siguen anclados en las normas que ya existen."
    ),
    "76.2. Criterios de tipificación orientativa": (
        "Sin perjuicio de cada procedimiento, se consideran especialmente graves la falsificación "
        "de datos o citas, la autoría fraudulenta, el uso de IA para eludir evaluaciones, "
        "la carga de datos sensibles en herramientas no autorizadas y el despliegue clandestino "
        "de sistemas de alto riesgo. La tipificación es orientativa, no un código penal interno. "
        "Las viñetas guían la ponderación de gravedad para que la respuesta sea proporcional "
        "y previsible."
    ),
    "Artículo 77°. Interpretación y lagunas": (
        "El Marco se interpreta de forma armónica, privilegiando dignidad, integridad, "
        "transparencia, equidad y responsabilidad humana. Ante silencio, rige el criterio "
        "pro persona y de prevención del riesgo. El Observatorio puede orientar; deciden "
        "los órganos competentes. Una práctica social extendida no convalida lo no autorizado. "
        "Las viñetas cierran el método interpretativo para que la laguna no se llene con "
        "improvisación ni con analogías laxas."
    ),
    "Artículo 78°. Anexos y desarrollo reglamentario": (
        "Los anexos —glosario, matriz, declaración, evaluación de riesgo, guías, política de "
        "consignas, incidentes, registro y guía preuniversitaria— forman parte operativa del "
        "Marco una vez aprobados. No crean derechos ni deberes contradictorios con el "
        "articulado: si chocan, prevalece el cuerpo normativo. Las viñetas identifican esos "
        "instrumentos para que la implementación no dependa de documentos sueltos, ni de "
        "modelos improvisados ajenos a este texto."
    ),
    "79.1. Período de adecuación": (
        "Durante ciento ochenta días desde la vigencia, las unidades identifican usos y "
        "sistemas existentes, los declaran, adecuan prácticas y suspenden lo manifiestamente "
        "prohibido. Ese plazo no es una moratoria de las prohibiciones graves. Las enumeraciones "
        "ordenan las tareas del período transitorio para que la entrada en vigor no sea "
        "simbólica ni, inversamente, una trampa de incumplimiento instantáneo e inevitable "
        "para quien actuaba de buena fe."
    ),
    "Artículo 81°. Difusión obligatoria": (
        "El Marco debe llegar a Rectorado y vicerrectorados, sedes, facultades, escuelas, "
        "institutos, secretarías, docentes, estudiantes y personal de gestión, por canales "
        "institucionales idóneos. La difusión no es cortesía: es condición de exigibilidad "
        "razonable y de equidad entre unidades. Las viñetas enumeran destinatarios mínimos "
        "para que ninguna sede o facultad pueda considerarse al margen por no haber sido "
        "formalmente notificada del texto."
    ),
}


def wc(s: str) -> int:
    return len(s.split())


def set_text(p, text: str):
    if p.runs:
        p.runs[0].text = text
        for r in p.runs[1:]:
            r.text = ""
    else:
        p.text = text
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.line_spacing = 1.5


def main():
    short = [(k, wc(v)) for k, v in INTROS.items() if not (60 <= wc(v) <= 80)]
    if short:
        print("WORD COUNT OUT OF RANGE:")
        for k, n in short:
            print(f"  {n:3}  {k}")
        raise SystemExit(1)

    doc = Document(TARGET)
    start = False
    replaced = missing = 0
    for i, p in enumerate(doc.paragraphs):
        t = p.text.strip().split("\t")[0]
        if "VERSIÓN 2026" in t or "VERSIÓN 2.0 DEPURADA" in t:
            start = True
            continue
        if not start or p.style.name.startswith("toc"):
            continue
        if not t.startswith("Este apartado desarrolla"):
            continue
        head = ""
        for k in range(i - 1, max(i - 6, -1), -1):
            pk = doc.paragraphs[k]
            tt = pk.text.strip().split("\t")[0]
            if pk.style.name in ("Heading 2", "Heading 3"):
                head = tt
                break
        if head in INTROS:
            set_text(p, INTROS[head])
            replaced += 1
        else:
            missing += 1
            print("NO INTRO FOR", repr(head))

    leftover = [k for k in INTROS if True]
    doc.save(TARGET)
    print("replaced", replaced)
    print("missing_heads", missing)
    print("defined", len(INTROS))


if __name__ == "__main__":
    main()
